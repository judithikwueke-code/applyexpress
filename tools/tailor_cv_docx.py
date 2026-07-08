"""
tailor_cv_docx.py — Tailor a CV for a specific job and render it to the house format.

Pipeline:
  1. Parse structured data from the base CV (.docx)
  2. LLM-rewrite: summary, skills, certifications, top-2 role titles + bullets
  3. Build a brand-new .docx to the exact house formatting specification

House format:
  Font:    Arial throughout
  Colours: Accent #1F4E79 | Dark #222222 | Grey #666666 | LGrey #555555
  Page:    A4, margins top/bottom 1080 twips (0.75") left/right 1440 twips (1")
  Headings: ALL CAPS, bold, 12pt accent, bottom border accent, 240/80 twip spacing
  Roles:   Title (11pt bold) — Org (11pt) [right tab] Dates (10pt italic grey)
  Bullets: • indent 360/200 twips, 40 twips after
"""

import sys
import os
import re
import copy
import time
import argparse
import logging
from pathlib import Path

from dotenv import load_dotenv
load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s [tailor_cv_docx] %(message)s")
logger = logging.getLogger(__name__)

sys.path.insert(0, str(Path(__file__).parent.parent))
from tools.llm_client import call_llm

from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_CV_PATH = Path(os.getenv("CV_PATH", ""))
PROFILE_PATH = Path(os.getenv("CANDIDATE_PROFILE_PATH",
                               str(Path(__file__).parent.parent / "candidate_profile.md")))
TMP_DIR = Path(os.getenv("TMP_DIR", str(Path(__file__).parent.parent / ".tmp")))

# ── House colours ──────────────────────────────────────────────────────────────
ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # deep navy   — name, headings
DARK   = RGBColor(0x22, 0x22, 0x22)   # near-black  — body, roles, bullets
GREY   = RGBColor(0x66, 0x66, 0x66)   # mid-grey    — dates (italic)
LGREY  = RGBColor(0x55, 0x55, 0x55)   # light-grey  — contact line

W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML = "http://www.w3.org/XML/1998/namespace"


# ── Prompts ────────────────────────────────────────────────────────────────────

_SYS = (
    "You are a top 0.1% executive CV writer and headhunter with 20 years of experience "
    "placing candidates at FTSE 100 and Fortune 500 companies. You know exactly what makes "
    "a senior HR director stop scrolling. Your cardinal rule: every bullet must show evidence "
    "of impact, not description of duty. You never invent experience — you reframe existing "
    "experience using the language of the job description. "
    "Return ONLY the rewritten content with no labels, explanations, or markdown."
)

_SUMMARY_PROMPT = """Rewrite this candidate's Professional Summary to position them for the specific role below.
A senior HR director will spend 6 seconds on this — make every word count.

CURRENT SUMMARY:
{current_summary}

CANDIDATE BACKGROUND (do not invent beyond this):
{profile_snippet}

TARGET JOB:
Title: {title}
Company: {company}
Description: {description}

RULES:
1. Write entirely from the CANDIDATE's perspective — their achievements, skills, and background. Never describe what the job requires.
2. Opening sentence must highlight the candidate's most relevant expertise for this role. Use strong function/domain language (e.g. "AML compliance expert" not "7+ years of experience").
3. NEVER mention specific years of experience as a number — do not write "5+ years" or "7 years" or copy any number from the job description. Use seniority language instead: "seasoned", "senior-level", "extensive background".
4. NEVER name specific certifications, qualifications, or acronyms (e.g. CAMS, ICA Diploma, ACAMS) in the summary — they belong only in the Certifications section.
5. NEVER copy phrases, sentences, or requirements verbatim from the job description — the summary must read as the candidate's own voice.
6. No passive voice. No first person. Start with a role function qualifier.
7. 3–4 sentences maximum. No buzzwords: no "passionate", "dynamic", "results-driven", "team player".
8. Do not invent any experience, achievements, or qualifications beyond what is in the current summary and candidate background.

Return ONLY the rewritten summary paragraph."""

_SKILLS_PROMPT = """Rewrite this candidate's skills list to maximise ATS match for the target job.
The output will be used as CORE COMPETENCIES shown as a single line separated by bullets.

CURRENT SKILLS:
{current_skills}

TARGET JOB:
Title: {title}
Company: {company}
Description: {description}

RULES:
1. The first 4–5 items MUST be verbatim or near-verbatim matches to SKILL TERMS (competencies, tools, methods, regulations) named in the job description — ATS scores heavily on exact-match proximity at the top.
2. After the top matches, reorder remaining skills from the current list by relevance to this role.
3. Use the job description's exact phrasing where the candidate genuinely has that skill (e.g. if JD says "UK GDPR" and candidate has "GDPR", use "UK GDPR").
4. NEVER add certification or qualification names (e.g. CAMS, ICA Diploma, ACAMS) to the skills list — certifications belong only in the Certifications section.
5. You may refine phrasing of existing skills to better match JD terminology, but do not invent entirely new competencies the candidate has not demonstrated.
6. Return one skill per line, no bullet characters, no headers.

Return ONLY the skills, one per line."""

_TRAINING_PROMPT = """Optimise this candidate's certifications and professional development list for the target role.

CURRENT CERTIFICATIONS (these are the ONLY certifications the candidate actually holds):
{current_training}

TARGET JOB:
Title: {title}
Company: {company}
Description: {description}

RULES:
1. Reorder so the most relevant qualifications from the CURRENT list appear first.
2. NEVER alter, rename, embellish, or claim new certifications as earned. The current list is fixed — do not add anything to it as if it is achieved.
3. CRITICAL — Adding certifications: If the JD explicitly names a specific certification that is NOT in the current list above, you MAY add it as "Certification Name (In Progress)". This is the ONLY circumstance in which you add anything new. Mark it "(In Progress)" — never as earned or complete.
4. Do NOT add more than ONE new "(In Progress)" entry regardless of how many the JD mentions.
5. Do NOT add degree programmes, MBAs, or multi-year qualifications as "in progress".
6. When in doubt, add nothing — the current list is sufficient.
7. Strip any leading bullet characters — return plain text, one entry per line.

Return ONLY the certifications, one per line, no bullet characters."""

_ROLE_TITLE_PROMPT = """A candidate is applying for "{title}" at "{company}".

Their CV shows this work history entry:
"{current_header}"

ABOUT THE EMPLOYER ({role_company}):
{employer_context}

Write a new job title that:
1. Clearly signals the SAME function as "{title}" — a hiring manager must immediately see the connection.
2. Preserves the same seniority (Senior stays Senior, Analyst stays Analyst).
3. Is truthfully plausible for someone working at {role_company} — strictly respect the employer context above. Do NOT imply the employer was a regulated financial institution if the context states otherwise.
4. Is specific: "AML Compliance Specialist" beats "Compliance Officer".

Return ONLY the new job title. No company, no dates, no explanation."""

_BULLETS_PROMPT = """Rewrite these CV bullet points to maximise ATS match and relevance for the target role.
Your goal is to make this candidate's existing experience read as directly relevant to this specific job.

CURRENT BULLET POINTS:
{current_bullets}

TARGET JOB:
Title: {title}
Company: {company}
Description: {description}

ABOUT THE EMPLOYER AT THIS ROLE ({role_company}):
{employer_context}

RULES:
1. Keep EXACTLY the same number of bullets as the original.
2. Reframe honestly — emphasise the parts of each bullet most relevant to the target job, using the job's terminology where the candidate's experience genuinely matches. The factual basis (what actually happened, where, at what level) must remain exactly true. Do not upgrade the seniority, scope, or regulatory status of the work.
3. Pull language DIRECTLY from the job description where the candidate's experience genuinely maps to it. ATS systems score on exact phrase matches — use the JD's vocabulary, not generic synonyms.
4. Every bullet must start with a strong action verb (Led, Managed, Delivered, Conducted, Developed, Ensured, Implemented, Oversaw, Directed, Drove, Spearheaded).
5. BANNED openers: "Responsible for", "Assisted with", "Supported", "Involved in", "Helped to", "Worked on".
6. NEVER invent new responsibilities that are not grounded in the original bullets — reframe what exists, do not fabricate.
7. NEVER add, change, or inflate numbers/metrics. If the original bullet has a number, keep it exactly. If it has no number, do NOT add one — anchor with scope language: "organisation-wide", "cross-functional", "high-risk portfolio", "regulatory deadline".
8. Do NOT include any bullet character at the start — return plain text, one bullet per line.
9. EMPLOYER ACCURACY: Strictly respect the employer context above. If the employer is stated as non-regulated, do NOT use FCA, PRA, regulated-firm, or FCA-authorised-entity language. Replace such references with terms appropriate to the employer's actual type (e.g. "platform governance standards", "regulatory alignment for marketplace partners", "contractual due diligence obligations", "compliance advisory to regulated clients").

Return ONLY the rewritten bullet lines, no bullet characters, same count as original."""


# ── Employer context ──────────────────────────────────────────────────────────

def _get_employer_context(org: str, context_map: dict) -> str:
    """Return framing constraints for a given employer, or empty string if unconfigured."""
    for key, ctx in context_map.items():
        if key.lower() in org.lower() or org.lower() in key.lower():
            parts = []
            if ctx.get("type"):
                parts.append(f"Employer type: {ctx['type']}")
            if ctx.get("constraints"):
                parts.append(f"Constraints: {ctx['constraints']}")
            if ctx.get("title_guidance"):
                parts.append(f"Title guidance: {ctx['title_guidance']}")
            return "\n".join(parts)
    return ""


# ── CV parser ──────────────────────────────────────────────────────────────────

def _is_section_heading(para) -> bool:
    text = para.text.strip()
    if not text or len(text) > 70:
        return False
    is_allcaps  = text == text.upper() and any(c.isalpha() for c in text)
    is_hd_style = "heading" in para.style.name.lower()
    is_bold_short = (
        len(text) < 50
        and any(run.bold for run in para.runs if run.text.strip())
        and text[0] not in ("•", "-", "–", "*", "▪", "◦")
    )
    return is_allcaps or is_hd_style or is_bold_short


def _find_section_paragraphs(doc, keywords: list):
    paras = doc.paragraphs
    heading_idx = None
    for i, para in enumerate(paras):
        if not _is_section_heading(para):
            continue
        up = para.text.strip().upper()
        if any(k.upper() in up for k in keywords):
            heading_idx = i
            break
    if heading_idx is None:
        return -1, []
    content = []
    for para in paras[heading_idx + 1:]:
        if _is_section_heading(para):
            break
        content.append(para)
    return heading_idx, content


def _parse_cv(doc) -> dict:
    """Extract structured data from the base CV docx."""
    bullet_chars = ("•", "-", "–", "*", "▪", "◦", "○", "·")

    def is_bullet(para):
        t = para.text.strip()
        return bool(t) and (t[0] in bullet_chars or "list" in para.style.name.lower())

    def is_role_header(para):
        t = para.text.strip()
        if not t or is_bullet(para):
            return False
        has_pipe = "|" in t
        has_date = bool(re.search(
            r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}", t
        ))
        has_bold = any(r.bold for r in para.runs if r.text.strip())
        return (has_bold or has_pipe or has_date) and len(t) < 120

    # Name + contact (packed into paragraph 0 with \n)
    p0 = doc.paragraphs[0].text.strip().split("\n") if doc.paragraphs else []
    name    = p0[0].strip() if p0 else ""
    contact = p0[1].strip() if len(p0) > 1 else ""

    # Summary
    _, sp = _find_section_paragraphs(doc, ["PROFESSIONAL SUMMARY", "SUMMARY", "PROFILE"])
    summary = next((p.text.strip() for p in sp if p.text.strip()), "")

    # Skills (stored as multi-line in one para or as separate paras)
    _, skp = _find_section_paragraphs(doc, ["KEY SKILLS", "SKILLS", "COMPETENCIES", "EXPERTISE"])
    raw_skills = "\n".join(p.text.strip() for p in skp if p.text.strip())
    skills = [s.lstrip("•-– ").strip() for s in raw_skills.split("\n") if s.strip()]

    # Roles
    _, sec = _find_section_paragraphs(
        doc, ["EXPERIENCE", "EMPLOYMENT", "CAREER HISTORY", "WORK HISTORY", "PROFESSIONAL EXPERIENCE"]
    )
    role_starts = [i for i, p in enumerate(sec) if is_role_header(p)]
    roles = []
    for j, start in enumerate(role_starts):
        end = role_starts[j + 1] if j + 1 < len(role_starts) else len(sec)
        header = sec[start].text.strip()
        parts  = [x.strip() for x in header.split("|")]
        r_title = parts[0] if parts else header
        r_org   = parts[1] if len(parts) > 1 else ""
        r_dates = parts[2] if len(parts) > 2 else ""
        raw_bullets = []
        for bp in sec[start:end]:
            if is_bullet(bp):
                for line in bp.text.strip().split("\n"):
                    if line.strip():
                        raw_bullets.append(line.lstrip("•-–*▪◦○· ").strip())
        roles.append({"title": r_title, "org": r_org, "dates": r_dates, "bullets": raw_bullets})

    # Education
    _, ep = _find_section_paragraphs(doc, ["EDUCATION"])
    education = [p.text.strip() for p in ep if p.text.strip()]

    # Certifications
    _, cp = _find_section_paragraphs(
        doc, ["CERTIFICATIONS", "TRAINING", "CERTIFICATION", "PROFESSIONAL DEVELOPMENT"]
    )
    raw_certs = "\n".join(p.text.strip() for p in cp if p.text.strip())
    certifications = [c.lstrip("•-– ").strip() for c in raw_certs.split("\n") if c.strip()]

    return {
        "name": name,
        "contact": contact,
        "summary": summary,
        "skills": skills,
        "roles": roles,
        "education": education,
        "certifications": certifications,
    }


# ── Document builder (house format) ───────────────────────────────────────────

def _set_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp  = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    if before: sp.set(qn("w:before"), str(before))
    if after:  sp.set(qn("w:after"),  str(after))


def _add_bottom_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), "1F4E79")
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_bullet_indent(para):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:left"),    "360")
    ind.set(qn("w:hanging"), "200")


def _add_right_tab(para, pos=9360):
    pPr = para._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos))
    tabs.append(tab)


def _run(para, text, size_hp=21, bold=False, italic=False, color=None):
    """Add a formatted run. size_hp is half-points (21 hp = 10.5 pt)."""
    run = para.add_run(text)
    run.font.name  = "Arial"
    run.font.size  = Pt(size_hp / 2)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color
    return run


def _add_tab_run(para):
    """Add a DOCX <w:tab/> element run (for right-aligned tab in role lines)."""
    from lxml import etree
    r_elem = etree.SubElement(para._p, f"{{{W}}}r")
    etree.SubElement(r_elem, f"{{{W}}}tab")


def _build_cv(data: dict, output_path: str) -> str:
    """Render structured CV data to a new .docx using the house format spec."""
    doc = Document()

    # Page setup — A4, margins 0.75" top/bottom, 1" left/right
    sec = doc.sections[0]
    sec.page_width    = Twips(11906)
    sec.page_height   = Twips(16838)
    sec.top_margin    = Twips(1080)
    sec.bottom_margin = Twips(1080)
    sec.left_margin   = Twips(1440)
    sec.right_margin  = Twips(1440)

    # Suppress the default Normal style spacing so we control all spacing explicitly
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)

    # Remove the single empty paragraph that Document() creates
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # ── Helper: new paragraph ──────────────────────────────────────────────────
    def new_para():
        from docx.oxml import OxmlElement
        p_elem = OxmlElement("w:p")
        doc.element.body.append(p_elem)
        from docx.text.paragraph import Paragraph
        return Paragraph(p_elem, doc.element.body)

    # ── Helper: section heading ────────────────────────────────────────────────
    def add_heading(text):
        p = new_para()
        _set_spacing(p, before=240, after=80)
        _add_bottom_border(p)
        _run(p, text.upper(), size_hp=24, bold=True, color=ACCENT)

    # ── Name block ─────────────────────────────────────────────────────────────
    p = new_para()
    _set_spacing(p, after=20)
    _run(p, data["name"], size_hp=36, bold=True, color=ACCENT)

    p = new_para()
    _set_spacing(p, after=20)
    _run(p, data["professional_title"], size_hp=24, color=DARK)

    p = new_para()
    _set_spacing(p, after=40)
    _run(p, data["contact"], size_hp=20, color=LGREY)

    # ── PROFESSIONAL SUMMARY ───────────────────────────────────────────────────
    add_heading("PROFESSIONAL SUMMARY")
    p = new_para()
    _set_spacing(p, after=80)
    _run(p, data["summary"], size_hp=21, color=DARK)

    # ── CORE COMPETENCIES ──────────────────────────────────────────────────────
    add_heading("CORE COMPETENCIES")
    p = new_para()
    _set_spacing(p, after=80)
    competencies = "  •  ".join(s for s in data["skills"] if s)
    _run(p, competencies, size_hp=21, color=DARK)

    # ── PROFESSIONAL EXPERIENCE ────────────────────────────────────────────────
    add_heading("PROFESSIONAL EXPERIENCE")
    for role in data["roles"]:
        # Role header: Title (bold) — Org  [tab]  Dates (italic grey)
        p = new_para()
        _set_spacing(p, before=160, after=20)
        _add_right_tab(p)
        _run(p, role["title"], size_hp=22, bold=True, color=DARK)
        _run(p, "  —  ",       size_hp=22, color=DARK)
        _run(p, role["org"],   size_hp=22, color=DARK)
        _add_tab_run(p)
        _run(p, role["dates"], size_hp=20, italic=True, color=GREY)

        for bullet in role["bullets"]:
            p = new_para()
            _set_spacing(p, after=40)
            _set_bullet_indent(p)
            _run(p, "•  " + bullet.lstrip("•-–*▪◦○· ").strip(), size_hp=21, color=DARK)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    add_heading("EDUCATION")
    for edu in data["education"]:
        p = new_para()
        _set_spacing(p, after=60)
        _add_right_tab(p)
        parts = [x.strip() for x in edu.split("|")]
        if len(parts) >= 3:
            _run(p, parts[0], size_hp=21, bold=True, color=DARK)
            _run(p, "  —  ",  size_hp=21, color=DARK)
            _run(p, parts[1], size_hp=21, color=DARK)
            _add_tab_run(p)
            _run(p, parts[2], size_hp=20, italic=True, color=GREY)
        else:
            _run(p, edu, size_hp=21, color=DARK)

    # ── CERTIFICATIONS & PROFESSIONAL DEVELOPMENT ──────────────────────────────
    add_heading("CERTIFICATIONS & PROFESSIONAL DEVELOPMENT")
    for cert in data["certifications"]:
        p = new_para()
        _set_spacing(p, after=40)
        _set_bullet_indent(p)
        _run(p, "•  " + cert.lstrip("•-– ").strip(), size_hp=21, color=DARK)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ── LLM tailoring ──────────────────────────────────────────────────────────────

def tailor_cv_docx(title: str, company: str, description: str, output_path: str = None) -> str:
    """Parse base CV, run LLM rewrites, build house-formatted output. Returns path."""
    import json as _json
    doc_base     = Document(str(BASE_CV_PATH))
    desc_short   = str(description)[:3000]
    profile_text = PROFILE_PATH.read_text()[:2500] if PROFILE_PATH.exists() else ""
    logger.info(f"Tailoring CV for: {title} @ {company}")

    # Load per-user employer context (truthful framing constraints per employer)
    employer_context_map: dict = {}
    ctx_path = Path(os.getenv("CV_PATH", "")).parent / "employer_context.json"
    if ctx_path.exists():
        try:
            employer_context_map = _json.loads(ctx_path.read_text())
            logger.info(f"Loaded employer context for {len(employer_context_map)} employers")
        except Exception as e:
            logger.warning(f"Could not load employer_context.json: {e}")

    data = _parse_cv(doc_base)

    # Professional title under name — the candidate's REAL current title.
    # Never the target job title: claiming the advertised title as your own
    # reads as fabrication (fatal in compliance/regulated hiring).
    data["professional_title"] = (data["roles"][0]["title"]
                                  if data.get("roles") and data["roles"][0].get("title")
                                  else "")

    # ── 1. Professional Summary ────────────────────────────────────────────────
    try:
        if data["summary"]:
            new_summary = call_llm(_SYS, _SUMMARY_PROMPT.format(
                current_summary=data["summary"], profile_snippet=profile_text,
                title=title, company=company, description=desc_short,
            ), max_tokens=400).strip()
            if new_summary:
                data["summary"] = new_summary
                logger.info(f"Summary rewritten: {new_summary[:80]}…")
        time.sleep(1)
    except Exception as e:
        logger.warning(f"Summary rewrite failed: {e}")

    # ── 2. Core Competencies (skills) ─────────────────────────────────────────
    try:
        skill_text = "\n".join(data["skills"])
        if skill_text:
            new_skills_raw = call_llm(_SYS, _SKILLS_PROMPT.format(
                current_skills=skill_text,
                title=title, company=company, description=desc_short,
            ), max_tokens=600).strip()
            if new_skills_raw:
                data["skills"] = [s.lstrip("•-– ").strip()
                                   for s in new_skills_raw.split("\n") if s.strip()]
                logger.info(f"Skills rewritten: {len(data['skills'])} items")
        time.sleep(1)
    except Exception as e:
        logger.warning(f"Skills rewrite failed: {e}")

    # ── 3. Certifications ─────────────────────────────────────────────────────
    try:
        cert_text = "\n".join(data["certifications"])
        if cert_text:
            new_certs_raw = call_llm(_SYS, _TRAINING_PROMPT.format(
                current_training=cert_text,
                title=title, company=company, description=desc_short,
            ), max_tokens=400).strip()
            if new_certs_raw:
                data["certifications"] = [c.lstrip("•-– ").strip()
                                           for c in new_certs_raw.split("\n") if c.strip()]
                logger.info("Certifications rewritten")
        time.sleep(1)
    except Exception as e:
        logger.warning(f"Certifications rewrite failed: {e}")

    # ── 4 & 5. Top 2 roles — rewrite title then bullets ───────────────────────
    for i, role in enumerate(data["roles"][:2]):
        role_company = role.get("org", "")
        emp_ctx = _get_employer_context(role_company, employer_context_map)
        ctx_block = emp_ctx if emp_ctx else "No special constraints — standard tailoring applies."

        # 4. Role title — NOT rewritten. Employer names, dates, and job titles
        # are the three facts UK financial-services background checks verify;
        # a title that fails referencing kills the offer. Tailoring lives in
        # the summary, skills ordering, and bullet emphasis only.

        # 5. Bullets
        try:
            bullet_text = "\n".join(role["bullets"])
            if bullet_text.strip():
                new_bullets_raw = call_llm(_SYS, _BULLETS_PROMPT.format(
                    role_title=role["title"],
                    current_bullets=bullet_text,
                    title=title, company=company, description=desc_short,
                    role_company=role_company,
                    employer_context=ctx_block,
                ), max_tokens=700).strip()
                if new_bullets_raw:
                    role["bullets"] = [b.lstrip("•-–*▪◦○· ").strip()
                                        for b in new_bullets_raw.split("\n") if b.strip()]
                    logger.info(f"Role {i+1} bullets rewritten: {len(role['bullets'])} bullets")
            time.sleep(1)
        except Exception as e:
            logger.warning(f"Role {i+1} bullets rewrite failed: {e}")

    # ── Build formatted output ─────────────────────────────────────────────────
    if not output_path:
        safe = re.sub(r"[^\w]", "_", title.lower())[:40]
        output_path = str(TMP_DIR / f"cv_{safe}.docx")

    TMP_DIR.mkdir(exist_ok=True)
    result = _build_cv(data, output_path)
    logger.info(f"Saved tailored CV: {result}")
    return result


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--job-title",   required=True)
    parser.add_argument("--company",     required=True)
    parser.add_argument("--description", default="")
    parser.add_argument("--output",      default=None)
    args = parser.parse_args()
    print(tailor_cv_docx(args.job_title, args.company, args.description, args.output))


if __name__ == "__main__":
    main()
