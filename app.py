"""
app.py — ApplyExpress SaaS Platform

A self-contained Flask web application:
  - User signup / login (SQLite, no external DB)
  - CV upload (stored in data/users/<id>/)
  - Profile & job-search preferences
  - Pipeline trigger — runs fetch → score → tailor CV → cover letter
    then emails the full report automatically
  - REST API for the Chrome extension

Deploy FREE on Fly.io (see fly.toml + Dockerfile).
Local dev:
    python app.py
"""

import os, sys, json, sqlite3, secrets, hashlib, threading, time, logging, re, fcntl, imaplib, email as _email
from pathlib import Path
from datetime import datetime, timedelta
from functools import wraps

try:
    import stripe
    stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "")
except ImportError:
    stripe = None

try:
    import sendgrid
    from sendgrid.helpers.mail import Mail
except ImportError:
    sendgrid = None

try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    _limiter_available = True
except ImportError:
    _limiter_available = False

from dotenv import load_dotenv
load_dotenv()

from flask import (Flask, request, jsonify, session,
                   redirect, url_for, g)
from flask_cors import CORS

# ── Project setup ─────────────────────────────────────────────────────────────

ROOT        = Path(__file__).parent
DATA_DIR    = Path(os.getenv("DATA_DIR", str(ROOT / "data")))
DB_PATH     = DATA_DIR / "autoapply.db"
TOOLS_DIR   = ROOT / "tools"

DATA_DIR.mkdir(parents=True, exist_ok=True)
sys.path.insert(0, str(ROOT))

logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s [autoapply] %(message)s")
log = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", secrets.token_hex(32))
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    # Enable once the site is served over HTTPS (Cloudflare/TLS): COOKIE_SECURE=1 in .env.production
    SESSION_COOKIE_SECURE=os.getenv("COOKIE_SECURE", "0") == "1",
)

# ── Credential encryption (Fernet symmetric) ──────────────────────────────────
# Job-board passwords are encrypted at rest in the DB.
# The key lives only in .env — never in code or DB.
try:
    from cryptography.fernet import Fernet as _Fernet
    _cred_key = os.getenv("CREDENTIAL_KEY", "").encode()
    _fernet   = _Fernet(_cred_key) if _cred_key else None
except Exception:
    _fernet = None

def _enc(plaintext: str) -> str:
    """Encrypt a credential string for DB storage. Returns plaintext if key missing."""
    if not plaintext or not _fernet:
        return plaintext
    # Don't double-encrypt values already encrypted (start with 'gAAAAA' — Fernet token prefix)
    if isinstance(plaintext, str) and plaintext.startswith("gAAAAA"):
        return plaintext
    return _fernet.encrypt(plaintext.encode()).decode()

def _dec(ciphertext: str) -> str:
    """Decrypt a credential from DB. Returns empty string on failure."""
    if not ciphertext:
        return ""
    if not _fernet:
        return ciphertext  # graceful fallback if key missing
    if not ciphertext.startswith("gAAAAA"):
        return ciphertext  # already plaintext (pre-migration rows)
    try:
        return _fernet.decrypt(ciphertext.encode()).decode()
    except Exception:
        return ""
CORS(app, supports_credentials=True, origins="*")

# Trust nginx reverse proxy (1 hop) so Flask-Limiter and url_for see real IPs / https
from werkzeug.middleware.proxy_fix import ProxyFix
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

if _limiter_available:
    limiter = Limiter(
        get_remote_address,
        app=app,
        default_limits=["200 per hour"],
        storage_uri="memory://",
    )
else:
    # Stub so @limiter.limit decorators don't crash if flask-limiter missing
    class _NoopLimiter:
        def limit(self, *a, **kw):
            return lambda f: f
        def exempt(self, f):
            return f
    limiter = _NoopLimiter()

# ── Database ──────────────────────────────────────────────────────────────────

def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(str(DB_PATH), check_same_thread=False)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA journal_mode=WAL")
        g.db.execute("PRAGMA foreign_keys=ON")
    return g.db

@app.teardown_appcontext
def close_db(_):
    db = g.pop("db", None)
    if db: db.close()

def init_db():
    db = sqlite3.connect(str(DB_PATH))
    db.executescript("""
    CREATE TABLE IF NOT EXISTS users (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        email           TEXT UNIQUE NOT NULL,
        password_hash   TEXT NOT NULL,
        first_name      TEXT DEFAULT '',
        last_name       TEXT DEFAULT '',
        phone           TEXT DEFAULT '',
        location        TEXT DEFAULT 'United Kingdom',
        api_key         TEXT UNIQUE NOT NULL,
        created_at      TEXT NOT NULL,

        -- Job search prefs
        keywords        TEXT DEFAULT 'compliance analyst',
        search_location TEXT DEFAULT 'London',
        threshold       INTEGER DEFAULT 7,
        email_subject   TEXT DEFAULT 'Job Application',

        -- Credentials (job boards)
        linkedin_email  TEXT DEFAULT '',
        linkedin_pass   TEXT DEFAULT '',
        reed_email      TEXT DEFAULT '',
        reed_pass       TEXT DEFAULT '',
        indeed_email    TEXT DEFAULT '',
        indeed_pass     TEXT DEFAULT '',

        -- Gmail app password for reports
        smtp_password   TEXT DEFAULT '',

        -- Pipeline state
        last_run        TEXT DEFAULT '',
        last_run_status TEXT DEFAULT ''
    );

    CREATE TABLE IF NOT EXISTS runs (
        id          INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id     INTEGER NOT NULL,
        run_id      TEXT NOT NULL,
        started_at  TEXT NOT NULL,
        finished_at TEXT DEFAULT '',
        status      TEXT DEFAULT 'running',
        jobs_found  INTEGER DEFAULT 0,
        jobs_applied INTEGER DEFAULT 0,
        report_json TEXT DEFAULT '',
        specialty_id INTEGER DEFAULT NULL,
        FOREIGN KEY(user_id) REFERENCES users(id)
    );

    CREATE TABLE IF NOT EXISTS specialties (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id         INTEGER NOT NULL,
        name            TEXT NOT NULL,
        slug            TEXT NOT NULL,
        keywords        TEXT DEFAULT '',
        search_location TEXT DEFAULT '',
        threshold       INTEGER DEFAULT 7,
        created_at      TEXT NOT NULL,
        UNIQUE(user_id, slug),
        FOREIGN KEY(user_id) REFERENCES users(id)
    );

    CREATE TABLE IF NOT EXISTS applications (
        id          INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id     INTEGER NOT NULL,
        run_db_id   INTEGER,
        title       TEXT DEFAULT '',
        company     TEXT DEFAULT '',
        url         TEXT DEFAULT '',
        status      TEXT DEFAULT 'applied',
        notes       TEXT DEFAULT '',
        applied_at  TEXT NOT NULL,
        FOREIGN KEY(user_id) REFERENCES users(id),
        FOREIGN KEY(run_db_id) REFERENCES runs(id)
    );

    CREATE TABLE IF NOT EXISTS blacklist (
        id         INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id    INTEGER NOT NULL,
        company    TEXT NOT NULL,
        added_at   TEXT NOT NULL,
        UNIQUE(user_id, company),
        FOREIGN KEY(user_id) REFERENCES users(id)
    );

    CREATE INDEX IF NOT EXISTS idx_applications_user_url
        ON applications(user_id, url);
    CREATE INDEX IF NOT EXISTS idx_applications_user_status
        ON applications(user_id, status, applied_at);
    """)
    # Migrate existing DBs — add new columns if missing
    existing = {r[1] for r in db.execute("PRAGMA table_info(users)").fetchall()}
    new_user_cols = [
        ("linkedin_email",   "TEXT DEFAULT ''"),
        ("reed_email",       "TEXT DEFAULT ''"),
        ("indeed_email",     "TEXT DEFAULT ''"),
        ("is_paid",          "INTEGER DEFAULT 0"),
        ("trial_ends_at",    "TEXT DEFAULT ''"),
        ("stripe_customer",  "TEXT DEFAULT ''"),
        ("stripe_sub_id",    "TEXT DEFAULT ''"),
        ("onboarding_done",  "INTEGER DEFAULT 0"),
        ("consented_at",     "TEXT DEFAULT ''"),
        ("twocaptcha_key",   "TEXT DEFAULT ''"),
        ("sponsor_check",    "INTEGER DEFAULT 0"),
    ]
    for col, defn in new_user_cols:
        if col not in existing:
            db.execute(f"ALTER TABLE users ADD COLUMN {col} {defn}")
    existing_run_cols = {r[1] for r in db.execute("PRAGMA table_info(runs)").fetchall()}
    if "specialty_id" not in existing_run_cols:
        db.execute("ALTER TABLE runs ADD COLUMN specialty_id INTEGER DEFAULT NULL")
    db.commit()
    db.close()
    log.info(f"DB ready: {DB_PATH}")

# ── Auth helpers ──────────────────────────────────────────────────────────────

def _hash(password: str) -> str:
    """Salted PBKDF2 hash for new/updated passwords."""
    from werkzeug.security import generate_password_hash
    return generate_password_hash(password)

def _verify_password(stored_hash: str, password: str) -> bool:
    """Verify against PBKDF2, falling back to legacy unsalted SHA-256."""
    if not stored_hash:
        return False
    if ":" in stored_hash or stored_hash.startswith(("pbkdf2", "scrypt")):
        from werkzeug.security import check_password_hash
        try:
            return check_password_hash(stored_hash, password)
        except Exception:
            return False
    # Legacy unsalted SHA-256 (pre-launch accounts)
    return secrets.compare_digest(stored_hash,
                                  hashlib.sha256(password.encode()).hexdigest())

def _user_dir(user_id: int) -> Path:
    d = DATA_DIR / "users" / str(user_id)
    d.mkdir(parents=True, exist_ok=True)
    (d / ".tmp").mkdir(exist_ok=True)
    return d

def _specialty_dir(user_id: int, slug: str) -> Path:
    d = _user_dir(user_id) / "specialties" / slug
    d.mkdir(parents=True, exist_ok=True)
    return d

def _extract_cv_profile(cv_path: Path, user, keywords: str, location: str) -> str:
    """Extract text from a .docx CV and return a candidate_profile.md string."""
    from docx import Document as DocxDocument
    doc = DocxDocument(str(cv_path))
    cv_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return f"""## Personal Info
- Name: {user['first_name']} {user['last_name']}
- Email: {user['email']}
- Phone: {user.get('phone','') or ''}
- Location: {user.get('location','United Kingdom') or 'United Kingdom'}

## Target Roles
- Job titles: {keywords}
- Work arrangement: Hybrid
- Right to work in UK: Yes

## Full CV Text
{cv_text}
"""

def _cv_parse_warning(cv_path) -> str:
    """Parse an uploaded CV with the pipeline's parser and warn about anything
    it cannot read — so users learn at upload time, not via blank applications."""
    try:
        from docx import Document as _DocxDocument
        from tools.tailor_cv_docx import _parse_cv
        d = _parse_cv(_DocxDocument(str(cv_path)))
        problems = []
        if not d["roles"]:
            problems.append("work experience entries")
        if not d["summary"]:
            problems.append("a professional summary")
        if not d["skills"]:
            problems.append("a skills list")
        if not d["contact"]:
            problems.append("contact details")
        if problems:
            return (" ⚠ We could not read " + ", ".join(problems) + " from this CV. "
                    "Tailored CVs may be sent with that content missing. Tip: use clear section "
                    "headings (e.g. PROFESSIONAL SUMMARY, KEY SKILLS, PROFESSIONAL EXPERIENCE) and "
                    "one line per role like 'Job Title | Company | Jan 2020 – Present'.")
    except Exception:
        pass
    return ""

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            if request.path.startswith("/api/"):
                return jsonify({"error": "Unauthorized"}), 401
            return redirect(url_for("login_page"))
        return f(*args, **kwargs)
    return decorated

def api_key_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        key = request.headers.get("X-API-Key") or request.args.get("api_key", "")
        db  = get_db()
        row = db.execute("SELECT * FROM users WHERE api_key=?", (key,)).fetchone()
        if not row:
            return jsonify({"error": "Unauthorized"}), 401
        g.api_user = dict(row)
        return f(*args, **kwargs)
    return decorated

def current_user():
    if "user_id" not in session:
        return None
    db  = get_db()
    row = db.execute("SELECT * FROM users WHERE id=?", (session["user_id"],)).fetchone()
    return dict(row) if row else None

def _is_active(u: dict) -> bool:
    """Return True if user has an active paid plan or is within the 14-day trial."""
    if u.get("is_paid"):
        return True
    trial_ends = u.get("trial_ends_at") or ""
    return bool(trial_ends) and trial_ends > datetime.utcnow().isoformat()

def paid_required(f):
    """Require an active subscription or unexpired trial."""
    @wraps(f)
    def decorated(*args, **kwargs):
        u = current_user()
        if not u or not _is_active(u):
            if request.path.startswith("/api/"):
                return jsonify({"error": "Subscription required", "upgrade": "/pricing"}), 402
            return redirect(url_for("pricing") + "?gate=1")
        return f(*args, **kwargs)
    return decorated

def _send_email(to: str, subject: str, html: str, smtp_password: str = ""):
    """Send transactional email. Tries SMTP (Gmail) first if a password is provided,
    then falls back to Sendgrid. Logs clearly if neither is configured."""

    # ── SMTP via Gmail (preferred — no extra API key needed) ──────────────────
    smtp_pwd = smtp_password or os.getenv("SMTP_PASSWORD", "")
    smtp_user = os.getenv("SMTP_EMAIL", to)
    if smtp_pwd:
        try:
            import smtplib
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText as _MIMEText
            msg = MIMEMultipart("alternative")
            msg["Subject"] = subject
            msg["From"]    = smtp_user
            msg["To"]      = to
            msg.attach(_MIMEText(html, "html"))
            with smtplib.SMTP("smtp.gmail.com", 587) as srv:
                srv.starttls()
                srv.login(smtp_user, smtp_pwd)
                srv.sendmail(smtp_user, [to], msg.as_string())
            log.info(f"[email] SMTP sent '{subject}' to {to}")
            return
        except Exception as e:
            log.warning(f"[email] SMTP failed: {e} — trying Sendgrid")

    # ── Sendgrid fallback ─────────────────────────────────────────────────────
    sg_key = os.getenv("SENDGRID_API_KEY", "")
    if sg_key and sendgrid:
        try:
            sg   = sendgrid.SendGridAPIClient(api_key=sg_key)
            mail = Mail(
                from_email=os.getenv("FROM_EMAIL", "noreply@applyexpress.io"),
                to_emails=to,
                subject=subject,
                html_content=html,
            )
            sg.send(mail)
            log.info(f"[email] Sendgrid sent '{subject}' to {to}")
            return
        except Exception as e:
            log.warning(f"[email] Sendgrid failed: {e}")

    log.warning(f"[email] No email sent — configure Gmail app password in Profile settings "
                f"(to: {to}, subject: {subject})")

# ── HTML helpers ──────────────────────────────────────────────────────────────

def _page(title, body, user=None, extra_head=""):
    nav_links = (
        f'<a href="/dashboard" class="nav-link">Dashboard</a>'
        f'<a href="/history" class="nav-link">History</a>'
        f'<a href="/profile" class="nav-link">Settings</a>'
        f'<a href="/logout" class="nav-link nav-link-muted">Sign out</a>'
    ) if user else (
        f'<a href="/login" class="nav-link">Log in</a>'
        f'<a href="/signup" class="btn-nav">Get started free</a>'
    )
    avatar = ""
    if user:
        try:
            fn = (user["first_name"] or "?")[:1].upper()
            ln = (user["last_name"]  or "?")[:1].upper()
            avatar = f'<div class="nav-avatar">{fn}{ln}</div>'
        except Exception:
            pass
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{title} · ApplyExpress</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:opsz,wght@9..40,400;9..40,500;9..40,600;9..40,700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
<link href="https://api.fontshare.com/v2/css?f[]=clash-display@500,600,700&display=swap" rel="stylesheet">
{extra_head}
<style>
/* ── Tokens ── */
:root{{
  --orange:#ff812d;
  --orange-light:rgba(255,180,110,.45);
  --orange-deep:rgba(200,90,10,.35);
  --green:#0c2718;
  --green-mid:#1a5c34;
  --green-light:rgba(12,39,24,.65);
  --green-hint:rgba(12,39,24,.12);
  --glass:rgba(255,255,255,.22);
  --glass-hover:rgba(255,255,255,.30);
  --glass-border:rgba(255,255,255,.42);
  --glass-border-bright:rgba(255,255,255,.62);
  --text-primary:#0c2718;
  --text-secondary:rgba(12,39,24,.65);
  --text-tertiary:rgba(12,39,24,.45);
  --accent:#1a5c34;
  --accent-dim:#0c2718;
  --success:#0c2718;
  --warning:#6b3000;
  --error:#7a0e0e;
  --r-sm:6px;--r:10px;--r-lg:16px;--r-xl:20px;
  --font-display:'Clash Display',sans-serif;
  --font-body:'DM Sans',sans-serif;
  --font-mono:'JetBrains Mono',monospace;
  --ease-spring:cubic-bezier(0.16,1,0.3,1);
}}
/* ── Reset ── */
*{{box-sizing:border-box;margin:0;padding:0}}
body{{
  font-family:var(--font-body);
  background:
    radial-gradient(ellipse at 20% 15%,var(--orange-light) 0%,transparent 45%),
    radial-gradient(ellipse at 82% 85%,var(--orange-deep) 0%,transparent 45%),
    var(--orange);
  background-attachment:fixed;
  color:var(--text-primary);
  line-height:1.6;font-size:14px;
  -webkit-font-smoothing:antialiased;min-height:100vh;
}}
::selection{{background:rgba(12,39,24,.2);color:#0c2718}}
a{{color:var(--green-mid);text-decoration:none;transition:opacity .15s}}
a:hover{{opacity:.75}}
/* ── Type ── */
h1{{font-family:var(--font-display);font-size:30px;font-weight:700;
    letter-spacing:-.02em;margin-bottom:8px;line-height:1.15;color:var(--text-primary)}}
h2{{font-family:var(--font-display);font-size:21px;font-weight:600;
    letter-spacing:-.01em;margin-bottom:14px;line-height:1.2;color:var(--text-primary)}}
h3{{font-size:15px;font-weight:600;margin-bottom:10px;color:var(--text-primary)}}
p{{color:var(--text-secondary)}}
.label{{font-size:11px;font-weight:600;color:var(--text-tertiary);margin-bottom:6px;
        display:block;letter-spacing:.07em;text-transform:uppercase}}
.mono{{font-family:var(--font-mono);font-size:12px;background:rgba(255,255,255,.25);
       color:var(--text-primary);padding:10px 14px;border-radius:var(--r);
       word-break:break-all;border:1px solid var(--glass-border)}}
/* ── Layout ── */
.container{{max-width:880px;margin:0 auto;padding:40px 24px}}
.container-lg{{max-width:1100px;margin:0 auto;padding:40px 24px}}
.container-sm{{max-width:560px;margin:0 auto;padding:40px 24px}}
.row{{display:grid;grid-template-columns:1fr 1fr;gap:16px}}
.row-3{{display:grid;grid-template-columns:1fr 1fr 1fr;gap:16px}}
.field{{margin-bottom:18px}}
hr{{border:none;border-top:1px solid var(--glass-border);margin:24px 0}}
/* ── Nav ── */
nav{{
  background:rgba(220,90,22,.52);
  backdrop-filter:blur(22px) saturate(1.8);
  -webkit-backdrop-filter:blur(22px) saturate(1.8);
  border-bottom:1px solid rgba(255,255,255,.22);
  padding:0 32px;display:flex;align-items:center;
  justify-content:space-between;height:62px;
  position:sticky;top:0;z-index:200;
  transition:background .2s;
}}
nav.scrolled{{background:rgba(200,75,10,.75)}}
.nav-brand{{
  font-family:var(--font-display);font-weight:700;font-size:18px;
  color:var(--text-primary)!important;letter-spacing:-.02em;
  display:flex;align-items:center;gap:10px;opacity:1!important;
  text-decoration:none!important;
}}
/* Logo B — pulsing rings */
.nav-logo{{width:26px;height:26px;flex-shrink:0}}
.nav-links{{display:flex;gap:2px;align-items:center}}
.nav-link{{
  color:rgba(12,39,24,.65)!important;font-size:13px;font-weight:500;
  padding:6px 13px;border-radius:var(--r-sm);
  transition:color .15s,background .15s;opacity:1!important;
  text-decoration:none!important;
}}
.nav-link:hover{{color:var(--text-primary)!important;background:rgba(255,255,255,.18);opacity:1!important}}
.nav-link-muted{{color:rgba(12,39,24,.45)!important}}
.nav-link-muted:hover{{color:#7a0e0e!important;background:rgba(122,14,14,.1);opacity:1!important}}
.nav-avatar{{
  width:32px;height:32px;border-radius:50%;
  background:linear-gradient(135deg,var(--green-mid),var(--green));
  color:#fdf8f0;font-size:11px;font-weight:800;
  display:flex;align-items:center;justify-content:center;
  margin-left:8px;font-family:var(--font-mono);flex-shrink:0;
  border:2px solid rgba(255,255,255,.35);
  box-shadow:0 2px 8px rgba(0,0,0,.18);
}}
.btn-nav{{
  background:linear-gradient(135deg,var(--green-mid),var(--green))!important;
  color:#fdf8f0!important;padding:8px 16px;
  border-radius:var(--r);font-weight:700;font-size:13px;
  transition:transform .12s,box-shadow .15s;border:none;cursor:pointer;
  display:inline-flex;align-items:center;
  box-shadow:0 2px 10px rgba(12,39,24,.3);
  text-decoration:none!important;
}}
.btn-nav:hover{{transform:scale(1.03);box-shadow:0 4px 16px rgba(12,39,24,.35);opacity:1!important}}
/* ── Glass card ── */
.card{{
  background:var(--glass);
  backdrop-filter:blur(20px) saturate(1.5);
  -webkit-backdrop-filter:blur(20px) saturate(1.5);
  border:1.5px solid var(--glass-border);
  border-radius:var(--r-lg);padding:28px;margin-bottom:20px;
  box-shadow:0 6px 28px rgba(0,0,0,.08),inset 0 1.5px 0 rgba(255,255,255,.55),inset 0 -1px 0 rgba(0,0,0,.03);
  transition:background .2s,transform .2s,box-shadow .2s;
}}
.card-interactive:hover{{
  background:var(--glass-hover);transform:translateY(-2px);
  box-shadow:0 12px 40px rgba(0,0,0,.12),inset 0 1.5px 0 rgba(255,255,255,.65);
}}
.card-accent{{border-color:rgba(12,39,24,.35);background:rgba(12,39,24,.1)}}
/* ── Forms ── */
input[type=text],input[type=email],input[type=password],input[type=tel],
input[type=number],input[type=file],select,textarea{{
  width:100%;padding:10px 13px;
  border:1.5px solid rgba(255,255,255,.38);
  border-radius:var(--r);font-size:14px;
  color:var(--text-primary);
  background:rgba(255,255,255,.25);
  outline:none;font-family:var(--font-body);
  transition:border-color .15s,box-shadow .15s;
  backdrop-filter:blur(8px);
}}
input::placeholder,textarea::placeholder{{color:var(--text-tertiary)}}
input:focus,select:focus,textarea:focus{{
  border-color:rgba(12,39,24,.55);
  box-shadow:0 0 0 3px rgba(12,39,24,.1);
  background:rgba(255,255,255,.32);
}}
select{{
  appearance:none;
  background-image:url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='12' height='8' viewBox='0 0 12 8'%3E%3Cpath fill='%230c2718' opacity='.5' d='M6 8 0 0h12z'/%3E%3C/svg%3E");
  background-repeat:no-repeat;background-position:right 13px center;padding-right:34px;
}}
textarea{{resize:vertical;min-height:72px}}
input[type=checkbox]{{width:16px;height:16px;accent-color:var(--green-mid);cursor:pointer}}
/* ── Buttons ── */
.btn{{
  display:inline-flex;align-items:center;justify-content:center;gap:6px;
  padding:10px 22px;border-radius:var(--r);font-size:14px;font-weight:700;
  cursor:pointer;border:none;transition:transform .12s,box-shadow .15s,opacity .15s;
  font-family:var(--font-body);white-space:nowrap;text-decoration:none!important;
  letter-spacing:.01em;
}}
.btn:hover{{transform:scale(1.02)}}
.btn:active{{transform:scale(.97)}}
.btn-primary{{
  background:linear-gradient(145deg,var(--green-mid),var(--green));
  color:#fdf8f0;
  box-shadow:0 3px 12px rgba(12,39,24,.35),inset 0 1px 0 rgba(255,255,255,.12);
}}
.btn-primary:hover{{box-shadow:0 6px 20px rgba(12,39,24,.4);color:#fdf8f0;opacity:1}}
.btn-success{{
  background:linear-gradient(145deg,#1a5c34,#0c2718);
  color:#fdf8f0;
  box-shadow:0 3px 12px rgba(12,39,24,.35);
}}
.btn-success:hover{{box-shadow:0 6px 20px rgba(12,39,24,.4);color:#fdf8f0;opacity:1}}
.btn-outline{{
  background:rgba(255,255,255,.22);color:var(--text-primary);
  border:1.5px solid var(--glass-border);
  box-shadow:inset 0 1px 0 rgba(255,255,255,.5);
}}
.btn-outline:hover{{background:rgba(255,255,255,.32);border-color:var(--glass-border-bright);opacity:1}}
.btn-danger{{background:rgba(122,14,14,.15);color:#7a0e0e;border:1.5px solid rgba(122,14,14,.2)}}
.btn-danger:hover{{background:rgba(122,14,14,.25);opacity:1}}
.btn-ghost{{background:transparent;color:var(--text-secondary)}}
.btn-ghost:hover{{background:rgba(255,255,255,.2);color:var(--text-primary);opacity:1}}
.btn-full{{width:100%}}
.btn-sm{{padding:7px 14px;font-size:12px}}
.btn-lg{{padding:13px 32px;font-size:15px}}
/* ── Badges ── */
.badge{{
  display:inline-flex;align-items:center;gap:4px;
  padding:3px 10px;border-radius:20px;
  font-size:11px;font-weight:700;letter-spacing:.04em;
}}
.badge-green{{background:rgba(12,39,24,.18);color:#0c2718;border:1px solid rgba(12,39,24,.22)}}
.badge-yellow{{background:rgba(107,48,0,.15);color:#6b3000;border:1px solid rgba(107,48,0,.2)}}
.badge-red{{background:rgba(122,14,14,.15);color:#7a0e0e;border:1px solid rgba(122,14,14,.2)}}
.badge-blue{{background:rgba(12,39,24,.12);color:#0c2718;border:1px solid rgba(12,39,24,.18)}}
.badge-gray{{background:rgba(255,255,255,.25);color:rgba(12,39,24,.65);border:1px solid var(--glass-border)}}
/* ── Alerts ── */
.alert{{padding:13px 17px;border-radius:var(--r);font-size:13px;margin-bottom:20px;line-height:1.5}}
.alert-error{{background:rgba(122,14,14,.12);color:#7a0e0e;border:1px solid rgba(122,14,14,.2)}}
.alert-success{{background:rgba(12,39,24,.12);color:#0c2718;border:1px solid rgba(12,39,24,.22)}}
.alert-info{{background:rgba(255,255,255,.22);color:var(--text-primary);border:1px solid var(--glass-border)}}
.info-tip{{
  background:rgba(255,255,255,.2);border:1px solid var(--glass-border);
  border-radius:var(--r);padding:13px 17px;font-size:13px;
  color:var(--text-secondary);margin-bottom:16px;
}}
/* ── Stats ── */
.stat-row{{display:grid;grid-template-columns:repeat(3,1fr);gap:14px;margin-bottom:24px}}
.stat{{
  background:var(--glass);
  backdrop-filter:blur(20px);-webkit-backdrop-filter:blur(20px);
  border:1.5px solid var(--glass-border);border-radius:var(--r-lg);
  padding:22px 24px;text-align:center;
  box-shadow:0 4px 16px rgba(0,0,0,.07),inset 0 1.5px 0 rgba(255,255,255,.55);
  transition:transform .2s,box-shadow .2s;
}}
.stat:hover{{transform:translateY(-2px);box-shadow:0 10px 30px rgba(0,0,0,.1),inset 0 1.5px 0 rgba(255,255,255,.65)}}
.stat-num{{
  font-family:var(--font-display);font-size:36px;font-weight:700;
  color:var(--text-primary);line-height:1;margin-bottom:6px;letter-spacing:-.02em;
}}
.stat-label{{font-size:11px;color:var(--text-tertiary);letter-spacing:.07em;
             text-transform:uppercase;font-weight:600}}
/* ── Table ── */
table{{width:100%;border-collapse:collapse;font-size:13px}}
th{{
  padding:11px 18px;text-align:left;
  border-bottom:1.5px solid var(--glass-border);
  font-weight:700;color:var(--text-tertiary);
  font-size:11px;text-transform:uppercase;letter-spacing:.08em;
  background:rgba(255,255,255,.15);
}}
td{{
  padding:13px 18px;border-bottom:1px solid rgba(255,255,255,.18);
  vertical-align:middle;color:var(--text-secondary);
}}
td:first-child{{color:var(--text-primary)}}
tr:hover td{{background:rgba(255,255,255,.15)}}
tr:last-child td{{border-bottom:none}}
/* ── Tabs ── */
.tabs{{display:flex;gap:0;border-bottom:1.5px solid rgba(255,255,255,.3);margin-bottom:24px}}
.tab{{
  padding:10px 16px;font-size:13px;font-weight:600;color:var(--text-secondary);
  border-bottom:2px solid transparent;margin-bottom:-1.5px;cursor:pointer;
  transition:color .15s,border-color .15s;text-decoration:none!important;
}}
.tab:hover{{color:var(--text-primary);opacity:1}}
.tab.active{{color:var(--text-primary);border-bottom-color:var(--green);opacity:1}}
/* ── Score chips ── */
.score-chip{{font-family:var(--font-mono);font-size:20px;font-weight:500;text-align:center}}
.score-high{{color:#0c2718}}
.score-mid{{color:#6b3000}}
.score-low{{color:#7a0e0e}}
/* ── Empty state ── */
.empty-state{{text-align:center;padding:56px 24px;color:var(--text-tertiary)}}
.empty-state .empty-icon{{font-size:36px;margin-bottom:14px;opacity:.5;display:block}}
.empty-state h3{{color:var(--text-secondary);font-size:15px;margin-bottom:8px}}
.empty-state p{{font-size:13px;line-height:1.6}}
/* ── Skeleton ── */
.skeleton{{
  background:linear-gradient(90deg,rgba(255,255,255,.15) 25%,rgba(255,255,255,.3) 50%,rgba(255,255,255,.15) 75%);
  background-size:200% 100%;animation:sk-pulse 1.5s linear infinite;border-radius:var(--r);
}}
@keyframes sk-pulse{{from{{background-position:200% 0}}to{{background-position:-200% 0}}}}
/* ── Toast ── */
#toast-container{{position:fixed;bottom:24px;right:24px;z-index:9999;
                  display:flex;flex-direction:column;gap:8px;pointer-events:none}}
.toast{{
  background:rgba(12,39,24,.92);backdrop-filter:blur(16px);
  border:1px solid rgba(255,255,255,.15);
  border-radius:var(--r-lg);padding:13px 17px;font-size:13px;color:#fdf8f0;
  min-width:260px;max-width:360px;pointer-events:all;
  display:flex;align-items:center;gap:10px;
  transform:translateY(16px);opacity:0;
  transition:transform .25s var(--ease-spring),opacity .25s ease;
  box-shadow:0 16px 48px rgba(0,0,0,.3);
}}
.toast.show{{transform:translateY(0);opacity:1}}
.toast-success{{border-left:3px solid #1a5c34}}
.toast-error{{border-left:3px solid #c0392b}}
.toast-info{{border-left:3px solid rgba(255,255,255,.5)}}
.toast-warning{{border-left:3px solid #9a5c00}}
/* ── Page entry animation ── */
@keyframes fade-up{{from{{opacity:0;transform:translateY(12px)}}to{{opacity:1;transform:none}}}}
/* ── Footer ── */
footer{{
  text-align:center;padding:32px 24px;
  color:rgba(12,39,24,.5);font-size:12px;
  border-top:1px solid rgba(255,255,255,.22);
  margin-top:48px;line-height:2;
}}
footer a{{color:rgba(12,39,24,.5)}}
footer a:hover{{color:var(--text-primary)}}
/* ── Reduced motion ── */
@media(prefers-reduced-motion:reduce){{*{{animation-duration:.01ms!important;transition-duration:.01ms!important}}}}
/* ── Mobile ── */
@media(max-width:640px){{
  nav{{padding:0 16px}}
  .container,.container-lg,.container-sm{{padding:28px 16px}}
  .row,.row-3{{grid-template-columns:1fr}}
  .stat-row{{grid-template-columns:1fr}}
}}
</style>
</head>
<body>
<nav id="main-nav">
  <a class="nav-brand" href="/">
    <svg class="nav-logo" viewBox="0 0 26 26" fill="none" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <radialGradient id="nl-g" cx="50%" cy="50%" r="50%">
          <stop offset="0%" stop-color="#1a5c34" stop-opacity=".35"/>
          <stop offset="100%" stop-color="#0c2718" stop-opacity="0"/>
        </radialGradient>
        <linearGradient id="nl-r" x1="0%" y1="0%" x2="100%" y2="100%">
          <stop offset="0%" stop-color="#1a5c34"/>
          <stop offset="100%" stop-color="#0c2718"/>
        </linearGradient>
      </defs>
      <circle cx="13" cy="13" r="12.5" fill="url(#nl-g)"/>
      <circle cx="13" cy="13" r="12" stroke="rgba(12,39,24,.2)" stroke-width="1"/>
      <circle cx="13" cy="13" r="8.5" stroke="#1a5c34" stroke-width="1.3" stroke-opacity=".5">
        <animate attributeName="r" values="8.5;10;8.5" dur="3s" repeatCount="indefinite"/>
        <animate attributeName="stroke-opacity" values=".5;.1;.5" dur="3s" repeatCount="indefinite"/>
      </circle>
      <circle cx="13" cy="13" r="5.5" stroke="url(#nl-r)" stroke-width="1.6">
        <animate attributeName="r" values="5.5;6.5;5.5" dur="2.4s" repeatCount="indefinite"/>
        <animate attributeName="stroke-opacity" values="1;.3;1" dur="2.4s" repeatCount="indefinite"/>
      </circle>
      <circle cx="13" cy="13" r="2.2" fill="url(#nl-r)">
        <animate attributeName="opacity" values="1;.45;1" dur="1.8s" repeatCount="indefinite"/>
      </circle>
    </svg>
    ApplyExpress
  </a>
  <div class="nav-links">{nav_links}{avatar}</div>
</nav>
<div id="toast-container"></div>
{body}
<footer>
  <div>© {datetime.now().year} ApplyExpress &nbsp;&middot;&nbsp; Precision job hunting, automated.</div>
  <div><a href="/privacy">Privacy</a> &nbsp;&middot;&nbsp; <a href="/pricing">Pricing</a> &nbsp;&middot;&nbsp; <a href="mailto:support@applyexpress.io">Support</a></div>
</footer>
<script>
(function(){{
  var rm=window.matchMedia('(prefers-reduced-motion:reduce)').matches;
  var nav=document.getElementById('main-nav');
  if(nav)window.addEventListener('scroll',function(){{nav.classList.toggle('scrolled',scrollY>8);}},{{passive:true}});
  if(!rm){{
    var els=[].slice.call(document.querySelectorAll('.card,.stat'));
    els.forEach(function(el,i){{
      el.style.cssText+='opacity:0;transform:translateY(10px);transition:opacity .38s cubic-bezier(.16,1,.3,1) '+(i*60)+'ms,transform .38s cubic-bezier(.16,1,.3,1) '+(i*60)+'ms;';
      setTimeout(function(){{el.style.opacity='1';el.style.transform='none';}},10+(i*60));
    }});
  }}
  var seen=new Set();
  var io=new IntersectionObserver(function(entries){{
    entries.forEach(function(e){{
      if(!e.isIntersecting||seen.has(e.target))return;
      seen.add(e.target);
      var el=e.target,target=parseFloat(el.dataset.val||el.textContent)||0;
      if(target<2)return;
      var dur=rm?0:800,t0=null;
      el.dataset.val=target;el.textContent='0';
      function step(ts){{
        if(!t0)t0=ts;
        var p=Math.min((ts-t0)/dur,1),ease=1-Math.pow(1-p,3);
        el.textContent=Math.round(target*ease);
        if(p<1)requestAnimationFrame(step);else el.textContent=target;
      }}
      requestAnimationFrame(step);
    }});
  }},{{threshold:.2}});
  document.querySelectorAll('.stat-num').forEach(function(el){{
    el.dataset.val=parseFloat(el.textContent)||0;io.observe(el);
  }});
  window.showToast=function(msg,type){{
    var c=document.getElementById('toast-container');if(!c)return;
    var icons={{success:'✓',error:'✕',info:'·',warning:'⚠'}};
    var t=document.createElement('div');
    t.className='toast toast-'+(type||'info');
    t.innerHTML='<span style="font-size:14px;flex-shrink:0">'+(icons[type]||'·')+'</span><span>'+msg+'</span>';
    c.appendChild(t);
    requestAnimationFrame(function(){{requestAnimationFrame(function(){{t.classList.add('show')}});}});
    setTimeout(function(){{t.classList.remove('show');setTimeout(function(){{t.remove();}},300);}},3800);
  }};
}})();
</script>
</html>"""


# ── Landing page ──────────────────────────────────────────────────────────────

@app.route("/")
def landing():
    u = current_user()
    if u:
        return redirect(url_for("dashboard"))

    how_steps = "".join(f"""
    <div class="lp-step-card card">
      <div class="lp-step-num">{num}</div>
      <h3 class="lp-step-title">{title}</h3>
      <p class="lp-step-desc">{desc}</p>
    </div>""" for num, title, desc in [
        ("01", "Upload your CV",
         "Upload your master CV once. We use it as a base for every tailored application."),
        ("02", "Pipeline searches for you",
         "Every 4 hours, ApplyExpress scans job boards and scores each role against your profile. Only the best matches go through."),
        ("03", "AI tailors every application",
         "A unique CV with a rewritten summary and a personalised cover letter is generated for each role."),
        ("04", "Extension applies automatically",
         "The Chrome extension submits your applications on Reed, LinkedIn, Indeed and more — then logs every result to your dashboard."),
    ])

    features = "".join(f"""
    <div class="lp-feat">
      <div class="lp-feat-icon">{icon}</div>
      <div>
        <div class="lp-feat-title">{title}</div>
        <div class="lp-feat-desc">{desc}</div>
      </div>
    </div>""" for icon, title, desc in [
        ("""<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10"/><path d="M12 8v4l3 3"/></svg>""",
         "Smart job scoring",
         "Every job is scored 1–10 against your profile. Only roles above your threshold get a tailored application."),
        ("""<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>""",
         "Tailored CV per job",
         "Your CV summary is rewritten to match each job description. Never send the same CV twice."),
        ("""<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M4 4h16c1.1 0 2 .9 2 2v12c0 1.1-.9 2-2 2H4c-1.1 0-2-.9-2-2V6c0-1.1.9-2 2-2z"/><polyline points="22,6 12,13 2,6"/></svg>""",
         "AI cover letters",
         "Personalised 3-paragraph cover letters, written from your profile — not a generic template."),
        ("""<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="2" y="3" width="20" height="14" rx="2"/><path d="M8 21h8M12 17v4"/></svg>""",
         "Fully autonomous",
         "Pipeline runs 3&times; per day without you lifting a finger. You get an email when it's done."),
        ("""<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="3" y="11" width="18" height="11" rx="2" ry="2"/><path d="M7 11V7a5 5 0 0110 0v4"/></svg>""",
         "Session-based applying",
         "Applications go through your own browser — no server-side bots, no ToS grey areas."),
        ("""<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/></svg>""",
         "Application dashboard",
         "Track every application: Applied, Failed, date, role, company. Searchable and filterable."),
        ("""<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="4.93" y1="4.93" x2="19.07" y2="19.07"/></svg>""",
         "Company blacklist",
         "Never want to apply to a specific company? Add them to your blacklist. Skipped automatically."),
        ("""<svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M12 22s8-4 8-10V5l-8-3-8 3v7c0 6 8 10 8 10z"/></svg>""",
         "Multi-ATS support",
         "Reed, LinkedIn, Indeed, Totaljobs, Greenhouse, Lever, Workable and Ashby — all handled."),
    ])

    faqs = "".join(f"""
    <details class="lp-faq">
      <summary class="lp-faq-q">{q} <span class="lp-faq-plus">+</span></summary>
      <div class="lp-faq-a">{a}</div>
    </details>""" for q, a in [
        ("Does this violate job board Terms of Service?",
         "ApplyExpress submits applications through your own Chrome browser using your own logged-in session — the same as if you clicked Apply yourself. The extension acts as an automation assistant on your machine, not a server-side bot."),
        ("Do I need to be at my computer for it to work?",
         "For the pipeline (fetching jobs, tailoring CVs) — no, that runs on our server automatically. For submitting applications — yes, Chrome needs to be open. Most users let it run in the background while they work."),
        ("Which job sites are supported?",
         "Reed, LinkedIn Easy Apply, Indeed, Totaljobs, Greenhouse, Lever, Workable and Ashby. More are added regularly."),
        ("Can I control which jobs it applies to?",
         "Yes. You set a match-score threshold (1–10). Only jobs that score above your threshold get a tailored application. You can also blacklist specific companies."),
        ("What happens after the free trial?",
         "After 14 days you'll be prompted to subscribe at &pound;19/month. If you don't subscribe, the pipeline pauses — your data is kept safe."),
    ])

    body = f"""
    <style>
    /* Landing-page overrides — nav is already rendered by _page */
    .lp-hero{{padding:96px 20px 80px;text-align:center;position:relative;overflow:hidden}}
    .lp-hero::before{{content:'';position:absolute;inset:0;
      background:radial-gradient(ellipse at 30% 40%,rgba(255,255,255,.07) 0%,transparent 55%),
                 radial-gradient(ellipse at 75% 70%,rgba(12,39,24,.25) 0%,transparent 50%);
      pointer-events:none}}
    .lp-pill{{display:inline-block;background:rgba(12,39,24,.2);border:1.5px solid rgba(26,92,52,.4);
      padding:6px 18px;border-radius:20px;font-size:11px;font-weight:700;
      letter-spacing:.08em;margin-bottom:28px;color:var(--green-mid);text-transform:uppercase}}
    .lp-hero h1{{font-family:'Clash Display',sans-serif;font-size:clamp(36px,6vw,58px);font-weight:800;
      color:var(--text-primary);line-height:1.1;margin-bottom:20px;
      max-width:680px;margin-left:auto;margin-right:auto}}
    .lp-hero h1 em{{font-style:normal;color:var(--green-mid)}}
    .lp-hero-sub{{font-size:17px;color:var(--text-secondary);max-width:500px;margin:0 auto 36px;line-height:1.65}}
    .lp-cta-row{{display:flex;gap:12px;justify-content:center;flex-wrap:wrap;margin-bottom:20px}}
    .lp-btn-primary{{background:linear-gradient(145deg,var(--green-mid),var(--green));color:#fdf8f0;
      padding:15px 36px;border-radius:12px;font-weight:700;font-size:15px;text-decoration:none;
      display:inline-block;box-shadow:0 4px 16px rgba(12,39,24,.35),inset 0 1px 0 rgba(255,255,255,.12);
      transition:all .15s ease}}
    .lp-btn-primary:hover{{transform:translateY(-1px);box-shadow:0 6px 20px rgba(12,39,24,.45)}}
    .lp-btn-ghost{{background:rgba(255,255,255,.16);border:1.5px solid rgba(255,255,255,.32);
      color:var(--text-primary);padding:14px 28px;border-radius:12px;font-weight:600;
      font-size:15px;text-decoration:none;display:inline-block;
      backdrop-filter:blur(8px);transition:all .15s ease}}
    .lp-btn-ghost:hover{{background:rgba(255,255,255,.24)}}
    .lp-hero-note{{font-size:12px;color:var(--text-tertiary);margin-top:4px}}
    .lp-proof-bar{{background:rgba(12,39,24,.12);border-top:1.5px solid rgba(12,39,24,.15);
      border-bottom:1.5px solid rgba(12,39,24,.15);padding:16px 20px;text-align:center}}
    .lp-proof-bar p{{font-size:13px;color:var(--text-secondary);margin:0}}
    .lp-proof-bar strong{{color:var(--text-primary)}}
    .lp-section{{max-width:980px;margin:0 auto;padding:64px 20px 0}}
    .lp-section-head{{text-align:center;margin-bottom:36px}}
    .lp-section-head h2{{font-family:'Clash Display',sans-serif;font-size:clamp(26px,4vw,34px);
      font-weight:700;color:var(--text-primary);margin:0 0 10px}}
    .lp-section-head p{{font-size:15px;color:var(--text-secondary);margin:0}}
    .lp-steps-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(210px,1fr));gap:18px;margin-bottom:64px}}
    .lp-step-card{{text-align:center;padding:30px 24px}}
    .lp-step-num{{font-family:'Clash Display',sans-serif;font-size:38px;font-weight:800;
      color:var(--green-mid);opacity:.55;margin-bottom:14px;line-height:1}}
    .lp-step-title{{font-family:'Clash Display',sans-serif;font-size:16px;font-weight:700;
      color:var(--text-primary);margin:0 0 10px}}
    .lp-step-desc{{font-size:13px;color:var(--text-secondary);line-height:1.65;margin:0}}
    .lp-feats-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(270px,1fr));gap:14px;margin-bottom:64px}}
    .lp-feat{{display:flex;gap:14px;align-items:flex-start;padding:20px;
      background:var(--glass);backdrop-filter:blur(16px) saturate(1.3);
      border:1.5px solid var(--glass-border);border-radius:14px;
      box-shadow:0 4px 16px rgba(0,0,0,.05),inset 0 1px 0 rgba(255,255,255,.45)}}
    .lp-feat-icon{{flex-shrink:0;width:36px;height:36px;border-radius:9px;
      background:rgba(12,39,24,.15);border:1.5px solid rgba(26,92,52,.25);
      display:flex;align-items:center;justify-content:center;color:var(--green-mid)}}
    .lp-feat-title{{font-weight:700;font-size:14px;color:var(--text-primary);margin-bottom:4px}}
    .lp-feat-desc{{font-size:13px;color:var(--text-secondary);line-height:1.55}}
    .lp-cta-card{{background:rgba(12,39,24,.15);border:1.5px solid rgba(26,92,52,.3);
      border-radius:20px;padding:48px 32px;text-align:center;margin-bottom:64px;
      box-shadow:inset 0 1.5px 0 rgba(255,255,255,.25)}}
    .lp-cta-price{{font-family:'Clash Display',sans-serif;font-size:48px;font-weight:800;
      color:var(--text-primary);margin-bottom:8px;line-height:1}}
    .lp-cta-sub{{font-size:15px;color:var(--text-secondary);margin-bottom:28px}}
    .lp-cta-note{{font-size:12px;color:var(--text-tertiary);margin-top:16px}}
    .lp-faq-wrap{{max-width:680px;margin:0 auto;padding-bottom:64px}}
    .lp-faq{{border:1.5px solid var(--glass-border);border-radius:12px;overflow:hidden;margin-bottom:10px;
      background:var(--glass);backdrop-filter:blur(14px)}}
    .lp-faq-q{{padding:16px 20px;cursor:pointer;font-weight:600;font-size:14px;color:var(--text-primary);
      list-style:none;display:flex;justify-content:space-between;align-items:center;
      user-select:none}}
    .lp-faq-q:hover{{background:rgba(255,255,255,.1)}}
    .lp-faq-plus{{color:var(--text-tertiary);font-size:18px;font-weight:400}}
    .lp-faq-a{{padding:0 20px 16px;font-size:14px;color:var(--text-secondary);line-height:1.65}}
    .lp-footer{{background:rgba(12,39,24,.25);border-top:1.5px solid rgba(12,39,24,.2);
      padding:32px 20px;text-align:center;font-size:13px;color:var(--text-tertiary)}}
    .lp-footer a{{color:var(--text-secondary);margin:0 12px;text-decoration:none}}
    .lp-footer a:hover{{color:var(--text-primary)}}
    .lp-footer p{{margin:4px 0}}
    </style>

    <!-- Hero -->
    <div class="lp-hero">
      <div class="lp-pill">Free 14-day trial &middot; No card required</div>
      <h1>Apply to jobs on autopilot.<br><em>Land interviews faster.</em></h1>
      <p class="lp-hero-sub">
        ApplyExpress searches Reed, LinkedIn, Indeed and Totaljobs, scores every role against
        your profile, writes a tailored CV and cover letter, then applies automatically &mdash; 3&times; a day.
      </p>
      <div class="lp-cta-row">
        <a href="/signup" class="lp-btn-primary">Start free trial &#8594;</a>
        <a href="/pricing" class="lp-btn-ghost">See pricing</a>
      </div>
      <p class="lp-hero-note">Works on Reed &middot; LinkedIn &middot; Indeed &middot; Totaljobs &middot; Greenhouse &middot; Lever &middot; Workable &middot; Ashby</p>
    </div>

    <!-- Proof bar -->
    <div class="lp-proof-bar">
      <p><strong>Automates the copy-paste grind</strong> &nbsp;&middot;&nbsp;
         Tailored CV for every single application &nbsp;&middot;&nbsp;
         Submitted from your own browser, using your own session</p>
    </div>

    <!-- How it works -->
    <div class="lp-section">
      <div class="lp-section-head">
        <h2>How it works</h2>
        <p>Set up once. Runs automatically from then on.</p>
      </div>
      <div class="lp-steps-grid">{how_steps}</div>

      <!-- Features -->
      <div class="lp-section-head">
        <h2>Everything included</h2>
      </div>
      <div class="lp-feats-grid">{features}</div>

      <!-- Pricing CTA -->
      <div class="lp-cta-card">
        <div class="lp-cta-price">&pound;19<span style="font-size:20px;color:var(--text-secondary)">/month</span></div>
        <p class="lp-cta-sub">14-day free trial. Cancel any time. No hidden fees.</p>
        <a href="/signup" class="lp-btn-primary" style="font-size:16px;padding:16px 44px">
          Start free trial &mdash; no card needed &#8594;
        </a>
        <p class="lp-cta-note">Join job seekers applying smarter, not harder.</p>
      </div>

      <!-- FAQ -->
      <div class="lp-section-head">
        <h2>Questions</h2>
      </div>
      <div class="lp-faq-wrap">{faqs}</div>
    </div>

    <!-- Footer -->
    <div class="lp-footer">
      <p>
        <a href="/privacy">Privacy Policy</a>
        <a href="/pricing">Pricing</a>
        <a href="mailto:support@applyexpress.io">Support</a>
      </p>
      <p>&copy; 2026 ApplyExpress &middot; Built to get you hired faster</p>
    </div>"""

    return _page("ApplyExpress — AI Job Applications on Autopilot", body,
                 extra_head='<meta name="description" content="ApplyExpress automatically applies to jobs on Reed, LinkedIn, Indeed and Totaljobs — with a tailored CV and cover letter for every role. 14-day free trial.">')


# ── Auth ──────────────────────────────────────────────────────────────────────

@app.route("/signup", methods=["GET", "POST"])
@limiter.limit("10 per hour")
def signup():
    error = ""
    if request.method == "POST":
        first  = request.form.get("first_name", "").strip()
        last   = request.form.get("last_name",  "").strip()
        email  = request.form.get("email",  "").strip().lower()
        pwd    = request.form.get("password", "")
        keywords = request.form.get("keywords", "").strip()
        location = request.form.get("location", "London").strip()

        consented = request.form.get("consent") == "1"
        if not all([first, last, email, pwd]):
            error = "All fields are required."
        elif len(pwd) < 6:
            error = "Password must be at least 6 characters."
        elif not consented:
            error = "Please confirm that you authorise ApplyExpress to submit applications on your behalf."
        else:
            db  = get_db()
            exists = db.execute("SELECT id FROM users WHERE email=?", (email,)).fetchone()
            if exists:
                error = "An account with that email already exists."
            else:
                api_key    = f"aa-{secrets.token_urlsafe(24)}"
                trial_ends = (datetime.utcnow() + timedelta(days=14)).isoformat()
                now        = datetime.utcnow().isoformat()
                db.execute("""
                    INSERT INTO users
                      (email, password_hash, first_name, last_name, api_key,
                       keywords, search_location, email_subject, created_at,
                       trial_ends_at, consented_at)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?)
                """, (email, _hash(pwd), first, last, api_key,
                      keywords or "compliance analyst", location,
                      (keywords.split(",")[0].strip().title() if keywords else "Job Application"),
                      now, trial_ends, now))
                db.commit()
                row = db.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
                session["user_id"] = row["id"]
                # Send welcome email
                _send_email(email, "Welcome to ApplyExpress 🎉",
                    f"<h2>Welcome, {first}!</h2>"
                    f"<p>Your 14-day free trial has started. Complete your setup to start applying automatically.</p>"
                    f"<p><a href='{request.host_url}setup'>Complete setup →</a></p>")
                return redirect(url_for("setup") + "?step=1")

    form = f"""
<style>
.auth-wrap{{min-height:calc(100vh - 62px);display:flex;align-items:center;
  justify-content:center;padding:40px 20px}}
.auth-card{{width:100%;max-width:480px}}
.auth-title{{font-family:var(--font-display);font-size:30px;font-weight:700;
  letter-spacing:-.02em;color:var(--text-primary);margin-bottom:6px}}
.auth-sub{{font-size:14px;color:var(--text-secondary);margin-bottom:28px}}
.auth-footer{{font-size:13px;color:var(--text-secondary);text-align:center;margin-top:20px}}
.consent-row{{display:flex;gap:10px;align-items:flex-start;margin-top:8px}}
.consent-row input{{width:16px;height:16px;margin-top:3px;flex-shrink:0}}
.consent-row label{{font-size:12px;color:var(--text-secondary);line-height:1.6}}
.trial-note{{font-size:11px;color:var(--text-tertiary);text-align:center;margin-top:10px;letter-spacing:.02em}}
</style>
<div class="auth-wrap">
  <div class="auth-card">
    <div class="card">
      <h1 class="auth-title">Create your account</h1>
      <p class="auth-sub">14-day free trial &middot; No credit card required</p>
      {"<div class='alert alert-error'>"+error+"</div>" if error else ""}
      <form method="POST">
        <div class="row">
          <div class="field"><label class="label">First name</label>
            <input type="text" name="first_name" required placeholder="Alice" autofocus></div>
          <div class="field"><label class="label">Last name</label>
            <input type="text" name="last_name" required placeholder="Smith"></div>
        </div>
        <div class="field"><label class="label">Email address</label>
          <input type="email" name="email" required placeholder="alice@gmail.com"></div>
        <div class="field"><label class="label">Password</label>
          <input type="password" name="password" required placeholder="At least 6 characters"></div>
        <div class="field"><label class="label">Job titles you want
          <span style="font-weight:400;opacity:.6">(comma-separated)</span></label>
          <input type="text" name="keywords" placeholder="e.g. AML Compliance, MLRO, Risk Analyst"></div>
        <div class="field"><label class="label">Location</label>
          <input type="text" name="location" value="London"></div>
        <div class="consent-row">
          <input type="checkbox" name="consent" value="1" id="consent" required>
          <label for="consent">
            I authorise ApplyExpress to submit job applications on my behalf.
            I agree to the <a href="/privacy" target="_blank">Privacy Policy</a>.
          </label>
        </div>
        <button type="submit" class="btn btn-primary btn-full" style="margin-top:20px;font-size:15px;padding:13px">
          Start free trial &rarr;
        </button>
        <p class="trial-note">14 days free &middot; Cancel any time</p>
      </form>
      <p class="auth-footer">Already have an account? <a href="/login">Log in</a></p>
    </div>
  </div>
</div>"""
    return _page("Sign up", form)


@app.route("/login", methods=["GET", "POST"])
@limiter.limit("20 per hour")
def login_page():
    error = ""
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        pwd   = request.form.get("password", "")
        db    = get_db()
        row   = db.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
        if row and _verify_password(row["password_hash"], pwd):
            # Transparently upgrade legacy SHA-256 hashes to salted PBKDF2
            if not (":" in row["password_hash"] or row["password_hash"].startswith(("pbkdf2", "scrypt"))):
                db.execute("UPDATE users SET password_hash=? WHERE id=?",
                           (_hash(pwd), row["id"]))
                db.commit()
            session["user_id"] = row["id"]
            return redirect(url_for("dashboard"))
        error = "Invalid email or password."

    form = f"""
<style>
.auth-wrap{{min-height:calc(100vh - 62px);display:flex;align-items:center;
  justify-content:center;padding:40px 20px}}
.auth-card{{width:100%;max-width:420px}}
.auth-title{{font-family:var(--font-display);font-size:30px;font-weight:700;
  letter-spacing:-.02em;color:var(--text-primary);margin-bottom:6px}}
.auth-sub{{font-size:14px;color:var(--text-secondary);margin-bottom:28px}}
.auth-footer{{font-size:13px;color:var(--text-secondary);text-align:center;margin-top:20px}}
</style>
<div class="auth-wrap">
  <div class="auth-card">
    <div class="card">
      <h1 class="auth-title">Welcome back</h1>
      <p class="auth-sub">Log in to your ApplyExpress account</p>
      {"<div class='alert alert-error'>"+error+"</div>" if error else ""}
      <form method="POST">
        <div class="field"><label class="label">Email address</label>
          <input type="email" name="email" required autofocus placeholder="alice@gmail.com"></div>
        <div class="field"><label class="label">Password</label>
          <input type="password" name="password" required placeholder="Your password"></div>
        <button type="submit" class="btn btn-primary btn-full" style="font-size:15px;padding:13px;margin-top:4px">
          Log in &rarr;
        </button>
      </form>
      <p class="auth-footer">No account yet? <a href="/signup">Start free trial</a></p>
    </div>
  </div>
</div>"""
    return _page("Log in", form)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("landing"))


# ── One-click extension connect ───────────────────────────────────────────────

@app.route("/connect/<token>")
def connect_extension(token):
    """Opened by dashboard button. content.js detects the meta tag and auto-configures."""
    db  = get_db()
    row = db.execute("SELECT id FROM users WHERE api_key=?", (token,)).fetchone()
    if not row:
        return "Invalid link.", 403
    server_url = request.host_url.rstrip("/")
    return f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<meta name="applyexpress-connect" data-server-url="{server_url}" data-api-key="{token}">
<title>Connecting ApplyExpress Extension…</title>
<style>body{{font-family:-apple-system,sans-serif;display:flex;flex-direction:column;
align-items:center;justify-content:center;height:100vh;background:#eff6ff;text-align:center;margin:0}}</style>
</head><body>
<div style="font-size:48px;margin-bottom:16px">⏳</div>
<h2 id="connect-status" style="color:#1e40af">Connecting extension…</h2>
<p id="connect-hint" style="color:#64748b">Make sure the ApplyExpress extension is installed in Chrome.</p>
<script>
// Fallback: if content.js hasn't replaced the page after 12 s,
// the extension probably isn't installed or didn't load.
setTimeout(function() {{
  var s = document.getElementById('connect-status');
  var h = document.getElementById('connect-hint');
  if (s) {{
    s.style.color = '#b45309';
    s.textContent = 'Extension not detected';
    if (h) h.innerHTML =
      'Make sure the ApplyExpress extension is loaded in Chrome (<b>chrome://extensions</b>), ' +
      'then <a href="" style="color:#2563eb">reload this page</a>.';
  }}
}}, 12000);
</script>
</body></html>"""


# ── Pricing ───────────────────────────────────────────────────────────────────

@app.route("/pricing")
def pricing():
    u    = current_user()
    gate = request.args.get("gate")
    gate_html = (
        '<div class="alert" style="background:#fef3c7;border:1px solid #f59e0b;color:#92400e;'
        'padding:12px 20px;border-radius:8px;margin-bottom:24px">'
        '⚠ Your trial has ended or you need an active plan to run the pipeline.</div>'
    ) if gate else ""

    active = _is_active(u) if u else False
    plan_badge = ""
    if u and active:
        label = "Paid" if u.get("is_paid") else f"Trial ends {(u.get('trial_ends_at') or '')[:10]}"
        plan_badge = f'<div class="alert" style="background:#dcfce7;border:1px solid #86efac;color:#15803d;padding:12px 20px;border-radius:8px;margin-bottom:24px">✓ Active plan: {label}</div>'

    cta = (
        "<a href='/subscribe' class='btn btn-primary btn-lg' style='font-size:16px;padding:14px 44px'>"
        "Start free trial &rarr;</a>"
        if not active else
        "<div class='badge badge-green' style='font-size:13px;padding:8px 18px'>&#10003; You have an active plan</div>"
    )
    features = [
        ("&#128269;", "Smart job scoring", "Every role scored 1–10 against your profile. Only the best matches get a tailored application."),
        ("&#128196;", "AI-tailored CV per job", "Your CV summary is rewritten to match each job description. Never send the same one twice."),
        ("&#128140;", "Personalised cover letters", "3-paragraph cover letters written from your profile — not generic templates."),
        ("&#129302;", "Fully autonomous pipeline", "Runs 3&times; per day. You get an email when it's done. No manual work."),
        ("&#128279;", "Multi-ATS support", "Reed, LinkedIn, Indeed, Totaljobs, Greenhouse, Lever, Workable, Ashby."),
        ("&#128202;", "Application dashboard", "Track every application: role, company, status, date. Searchable."),
    ]
    feat_html = "".join(f"""
    <div style="display:flex;gap:14px;align-items:flex-start">
      <div style="font-size:20px;flex-shrink:0;margin-top:2px">{icon}</div>
      <div>
        <div style="font-weight:700;font-size:14px;color:var(--text-primary);margin-bottom:3px">{title}</div>
        <div style="font-size:13px;color:var(--text-secondary);line-height:1.6">{desc}</div>
      </div>
    </div>""" for icon, title, desc in features)

    body = f"""
<style>
.pricing-wrap{{max-width:720px;margin:0 auto;padding:48px 24px 64px;
  animation:fade-up .4s var(--ease-spring) both}}
.pricing-hero{{text-align:center;margin-bottom:40px}}
.pricing-hero h1{{font-size:38px;margin-bottom:10px}}
.price-card{{text-align:center;padding:0;overflow:hidden;margin-bottom:24px}}
.price-header{{
  background:rgba(12,39,24,.1);border-bottom:1.5px solid rgba(255,255,255,.3);
  padding:14px 28px;
}}
.price-tag{{font-family:var(--font-display);font-size:52px;font-weight:700;
  letter-spacing:-.03em;color:var(--text-primary);line-height:1;margin:28px 0 8px}}
.price-period{{font-size:16px;font-weight:400;color:var(--text-secondary)}}
.price-sub{{font-size:14px;color:var(--text-secondary);margin-bottom:28px}}
.feat-grid{{display:grid;grid-template-columns:1fr 1fr;gap:20px;
  text-align:left;padding:0 28px 28px}}
.price-cta{{padding:0 28px 32px}}
.price-note{{font-size:12px;color:var(--text-tertiary);margin-top:12px}}
@media(max-width:560px){{.feat-grid{{grid-template-columns:1fr}}}}
</style>
<div class="pricing-wrap">
  {gate_html}{plan_badge}
  <div class="pricing-hero">
    <h1>Simple, honest pricing</h1>
    <p>14-day free trial. Cancel any time. No hidden fees.</p>
  </div>

  <div class="card price-card">
    <div class="price-header">
      <span style="font-size:11px;font-weight:700;color:var(--text-primary);
        text-transform:uppercase;letter-spacing:.1em">&#127881; Most popular</span>
    </div>
    <div class="price-tag">£19<span class="price-period">/mo</span></div>
    <p class="price-sub">Everything you need to land your next role faster</p>

    <div class="feat-grid">{feat_html}</div>

    <div class="price-cta">
      {cta}
      <p class="price-note">14 days free, then £19/month. Cancel any time from your account.</p>
    </div>
  </div>

  <div style="text-align:center;margin-top:20px">
    <p style="font-size:13px;color:var(--text-secondary)">
      Questions? <a href="mailto:support@applyexpress.io">support@applyexpress.io</a>
    </p>
  </div>
</div>"""
    return _page("Pricing", body, user=u)


@app.route("/subscribe")
@login_required
def subscribe():
    u = current_user()
    if not stripe:
        return "Stripe not configured — set STRIPE_SECRET_KEY in .env", 500

    price_id = os.getenv("STRIPE_PRICE_ID", "")
    if not price_id:
        return "STRIPE_PRICE_ID not set in .env", 500

    try:
        checkout = stripe.checkout.Session.create(
            payment_method_types=["card"],
            mode="subscription",
            customer_email=u["email"],
            line_items=[{"price": price_id, "quantity": 1}],
            subscription_data={"trial_period_days": 14},
            success_url=request.host_url + "dashboard?subscribed=1",
            cancel_url=request.host_url + "pricing",
            metadata={"user_id": str(u["id"])},
        )
        return redirect(checkout.url)
    except Exception as e:
        log.error(f"Stripe checkout error: {e}")
        return f"Payment error: {e}", 500


@app.route("/stripe/webhook", methods=["POST"])
@limiter.exempt
def stripe_webhook():
    if not stripe:
        return "Stripe not configured", 400

    payload    = request.get_data()
    sig_header = request.headers.get("Stripe-Signature", "")
    wh_secret  = os.getenv("STRIPE_WEBHOOK_SECRET", "")

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, wh_secret)
    except Exception as e:
        log.warning(f"Stripe webhook signature error: {e}")
        return jsonify({"error": str(e)}), 400

    db = sqlite3.connect(str(DB_PATH))
    db.row_factory = sqlite3.Row

    evt_type = event["type"]
    log.info(f"Stripe event: {evt_type}")

    if evt_type in ("customer.subscription.created", "customer.subscription.updated",
                    "invoice.payment_succeeded"):
        sub  = event["data"]["object"]
        # customer.subscription.* has customer field; invoice has customer too
        cust_id = sub.get("customer") or sub.get("customer_id", "")
        sub_id  = sub.get("id") or sub.get("subscription", "")
        # Find user by stripe_customer or by metadata on checkout session
        row = db.execute("SELECT id FROM users WHERE stripe_customer=?", (cust_id,)).fetchone()
        if not row:
            # Try metadata lookup via checkout session
            uid_meta = (event.get("data", {}).get("object", {})
                        .get("metadata", {}).get("user_id", ""))
            if uid_meta:
                row = db.execute("SELECT id FROM users WHERE id=?", (uid_meta,)).fetchone()
        if row:
            db.execute(
                "UPDATE users SET is_paid=1, stripe_customer=?, stripe_sub_id=? WHERE id=?",
                (cust_id, sub_id, row["id"])
            )
            db.commit()
            log.info(f"User {row['id']} marked as paid (sub={sub_id})")

    elif evt_type in ("customer.subscription.deleted", "invoice.payment_failed"):
        sub     = event["data"]["object"]
        cust_id = sub.get("customer", "")
        db.execute("UPDATE users SET is_paid=0 WHERE stripe_customer=?", (cust_id,))
        db.commit()
        log.info(f"Subscription cancelled/failed for customer {cust_id}")

    # Handle first checkout: link customer to user via metadata
    elif evt_type == "checkout.session.completed":
        sess    = event["data"]["object"]
        uid     = sess.get("metadata", {}).get("user_id", "")
        cust_id = sess.get("customer", "")
        sub_id  = sess.get("subscription", "")
        if uid and cust_id:
            db.execute(
                "UPDATE users SET stripe_customer=?, stripe_sub_id=? WHERE id=?",
                (cust_id, sub_id, uid)
            )
            db.commit()

    db.close()
    return jsonify({"ok": True})


@app.route("/cancel-subscription")
@login_required
def cancel_subscription():
    u = current_user()
    if not stripe or not u.get("stripe_sub_id"):
        return redirect(url_for("dashboard"))
    try:
        stripe.Subscription.modify(u["stripe_sub_id"], cancel_at_period_end=True)
        return _page("Subscription cancelled",
            '<div class="container"><div class="card"><h2>Subscription cancelled</h2>'
            '<p style="margin-top:12px;color:#64748b">Your access continues until the end of the current billing period.</p>'
            '<p style="margin-top:16px"><a href="/dashboard">← Back to dashboard</a></p></div></div>', user=u)
    except Exception as e:
        return f"Error: {e}", 500


# ── Dashboard ─────────────────────────────────────────────────────────────────

@app.route("/dashboard")
@login_required
def dashboard():
    u    = current_user()
    db   = get_db()
    runs = db.execute(
        "SELECT * FROM runs WHERE user_id=? ORDER BY started_at DESC LIMIT 10",
        (u["id"],)
    ).fetchall()

    welcome = '<div class="alert alert-success">🎉 Welcome! Upload your CV and set your preferences to get started.</div>' \
              if request.args.get("welcome") else ""

    # Stats
    total_applied = sum(r["jobs_applied"] for r in runs)
    total_found   = sum(r["jobs_found"]   for r in runs)
    run_count     = len(runs)

    # stats rendered inline in body below

    # CV status
    cv_path = _user_dir(u["id"]) / "cv.docx"
    cv_status = (
        f'<span class="badge badge-green">✓ CV uploaded</span>'
        if cv_path.exists() else
        f'<span class="badge badge-red">No CV — <a href="/profile#cv">upload one</a></span>'
    )

    # Run button — with specialty picker if user has specialties
    dash_specialties = db.execute(
        "SELECT * FROM specialties WHERE user_id=? ORDER BY name", (u["id"],)).fetchall()
    spec_options = '<option value="">— Default CV —</option>' + \
        "".join(f'<option value="{s["id"]}">{s["name"]}</option>' for s in dash_specialties)
    run_btn = f"""
    <form action="/run" method="POST" style="display:flex;gap:10px;align-items:center;flex-wrap:wrap">
      <select name="specialty_id">{spec_options}</select>
      <button type="submit" class="btn btn-success btn-lg">&#9654;&nbsp; Run now</button>
    </form>"""

    # Last run status
    last_run_badge = ""
    if runs:
        last = runs[0]
        status_map = {"running":"badge-blue","completed":"badge-green","failed":"badge-red","interrupted":"badge-yellow"}
        cls = status_map.get(last["status"], "badge-gray")
        last_run_badge = f'Last run: <span class="badge {cls}">{last["status"]}</span> · {last["started_at"][:16]}'

    # Runs table
    rows_html = ""
    for r in runs:
        status_map = {"running":"badge-blue","completed":"badge-green","failed":"badge-red","interrupted":"badge-yellow"}
        cls = status_map.get(r["status"], "badge-gray")
        applied_pct = int((r["jobs_applied"] / r["jobs_found"] * 100)) if r["jobs_found"] else 0
        score_bar = f"""<div class="score-bar-wrap">
          <div class="score-bar-track"><div class="score-bar-fill" style="width:{applied_pct}%"></div></div>
          <span class="score-bar-val">{r["jobs_applied"]}</span>
        </div>""" if r["jobs_found"] else f'<span style="color:var(--text-tertiary);font-size:12px">—</span>'
        view_link = f'<a href="/run/{r["id"]}" style="font-size:12px;font-weight:700;color:rgba(12,39,24,.7)">View &rarr;</a>' if r["report_json"] else ""
        rows_html += f"""<tr>
          <td style="font-family:var(--font-mono);font-size:12px">{r["started_at"][:16]}</td>
          <td><span class="badge {cls}">{r["status"].capitalize()}</span></td>
          <td style="color:var(--text-secondary)">{r["jobs_found"]}</td>
          <td>{score_bar}</td>
          <td>{view_link}</td>
        </tr>"""
    empty_row = '<tr><td colspan="5" style="text-align:center;padding:32px;color:var(--text-tertiary)">No runs yet — click <strong style="color:var(--text-primary)">Run now</strong> to start.</td></tr>'
    runs_table = f"""
    <table>
      <thead><tr>
        <th>Date &amp; time</th><th>Status</th><th>Found</th><th>Applied</th><th></th>
      </tr></thead>
      <tbody>{rows_html if rows_html else empty_row}</tbody>
    </table>"""

    body = f"""
<style>
.dash-wrap{{max-width:1100px;margin:0 auto;padding:40px 28px 64px;
  animation:fade-up .4s var(--ease-spring) both}}
.dash-hero{{display:flex;align-items:flex-start;justify-content:space-between;
  flex-wrap:wrap;gap:20px;margin-bottom:32px}}
.dash-greeting{{font-family:var(--font-display);font-size:36px;font-weight:700;
  letter-spacing:-.025em;color:var(--text-primary);margin-bottom:8px}}
.dash-subline{{display:flex;align-items:center;gap:10px;font-size:13px;
  color:var(--text-secondary);flex-wrap:wrap}}
.dash-dot{{width:8px;height:8px;border-radius:50%;flex-shrink:0;background:#0c2718;
  box-shadow:0 0 0 3px rgba(12,39,24,.18);animation:ddot 2.5s ease infinite}}
@keyframes ddot{{0%,100%{{box-shadow:0 0 0 3px rgba(12,39,24,.18)}}
  50%{{box-shadow:0 0 0 8px rgba(12,39,24,.05)}}}}
.run-form{{display:flex;gap:10px;align-items:center;flex-wrap:wrap}}
.run-form select{{width:auto;min-width:180px}}
.stat-grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-bottom:28px}}
.stat-card{{padding:26px 24px}}
.stat-card .stat-icon{{font-size:20px;display:block;margin-bottom:14px}}
.stat-card .stat-num{{font-family:var(--font-display);font-size:44px;font-weight:700;
  letter-spacing:-.03em;line-height:1;margin-bottom:6px}}
.stat-card .stat-label{{font-size:11px;color:var(--text-tertiary);
  text-transform:uppercase;letter-spacing:.08em;font-weight:600}}
.stat-delta{{display:inline-flex;align-items:center;gap:4px;margin-top:12px;
  font-size:12px;font-weight:700;padding:3px 10px;border-radius:20px;
  background:rgba(12,39,24,.12);color:var(--text-primary)}}
.section-row{{display:flex;align-items:center;justify-content:space-between;margin-bottom:14px}}
.section-title{{font-family:var(--font-display);font-size:17px;font-weight:700;
  color:var(--text-primary);letter-spacing:-.01em}}
.section-link{{font-size:12px;font-weight:700;color:rgba(12,39,24,.6);letter-spacing:.02em}}
.section-link:hover{{color:var(--text-primary)}}
.table-wrap{{padding:0;overflow:hidden;margin-bottom:24px}}
.score-bar-wrap{{display:flex;align-items:center;gap:10px}}
.score-bar-track{{flex:1;max-width:80px;height:5px;background:rgba(12,39,24,.1);
  border-radius:3px;overflow:hidden}}
.score-bar-fill{{height:100%;border-radius:3px;
  background:linear-gradient(90deg,#1a5c34,#0c2718);
  transition:width .9s var(--ease-spring)}}
.score-bar-val{{font-family:var(--font-mono);font-size:12px;
  font-weight:700;color:#0c2718;min-width:28px}}
.ext-card{{padding:34px;display:flex;align-items:center;gap:28px;flex-wrap:wrap}}
.ext-icon{{width:68px;height:68px;border-radius:18px;font-size:30px;flex-shrink:0;
  background:rgba(12,39,24,.1);border:1.5px solid rgba(12,39,24,.15);
  display:flex;align-items:center;justify-content:center;
  box-shadow:inset 0 1px 0 rgba(255,255,255,.3)}}
.ext-body{{flex:1;min-width:220px}}
.ext-title{{font-family:var(--font-display);font-size:20px;font-weight:700;
  color:var(--text-primary);margin-bottom:10px}}
.ext-desc{{font-size:13px;color:var(--text-secondary);margin-bottom:22px;
  max-width:480px;line-height:1.7}}
.ext-status{{text-align:right;flex-shrink:0}}
.ext-status-lbl{{font-size:11px;color:var(--text-tertiary);text-transform:uppercase;
  letter-spacing:.07em;margin-bottom:8px}}
@media(max-width:680px){{
  .stat-grid{{grid-template-columns:1fr}}
  .dash-hero{{flex-direction:column}}
  .ext-card{{flex-direction:column;gap:20px;padding:24px}}
  .dash-wrap{{padding:28px 16px 56px}}
}}
</style>

<div class="dash-wrap">
  {welcome}

  <!-- Hero row -->
  <div class="dash-hero">
    <div>
      <h1 class="dash-greeting">Hi, {u["first_name"]} &#x1F44B;</h1>
      <div class="dash-subline">
        <span class="dash-dot"></span>
        {cv_status}
        <span style="color:var(--text-tertiary)">{last_run_badge}</span>
      </div>
    </div>
    <div class="run-form">
      {run_btn}
    </div>
  </div>

  <!-- Stats -->
  <div class="stat-grid">
    <div class="card stat-card">
      <span class="stat-icon">&#127919;</span>
      <div class="stat-num">{total_applied}</div>
      <div class="stat-label">Applications sent</div>
    </div>
    <div class="card stat-card">
      <span class="stat-icon">&#128269;</span>
      <div class="stat-num">{total_found}</div>
      <div class="stat-label">Jobs found</div>
    </div>
    <div class="card stat-card">
      <span class="stat-icon">&#9889;</span>
      <div class="stat-num">{run_count}</div>
      <div class="stat-label">Pipeline runs</div>
    </div>
  </div>

  <!-- Runs table -->
  <div class="section-row">
    <span class="section-title">Recent runs</span>
    <a class="section-link" href="/history">View all &rarr;</a>
  </div>
  <div class="card table-wrap">
    {runs_table}
  </div>

  <!-- Extension connect -->
  <div class="card ext-card">
    <div class="ext-icon">&#128279;</div>
    <div class="ext-body">
      <h2 class="ext-title">Connect the Chrome extension</h2>
      <p class="ext-desc">
        One click — opens a page that automatically configures the extension
        with your server and API key. No copying or pasting needed. Then press
        <strong style="color:var(--text-primary)">Start</strong> in the popup to begin applying.
      </p>
      <a href="/connect/{u["api_key"]}" target="_blank" class="btn btn-primary btn-lg">
        Connect extension in one click &rarr;
      </a>
    </div>
    <div class="ext-status">
      <div class="ext-status-lbl">Extension</div>
      <span class="badge badge-green">&#8226; Ready</span>
    </div>
  </div>
</div>"""

    is_running = runs and runs[0]["status"] == "running"
    active_run_id = runs[0]["id"] if is_running else None

    # When a run is active, inject a banner + silent JS poller (no meta-refresh)
    running_banner = ""
    poll_script    = ""
    if is_running:
        started_at = runs[0]["started_at"] or ""
        running_banner = f"""
        <div class="run-banner" id="run-banner">
          <div class="run-banner-left">
            <div class="run-banner-spinner"></div>
            <div>
              <div class="run-banner-title">Pipeline running</div>
              <div class="run-banner-sub">
                Finding jobs &rarr; scoring &rarr; tailoring CVs &rarr; generating cover letters
              </div>
            </div>
          </div>
          <div class="run-banner-right">
            <div class="run-banner-eta">Typically 5 &ndash; 15 min depending on jobs found</div>
            <div class="run-banner-elapsed" id="run-elapsed">Elapsed: calculating…</div>
          </div>
        </div>"""
        poll_script = f"""
        <script>
        (function(){{
          var startedAt = new Date("{started_at}Z");
          var runId     = {active_run_id};
          var reloaded  = false;

          function fmtElapsed(ms){{
            var s = Math.floor(ms/1000), m = Math.floor(s/60);
            s = s % 60;
            return m + 'm ' + (s<10?'0':'')+s+'s';
          }}

          // Tick elapsed time every second
          var ticker = setInterval(function(){{
            var el = document.getElementById('run-elapsed');
            if(el) el.textContent = 'Elapsed: ' + fmtElapsed(Date.now() - startedAt.getTime());
          }}, 1000);

          // Poll status every 10s — only reload once when run finishes
          var poller = setInterval(function(){{
            if(reloaded) return;
            fetch('/api/run-status', {{credentials:'same-origin'}})
              .then(function(r){{ return r.json(); }})
              .then(function(d){{
                if(d.status && d.status !== 'running'){{
                  reloaded = true;
                  clearInterval(ticker);
                  clearInterval(poller);
                  // Soft reload — no jarring flash, just update the page
                  location.reload();
                }}
              }})
              .catch(function(){{}});
          }}, 10000);
        }})();
        </script>"""

    # Inject running banner right after dash-wrap opens
    body = body.replace(
        '<div class="dash-wrap">',
        '<div class="dash-wrap">' + running_banner,
        1
    )
    body += poll_script

    # Add banner CSS to existing style block
    banner_css = """
    .run-banner{display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:16px;
      background:rgba(12,39,24,.14);border:1.5px solid rgba(26,92,52,.3);border-radius:16px;
      padding:20px 24px;margin-bottom:24px;
      box-shadow:0 4px 20px rgba(0,0,0,.06),inset 0 1px 0 rgba(255,255,255,.3)}
    .run-banner-left{display:flex;align-items:center;gap:16px}
    .run-banner-spinner{width:22px;height:22px;border-radius:50%;flex-shrink:0;
      border:2.5px solid rgba(26,92,52,.25);border-top-color:var(--green-mid);
      animation:spin .8s linear infinite}
    @keyframes spin{to{transform:rotate(360deg)}}
    .run-banner-title{font-family:'Clash Display',sans-serif;font-size:15px;font-weight:700;
      color:var(--text-primary);margin-bottom:3px}
    .run-banner-sub{font-size:12px;color:var(--text-secondary);font-style:italic}
    .run-banner-right{text-align:right}
    .run-banner-eta{font-size:12px;color:var(--text-secondary);margin-bottom:4px}
    .run-banner-elapsed{font-size:12px;color:var(--green-mid);font-family:'JetBrains Mono',monospace;font-weight:600}
    """
    body = body.replace("<style>", "<style>" + banner_css, 1)

    return _page("Dashboard", body, user=u)


# ── Run detail ────────────────────────────────────────────────────────────────

@app.route("/extension.zip")
def download_extension():
    """Serve the Chrome extension zip — no login required so new users can download it."""
    from flask import send_file
    p = ROOT / "chrome-extension.zip"
    if not p.exists():
        import zipfile
        with zipfile.ZipFile(str(p), "w") as z:
            for f in (ROOT / "chrome-extension").rglob("*"):
                if f.is_file():
                    z.write(f, f.relative_to(ROOT))
    return send_file(str(p), as_attachment=True, download_name="applyexpress-extension.zip")


@app.route("/download")
@login_required
def download_file():
    """Serve a file from the user's data directory as a download."""
    from flask import send_file, abort
    u    = current_user()
    path = request.args.get("path","")
    f    = Path(path)
    user_data = (DATA_DIR / "users" / str(u["id"])).resolve()
    if not str(f.resolve()).startswith(str(user_data)):
        abort(403)
    if not f.exists():
        abort(404)
    return send_file(str(f), as_attachment=True, download_name=f.name)


@app.route("/view")
@login_required
def view_file():
    """Serve a file inline in the browser (PDF viewer, plain text)."""
    from flask import send_file, abort
    u    = current_user()
    path = request.args.get("path","")
    f    = Path(path)
    user_data = (DATA_DIR / "users" / str(u["id"])).resolve()
    if not str(f.resolve()).startswith(str(user_data)):
        abort(403)
    if not f.exists():
        abort(404)
    mime = "application/pdf" if f.suffix == ".pdf" else "text/plain"
    return send_file(str(f), as_attachment=False, mimetype=mime)


@app.route("/run/<int:run_db_id>")
@login_required
def run_detail(run_db_id):
    u  = current_user()
    db = get_db()
    r  = db.execute("SELECT * FROM runs WHERE id=? AND user_id=?", (run_db_id, u["id"])).fetchone()
    if not r:
        return redirect(url_for("dashboard"))

    report = json.loads(r["report_json"]) if r["report_json"] else {"jobs": []}
    jobs   = report.get("jobs", [])

    status_map = {"running":"badge-blue","completed":"badge-green","failed":"badge-red","interrupted":"badge-yellow"}
    cls = status_map.get(r["status"], "badge-gray")

    # Pull per-application outcomes recorded by the Chrome extension
    app_rows = db.execute(
        "SELECT * FROM applications WHERE run_db_id=? ORDER BY id ASC", (run_db_id,)
    ).fetchall()
    app_outcomes = {(a["company"].lower(), a["title"].lower()): a for a in app_rows}

    # Augment cards with extension-reported status badge
    def _outcome_badge(j):
        key = (j.get("company","").lower(), j.get("title","").lower())
        a = app_outcomes.get(key)
        if not a:
            return '<span class="rd-pending">Pending extension</span>'
        if a["status"] == "applied":
            return '<span class="badge-green">&#10003; Applied</span>'
        return '<span class="badge-red">&#10007; Failed</span>'

    def _score_cls(score):
        if isinstance(score, int) and score >= 8: return "score-high"
        if isinstance(score, int) and score >= 6: return "score-mid"
        return "score-low"

    # Rebuild cards with outcome badge
    cards_with_status = ""
    for j in jobs:
        score   = j.get("score", "—")
        cv_name = Path(j["cv_docx"]).name if j.get("cv_docx") else "original CV"
        notes   = j.get("notes", "") or ""
        cl_text = ""
        if j.get("cover_letter_path") and Path(j["cover_letter_path"]).exists():
            cl_text = Path(j["cover_letter_path"]).read_text().strip()
        notes_html = f'<div class="rd-notes">&#9888; {notes}</div>' if notes else ""
        cv_viewer = (
            f'<details class="rd-details">'
            f'<summary class="rd-summary">&#128065; View CV (PDF)</summary>'
            f'<div class="rd-details-body">'
            f'<iframe src="/view?path={j["cv_docx"].replace(".docx",".pdf")}" class="rd-iframe"></iframe>'
            f'</div></details>'
            if j.get("cv_docx") and Path(j["cv_docx"].replace(".docx",".pdf")).exists() else ""
        )
        cv_dl = (
            f'<a href="/download?path={j["cv_docx"]}" class="rd-dl-btn">&#8595; Download CV (.docx)</a>'
            if j.get("cv_docx") and Path(j["cv_docx"]).exists() else ""
        )
        cl_viewer = (
            f'<details class="rd-details">'
            f'<summary class="rd-summary">&#128065; View Cover Letter</summary>'
            f'<div class="rd-details-body rd-cl-body">{cl_text}</div>'
            f'</details>'
            if cl_text else ""
        )
        cl_dl = (
            f'<a href="/download?path={j["cover_letter_path"]}" class="rd-dl-btn">&#8595; Download Cover Letter</a>'
            if j.get("cover_letter_path") and Path(j["cover_letter_path"]).exists() else ""
        )
        cl_empty = '<p class="rd-empty-cl">No cover letter generated.</p>' if not cl_text else ""
        cards_with_status += f"""
        <div class="card rd-card">
          <div class="rd-card-top">
            <div class="rd-card-info">
              <h3 class="rd-job-title">{j.get('title','—')}</h3>
              <p class="rd-job-meta">{j.get('company','—')} &middot; {j.get('source','')}
                &nbsp;&middot;&nbsp; <a href="{j.get('url','#')}" target="_blank" class="rd-job-link">View job &#8594;</a></p>
            </div>
            <div class="rd-score-wrap">
              <div class="rd-score {_score_cls(score)}">{score}<span class="rd-score-denom">/10</span></div>
              <div class="rd-score-label">match score</div>
              <div style="margin-top:8px">{_outcome_badge(j)}</div>
            </div>
          </div>
          {notes_html}
          <div class="rd-divider"></div>
          <div class="rd-section">
            <p class="rd-section-title">Tailored CV: <code class="rd-code">{cv_name}</code></p>
            {cv_viewer}
            {cv_dl}
          </div>
          <div class="rd-section">
            <p class="rd-section-title">Cover letter</p>
            {cl_viewer}
            {cl_dl}
            {cl_empty}
          </div>
        </div>"""

    # Also show a separate applications-only section if extension reported outcomes
    # for jobs not in pipeline report (e.g. manually fetched from Google Sheets)
    ext_only_rows = [a for a in app_rows
                     if not any(a["company"].lower() == j.get("company","").lower()
                                and a["title"].lower() == j.get("title","").lower()
                                for j in jobs)]
    ext_only_html = ""
    if ext_only_rows:
        ext_only_html = '<h2 class="rd-ext-heading">Additional extension-submitted jobs</h2>'
        for a in ext_only_rows:
            badge = ('<span class="badge-green">&#10003; Applied</span>'
                     if a["status"] == "applied" else
                     '<span class="badge-red">&#10007; Failed</span>')
            ext_only_html += f"""
            <div class="card rd-ext-card">
              <div>
                <strong class="rd-ext-title">{a["title"] or "—"}</strong>
                <span class="rd-ext-company"> @ {a["company"] or "—"}</span>
                {f'<br><a href="{a["url"]}" target="_blank" class="rd-job-link" style="font-size:12px">{a["url"][:60]}…</a>' if a["url"] else ''}
              </div>
              <div class="rd-ext-right">
                {badge}
                <span class="rd-ext-date">{a["applied_at"][:16]}</span>
              </div>
            </div>"""

    n_applied = sum(1 for a in app_rows if a["status"] == "applied")
    n_failed  = sum(1 for a in app_rows if a["status"] == "failed")

    spec_badge = ""
    if r["specialty_id"]:
        spec_row = db.execute("SELECT name FROM specialties WHERE id=?", (r["specialty_id"],)).fetchone()
        spec_badge = f'<span class="badge badge-blue">{spec_row["name"] if spec_row else "Specialty"}</span>'

    no_jobs_html = """<div class="rd-no-jobs">
      <svg width="36" height="36" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" style="opacity:.4;margin-bottom:10px">
        <circle cx="12" cy="12" r="10"/><path d="M12 8v4m0 4h.01"/>
      </svg>
      <p>No job details available for this run.</p>
    </div>"""

    body = f"""
    <style>
    .rd-wrap{{max-width:820px;margin:0 auto;padding:32px 20px 60px;animation:fadeUp .45s ease both}}
    .rd-back{{margin-bottom:20px}}
    .rd-back a{{color:var(--green-mid);font-size:14px;text-decoration:none;font-weight:500;opacity:.8}}
    .rd-back a:hover{{opacity:1}}
    .rd-header{{display:flex;align-items:center;gap:14px;flex-wrap:wrap;margin-bottom:24px}}
    .rd-header h1{{font-family:'Clash Display',sans-serif;font-size:26px;font-weight:700;color:var(--text-primary);margin:0}}
    .rd-stats{{display:flex;gap:12px;flex-wrap:wrap;margin-bottom:28px}}
    .rd-stat{{background:var(--glass);backdrop-filter:blur(16px) saturate(1.4);border:1.5px solid var(--glass-border);border-radius:14px;padding:14px 20px;min-width:88px;text-align:center;box-shadow:0 4px 16px rgba(0,0,0,.06),inset 0 1px 0 rgba(255,255,255,.5)}}
    .rd-stat-num{{font-family:'Clash Display',sans-serif;font-size:26px;font-weight:700;color:var(--text-primary);line-height:1}}
    .rd-stat-num.green{{color:var(--green-mid)}}
    .rd-stat-num.red{{color:#7a0e0e}}
    .rd-stat-label{{font-size:11px;font-weight:600;letter-spacing:.06em;text-transform:uppercase;color:var(--text-secondary);margin-top:4px}}
    .rd-card{{margin-bottom:20px}}
    .rd-card-top{{display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:14px}}
    .rd-card-info{{flex:1;min-width:0}}
    .rd-job-title{{font-family:'Clash Display',sans-serif;font-size:18px;font-weight:700;color:var(--text-primary);margin:0 0 5px}}
    .rd-job-meta{{font-size:13px;color:var(--text-secondary);margin:0}}
    .rd-job-link{{color:var(--green-mid);text-decoration:none;opacity:.85}}
    .rd-job-link:hover{{opacity:1;text-decoration:underline}}
    .rd-score-wrap{{text-align:right;flex-shrink:0}}
    .rd-score{{font-family:'Clash Display',sans-serif;font-size:32px;font-weight:800;line-height:1}}
    .rd-score.score-high{{color:var(--green-mid)}}
    .rd-score.score-mid{{color:#a05c00}}
    .rd-score.score-low{{color:#7a0e0e}}
    .rd-score-denom{{font-size:14px;color:var(--text-tertiary)}}
    .rd-score-label{{font-size:11px;color:var(--text-tertiary);letter-spacing:.04em}}
    .rd-notes{{font-size:12px;color:#a05c00;margin-top:10px;padding:8px 12px;background:rgba(245,166,35,.12);border:1px solid rgba(245,166,35,.25);border-radius:8px}}
    .rd-divider{{height:1px;background:var(--glass-border);margin:16px 0}}
    .rd-section{{margin-bottom:14px}}
    .rd-section-title{{font-size:13px;font-weight:600;color:var(--text-primary);margin:0 0 8px}}
    .rd-code{{background:rgba(255,255,255,.2);padding:2px 8px;border-radius:5px;font-family:'JetBrains Mono',monospace;font-size:12px;margin-left:4px;color:var(--text-primary)}}
    .rd-details{{border:1.5px solid var(--glass-border);border-radius:10px;overflow:hidden;margin-bottom:8px}}
    .rd-summary{{padding:9px 14px;cursor:pointer;font-size:13px;color:var(--green-mid);background:rgba(255,255,255,.14);user-select:none;list-style:none;display:flex;align-items:center;gap:6px;font-weight:500}}
    .rd-summary:hover{{background:rgba(255,255,255,.22)}}
    .rd-details-body{{padding:10px}}
    .rd-iframe{{width:100%;height:520px;border:none;border-radius:6px}}
    .rd-cl-body{{background:rgba(255,255,255,.12);padding:16px;font-size:13px;line-height:1.75;white-space:pre-wrap;color:var(--text-primary)}}
    .rd-dl-btn{{display:inline-flex;align-items:center;gap:6px;padding:7px 16px;border:1.5px solid rgba(26,92,52,.35);border-radius:8px;font-size:13px;color:var(--green-mid);background:rgba(26,92,52,.1);text-decoration:none;font-weight:500;margin-right:8px;transition:all .15s ease}}
    .rd-dl-btn:hover{{background:rgba(26,92,52,.2);border-color:rgba(26,92,52,.55)}}
    .rd-empty-cl{{font-size:13px;color:var(--text-tertiary);margin-top:6px}}
    .rd-pending{{font-size:12px;color:var(--text-tertiary);font-style:italic}}
    .rd-no-jobs{{padding:60px 20px;text-align:center;color:var(--text-secondary)}}
    .rd-no-jobs p{{margin:0;font-size:15px}}
    .rd-ext-heading{{font-family:'Clash Display',sans-serif;font-size:18px;font-weight:600;color:var(--text-primary);margin:28px 0 14px}}
    .rd-ext-card{{display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px;margin-bottom:12px;padding:14px 18px}}
    .rd-ext-title{{font-weight:600;font-size:14px;color:var(--text-primary)}}
    .rd-ext-company{{font-size:14px;color:var(--text-secondary)}}
    .rd-ext-right{{display:flex;align-items:center;gap:10px}}
    .rd-ext-date{{font-size:11px;color:var(--text-tertiary);font-family:'JetBrains Mono',monospace}}
    </style>
    <div class="rd-wrap">
      <div class="rd-back"><a href="/dashboard">&#8592; Dashboard</a></div>
      <div class="rd-header">
        <h1>Run {r["started_at"][:16]}</h1>
        <span class="badge {cls}">{r["status"]}</span>
        {spec_badge}
      </div>
      <div class="rd-stats">
        <div class="rd-stat">
          <div class="rd-stat-num">{r["jobs_found"]}</div>
          <div class="rd-stat-label">Jobs found</div>
        </div>
        <div class="rd-stat">
          <div class="rd-stat-num">{len(jobs)}</div>
          <div class="rd-stat-label">CVs tailored</div>
        </div>
        <div class="rd-stat">
          <div class="rd-stat-num green">{n_applied}</div>
          <div class="rd-stat-label">Applied &#10003;</div>
        </div>
        <div class="rd-stat">
          <div class="rd-stat-num red">{n_failed}</div>
          <div class="rd-stat-label">Failed &#10007;</div>
        </div>
      </div>
      {no_jobs_html if not jobs else cards_with_status}
      {ext_only_html}
    </div>"""
    return _page(f"Run {r['started_at'][:10]}", body, user=u)


# ── Onboarding wizard ─────────────────────────────────────────────────────────

@app.route("/setup", methods=["GET", "POST"])
@login_required
def setup():
    u    = current_user()
    step = int(request.args.get("step", 1))
    msg  = ""
    db   = get_db()

    if request.method == "POST":
        action = request.form.get("action", "")

        if action == "step1_cv":
            cv_file = request.files.get("cv")
            if cv_file and cv_file.filename.endswith(".docx"):
                cv_path = _user_dir(u["id"]) / "cv.docx"
                cv_file.save(str(cv_path))
                try:
                    _extract_cv_profile(cv_path, u, u["keywords"], u["search_location"])
                except Exception:
                    pass
            return redirect(url_for("setup") + "?step=2")

        elif action == "step2_prefs":
            keywords = request.form.get("keywords", "").strip()
            location = request.form.get("location", "London").strip()
            phone    = request.form.get("phone", "").strip()
            db.execute(
                "UPDATE users SET keywords=?, search_location=?, phone=?, email_subject=? WHERE id=?",
                (keywords, location, phone,
                 (keywords.split(",")[0].strip().title() if keywords else "Job Application"),
                 u["id"])
            )
            db.commit()
            return redirect(url_for("setup") + "?step=3")

        elif action == "step3_done":
            db.execute("UPDATE users SET onboarding_done=1 WHERE id=?", (u["id"],))
            db.commit()
            return redirect(url_for("dashboard") + "?welcome=1")

    u = current_user()  # refresh after possible commit
    cv_exists = (_user_dir(u["id"]) / "cv.docx").exists()

    def _step_node(i, label):
        if i < step:
            cls, num_cls = "sw-step sw-done", "sw-num sw-num-done"
            icon = "&#10003;"
        elif i == step:
            cls, num_cls = "sw-step sw-active", "sw-num sw-num-active"
            icon = str(i)
        else:
            cls, num_cls = "sw-step", "sw-num sw-num-future"
            icon = str(i)
        return (f'<div class="{cls}">'
                f'<div class="{num_cls}">{icon}</div>'
                f'<span class="sw-label">{label}</span>'
                f'</div>')

    steps_html = (
        _step_node(1, "Upload CV") +
        '<div class="sw-connector"></div>' +
        _step_node(2, "Preferences") +
        '<div class="sw-connector"></div>' +
        _step_node(3, "Extension")
    )

    if step == 1:
        cv_ready = (
            '<div class="sw-success-banner">&#10003; CV uploaded &mdash; you can replace it below</div>'
            if cv_exists else ""
        )
        keep_btn = (
            "<a href='/setup?step=2' class='btn sw-ghost-btn'>Keep current CV &#8594;</a>"
            if cv_exists else ""
        )
        content = f"""
        <h2 class="sw-step-title">Upload your CV</h2>
        <p class="sw-step-sub">We'll tailor a unique CV for every job you apply to.</p>
        {cv_ready}
        <form method="POST" enctype="multipart/form-data">
          <input type="hidden" name="action" value="step1_cv">
          <div class="field">
            <label class="label">CV file (.docx only)</label>
            <input type="file" name="cv" accept=".docx">
          </div>
          <div class="sw-btn-row">
            <button type="submit" class="btn btn-primary">{'Replace CV &amp; continue' if cv_exists else 'Upload &amp; continue'} &#8594;</button>
            {keep_btn}
          </div>
        </form>"""

    elif step == 2:
        content = f"""
        <h2 class="sw-step-title">Set your job preferences</h2>
        <p class="sw-step-sub">We search these terms and score jobs against your profile.</p>
        <form method="POST">
          <input type="hidden" name="action" value="step2_prefs">
          <div class="field">
            <label class="label">Job titles <span class="sw-hint">(comma-separated)</span></label>
            <input type="text" name="keywords" value="{u.get('keywords','')}" required
              placeholder="e.g. AML Compliance, MLRO, Risk Analyst">
          </div>
          <div class="field" style="margin-top:14px">
            <label class="label">Location</label>
            <input type="text" name="location" value="{u.get('search_location','London')}" required>
          </div>
          <div class="field" style="margin-top:14px">
            <label class="label">Phone number</label>
            <input type="tel" name="phone" value="{u.get('phone','')}" placeholder="+44 7700 900000">
          </div>
          <div class="sw-btn-row" style="margin-top:20px">
            <button type="submit" class="btn btn-primary">Save &amp; continue &#8594;</button>
          </div>
        </form>"""

    else:  # step 3
        api_key = u.get("api_key", "")
        server_url = request.host_url.rstrip("/")
        content = f"""
        <h2 class="sw-step-title">Install the Chrome extension</h2>
        <p class="sw-step-sub">The extension runs in your browser and submits applications on your behalf.</p>
        <ol class="sw-steps-list">
          <li>Download the extension zip from your dashboard</li>
          <li>Go to <code class="sw-inline-code">chrome://extensions</code> &#8594; enable Developer mode &#8594; Load unpacked</li>
          <li>Select the unzipped extension folder</li>
          <li>Click the ApplyExpress icon and paste these settings:</li>
        </ol>
        <div class="sw-creds-box">
          <div class="sw-cred-row">
            <span class="label">Server URL</span>
            <code class="sw-cred-code sw-cred-cyan">{server_url}</code>
          </div>
          <div class="sw-cred-row">
            <span class="label">Your API key</span>
            <code class="sw-cred-code sw-cred-green">{api_key}</code>
          </div>
        </div>
        <form method="POST">
          <input type="hidden" name="action" value="step3_done">
          <div class="sw-btn-row" style="margin-top:8px">
            <button type="submit" class="btn btn-primary">I'm done &mdash; go to dashboard &#8594;</button>
          </div>
        </form>"""

    body = f"""
    <style>
    .sw-wrap{{max-width:580px;margin:0 auto;padding:48px 20px 80px;animation:fadeUp .45s ease both}}
    .sw-skip{{margin-bottom:20px}}
    .sw-skip a{{font-size:13px;color:var(--text-tertiary);text-decoration:none;font-weight:500}}
    .sw-skip a:hover{{color:var(--text-secondary)}}
    .sw-rail{{display:flex;align-items:center;gap:0;margin-bottom:32px}}
    .sw-step{{display:flex;align-items:center;gap:9px}}
    .sw-connector{{flex:1;height:1.5px;background:var(--glass-border);margin:0 10px}}
    .sw-num{{width:30px;height:30px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:13px;font-weight:700;flex-shrink:0;transition:all .2s ease}}
    .sw-num-active{{background:var(--green-mid);color:#fdf8f0;box-shadow:0 0 0 3px rgba(26,92,52,.2)}}
    .sw-num-done{{background:rgba(12,39,24,.2);color:var(--green-mid);border:1.5px solid rgba(26,92,52,.4)}}
    .sw-num-future{{background:rgba(255,255,255,.15);color:var(--text-tertiary);border:1.5px solid var(--glass-border)}}
    .sw-label{{font-size:13px;font-weight:500;color:var(--text-secondary);white-space:nowrap}}
    .sw-active .sw-label{{color:var(--text-primary);font-weight:600}}
    .sw-done .sw-label{{color:var(--green-mid)}}
    .sw-card{{background:var(--glass);backdrop-filter:blur(20px) saturate(1.5);border:1.5px solid var(--glass-border);border-radius:20px;padding:32px;box-shadow:0 8px 32px rgba(0,0,0,.08),inset 0 1.5px 0 rgba(255,255,255,.55)}}
    .sw-step-title{{font-family:'Clash Display',sans-serif;font-size:22px;font-weight:700;color:var(--text-primary);margin:0 0 6px}}
    .sw-step-sub{{font-size:14px;color:var(--text-secondary);margin:0 0 24px}}
    .sw-hint{{font-weight:400;color:var(--text-tertiary);font-size:12px}}
    .sw-success-banner{{background:rgba(12,39,24,.15);border:1px solid rgba(26,92,52,.3);border-radius:8px;padding:10px 16px;font-size:13px;color:var(--green-mid);margin-bottom:18px;font-weight:500}}
    .sw-btn-row{{display:flex;gap:10px;align-items:center;flex-wrap:wrap;margin-top:20px}}
    .sw-ghost-btn{{background:rgba(255,255,255,.16);border:1.5px solid var(--glass-border);color:var(--text-primary)!important;padding:9px 18px;border-radius:10px;font-size:14px;font-weight:500;text-decoration:none;transition:all .15s ease}}
    .sw-ghost-btn:hover{{background:rgba(255,255,255,.26)}}
    .sw-steps-list{{padding-left:20px;line-height:2.2;color:var(--text-primary);font-size:14px;margin-bottom:18px}}
    .sw-inline-code{{background:rgba(255,255,255,.2);padding:2px 7px;border-radius:5px;font-family:'JetBrains Mono',monospace;font-size:12px}}
    .sw-creds-box{{background:rgba(12,39,24,.1);border:1.5px solid rgba(26,92,52,.25);border-radius:12px;padding:18px 20px;margin-bottom:20px}}
    .sw-cred-row{{margin-bottom:14px}}
    .sw-cred-row:last-child{{margin-bottom:0}}
    .sw-cred-code{{display:block;padding:10px 14px;border-radius:8px;font-family:'JetBrains Mono',monospace;font-size:13px;margin-top:5px;word-break:break-all}}
    .sw-cred-cyan{{background:rgba(12,39,24,.25);color:#5fceaa}}
    .sw-cred-green{{background:rgba(12,39,24,.25);color:#a0e0c0}}
    </style>
    <div class="sw-wrap">
      <div class="sw-skip"><a href="/dashboard">&#8592; Skip for now</a></div>
      <div class="sw-rail">{steps_html}</div>
      <div class="sw-card">
        {content}
      </div>
    </div>"""
    return _page("Setup", body, user=u)


# ── Profile / settings ────────────────────────────────────────────────────────

@app.route("/profile/extract-cv-keywords", methods=["POST"])
@login_required
def extract_cv_keywords():
    """Read a .docx file and return the Skills section text as keywords."""
    from docx import Document as DocxDocument
    import tempfile, os as _os

    cv_file = request.files.get("specialty_cv")
    if not cv_file or not cv_file.filename.endswith(".docx"):
        return jsonify({"keywords": ""})

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            cv_file.save(tmp.name)
            tmp_path = tmp.name

        doc = DocxDocument(tmp_path)
        _os.unlink(tmp_path)

        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Find Skills section
        skill_keywords = {"SKILLS", "KEY SKILLS", "CORE SKILLS", "COMPETENCIES", "EXPERTISE"}
        keywords_text = ""
        in_section = False
        for para in paras:
            upper = para.upper()
            if any(upper == kw or upper.startswith(kw + " ") for kw in skill_keywords):
                in_section = True
                continue
            if in_section:
                # Stop at next all-caps heading
                if para == para.upper() and any(c.isalpha() for c in para) and len(para) < 60:
                    break
                keywords_text += para + " "

        # Fallback: use first paragraph (usually job title/summary) if no skills section
        if not keywords_text.strip() and paras:
            keywords_text = paras[0]

        # Normalise: strip bullets, extra spaces, truncate
        import re as _re
        keywords_text = _re.sub(r"[•\-–*▪◦○·]", ",", keywords_text)
        keywords_text = _re.sub(r"\s+", " ", keywords_text).strip().strip(",").strip()
        if len(keywords_text) > 300:
            keywords_text = keywords_text[:300]

        return jsonify({"keywords": keywords_text})
    except Exception as e:
        return jsonify({"keywords": "", "error": str(e)})


@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile():
    u   = current_user()
    db  = get_db()
    msg = ""

    if request.method == "POST":
        action = request.form.get("action", "profile")

        if action == "blacklist_add":
            company = request.form.get("company", "").strip()
            if company:
                try:
                    db.execute(
                        "INSERT OR IGNORE INTO blacklist (user_id, company, added_at) VALUES (?,?,?)",
                        (u["id"], company, datetime.utcnow().isoformat())
                    )
                    db.commit()
                    msg = f"'{company}' added to blacklist."
                except Exception:
                    pass
            return redirect(url_for("profile") + "#blacklist")

        elif action == "blacklist_remove":
            company = request.form.get("company", "").strip()
            if company:
                db.execute("DELETE FROM blacklist WHERE user_id=? AND company=?", (u["id"], company))
                db.commit()
            return redirect(url_for("profile") + "#blacklist")

        elif action == "profile":
            db.execute("""UPDATE users SET first_name=?,last_name=?,phone=?,location=?,
                          keywords=?,search_location=?,threshold=?,email_subject=?
                          WHERE id=?""",
                (request.form.get("first_name","").strip(),
                 request.form.get("last_name","").strip(),
                 request.form.get("phone","").strip(),
                 request.form.get("location","").strip(),
                 request.form.get("keywords","").strip(),
                 request.form.get("search_location","").strip(),
                 int(request.form.get("threshold", 7)),
                 request.form.get("keywords","Job Application").split(",")[0].strip().title(),
                 u["id"]))
            db.commit()
            msg = "Profile saved."
            u = current_user()

        elif action == "cv":
            cv_file = request.files.get("cv")
            if cv_file and cv_file.filename.endswith(".docx"):
                cv_path = _user_dir(u["id"]) / "cv.docx"
                cv_file.save(str(cv_path))
                try:
                    profile_text_content = _extract_cv_profile(cv_path, u, u["keywords"], u["search_location"])
                    (_user_dir(u["id"]) / "candidate_profile.md").write_text(profile_text_content)
                    msg = "CV uploaded and profile updated from CV content." + _cv_parse_warning(cv_path)
                except Exception as e:
                    msg = f"CV uploaded. (Profile extraction failed: {e})"
            else:
                msg = "Please upload a .docx file."

        elif action == "add_specialty":
            spec_name  = request.form.get("specialty_name","").strip()
            spec_slug  = re.sub(r"[^\w]", "_", spec_name.lower()).strip("_")[:30]
            spec_kw    = request.form.get("spec_keywords", u["keywords"]).strip()
            # Location and threshold inherited from global settings
            spec_loc   = u["search_location"]
            spec_thr   = u["threshold"]
            cv_file    = request.files.get("specialty_cv")
            if not spec_name or not spec_slug:
                msg = "Please enter a specialty name."
            elif not cv_file or not cv_file.filename.endswith(".docx"):
                msg = "Please upload a .docx CV file for this specialty."
            else:
                db.execute("""INSERT OR IGNORE INTO specialties
                    (user_id,name,slug,keywords,search_location,threshold,created_at)
                    VALUES (?,?,?,?,?,?,?)""",
                    (u["id"], spec_name, spec_slug, spec_kw, spec_loc, spec_thr,
                     datetime.utcnow().isoformat()))
                db.commit()
                sdir    = _specialty_dir(u["id"], spec_slug)
                cv_path = sdir / "cv.docx"
                cv_file.save(str(cv_path))
                try:
                    profile_text_content = _extract_cv_profile(cv_path, u, spec_kw, spec_loc)
                    (sdir / "profile.md").write_text(profile_text_content)
                    (_user_dir(u["id"]) / "candidate_profile.md").write_text(profile_text_content)
                except Exception:
                    pass
                msg = f'Specialty "{spec_name}" added.'

        elif action == "upload_specialty_cv":
            spec_id = request.form.get("spec_id","")
            spec    = db.execute("SELECT * FROM specialties WHERE id=? AND user_id=?",
                                 (spec_id, u["id"])).fetchone() if spec_id else None
            if spec:
                sdir    = _specialty_dir(u["id"], spec["slug"])
                cv_file = request.files.get("specialty_cv")
                if cv_file and cv_file.filename.endswith(".docx"):
                    cv_path = sdir / "cv.docx"
                    cv_file.save(str(cv_path))
                    try:
                        profile_text_content = _extract_cv_profile(cv_path, u, spec["keywords"], spec["search_location"])
                        (sdir / "profile.md").write_text(profile_text_content)
                        # Also update global candidate profile so Profile Content section refreshes
                        (_user_dir(u["id"]) / "candidate_profile.md").write_text(profile_text_content)
                    except Exception:
                        pass
                    msg = f'CV uploaded for "{spec["name"]}". Profile content updated.' + _cv_parse_warning(cv_path)
                else:
                    msg = "Please upload a .docx file."
            else:
                msg = "Specialty not found."

        elif action == "save_specialty_profile":
            spec_id = request.form.get("spec_id","")
            spec    = db.execute("SELECT * FROM specialties WHERE id=? AND user_id=?",
                                 (spec_id, u["id"])).fetchone() if spec_id else None
            if spec:
                (_specialty_dir(u["id"], spec["slug"]) / "profile.md").write_text(
                    request.form.get("profile_text","").replace("\r\n","\n"))
                msg = f'Profile saved for "{spec["name"]}".'
            else:
                msg = "Specialty not found."

        elif action == "delete_specialty":
            spec_id = request.form.get("spec_id","")
            spec    = db.execute("SELECT * FROM specialties WHERE id=? AND user_id=?",
                                 (spec_id, u["id"])).fetchone() if spec_id else None
            if spec:
                import shutil
                shutil.rmtree(str(_specialty_dir(u["id"], spec["slug"])), ignore_errors=True)
                db.execute("DELETE FROM specialties WHERE id=?", (spec_id,))
                db.commit()
                msg = f'Specialty "{spec["name"]}" deleted.'

        elif action == "profile_text":
            profile_path = _user_dir(u["id"]) / "candidate_profile.md"
            profile_path.write_text(request.form.get("profile_text", "").replace("\r\n", "\n"))
            msg = "Profile content saved."

        elif action == "credentials":
            db.execute("""UPDATE users SET smtp_password=?,
                          linkedin_email=?,linkedin_pass=?,
                          reed_email=?,reed_pass=?,
                          indeed_email=?,indeed_pass=?
                          WHERE id=?""",
                (_enc(request.form.get("smtp_password","").strip()),
                 request.form.get("linkedin_email","").strip(),
                 _enc(request.form.get("linkedin_pass","").strip()),
                 request.form.get("reed_email","").strip(),
                 _enc(request.form.get("reed_pass","").strip()),
                 request.form.get("indeed_email","").strip(),
                 _enc(request.form.get("indeed_pass","").strip()),
                 u["id"]))
            db.commit()
            msg = "Credentials saved."
            u = current_user()

    profile_md_path = _user_dir(u["id"]) / "candidate_profile.md"
    profile_md_text = profile_md_path.read_text() if profile_md_path.exists() else ""
    specialties_list = db.execute(
        "SELECT * FROM specialties WHERE user_id=? ORDER BY name", (u["id"],)).fetchall()

    # Fallback: if no global profile yet, use the most recently updated specialty profile
    if not profile_md_text:
        spec_profiles = sorted(
            [_specialty_dir(u["id"], s["slug"]) / "profile.md" for s in specialties_list],
            key=lambda p: p.stat().st_mtime if p.exists() else 0, reverse=True
        )
        for sp in spec_profiles:
            if sp.exists():
                profile_md_text = sp.read_text()
                break

    alert = f'<div class="alert alert-success">{msg}</div>' if msg else ""

    blacklist_rows = db.execute(
        "SELECT company FROM blacklist WHERE user_id=? ORDER BY company", (u["id"],)
    ).fetchall()

    blacklist_html = "".join(
        f'<div class="pf-bl-row">'
        f'<span class="pf-bl-name">{b["company"]}</span>'
        f'<form method="POST" style="margin:0">'
        f'<input type="hidden" name="action" value="blacklist_remove">'
        f'<input type="hidden" name="company" value="{b["company"]}">'
        f'<button type="submit" class="pf-bl-remove">Remove &#10005;</button>'
        f'</form></div>'
        for b in blacklist_rows
    ) or '<p class="pf-muted">No companies blacklisted yet.</p>'

    def _spec_profile_text(uid, slug):
        p = _specialty_dir(uid, slug) / "profile.md"
        return p.read_text() if p.exists() else ""

    spec_profiles_html = "".join(f"""
    <div class="pf-spec-profile">
      <div class="pf-spec-profile-label">{s['name']}</div>
      <form method="POST">
        <input type="hidden" name="action" value="save_specialty_profile">
        <input type="hidden" name="spec_id" value="{s['id']}">
        <textarea name="profile_text" rows="10" class="pf-textarea"
          placeholder="Upload a CV for this specialty to auto-populate the profile.">{_spec_profile_text(u['id'], s['slug'])}</textarea>
        <button type="submit" class="btn btn-primary" style="margin-top:8px">Save</button>
      </form>
    </div>""" for s in specialties_list
    ) or '<p class="pf-muted">No specialties yet — add one above to see profile content here.</p>'

    def _spec_cv_label(uid, slug):
        p = _specialty_dir(uid, slug) / "cv.docx"
        if p.exists():
            kb = p.stat().st_size // 1024
            return f'<span class="spec-cv-ok">&#10003; cv.docx ({kb}KB)</span>'
        return '<span class="spec-cv-missing">No CV yet</span>'

    spec_cards_html = ("".join(f"""
    <div class="pf-spec-row">
      <div class="pf-spec-top">
        <span class="pf-spec-name">{s['name']}</span>
        <div class="pf-spec-actions">
          {_spec_cv_label(u['id'], s['slug'])}
          <form method="POST" enctype="multipart/form-data" id="ucv-{s['id']}" style="margin:0;display:inline">
            <input type="hidden" name="action" value="upload_specialty_cv">
            <input type="hidden" name="spec_id" value="{s['id']}">
            <label class="btn-link-sm">Replace CV
              <input type="file" name="specialty_cv" accept=".docx" required
                style="display:none" onchange="document.getElementById('ucv-{s['id']}').submit()">
            </label>
          </form>
        </div>
      </div>
      <div class="pf-spec-keywords">{s['keywords'] or '—'}</div>
      <form method="POST" style="margin:4px 0 0">
        <input type="hidden" name="action" value="delete_specialty">
        <input type="hidden" name="spec_id" value="{s['id']}">
        <button type="submit" class="pf-del-btn" onclick="return confirm('Delete specialty &quot;{s['name']}&quot;?')">Delete</button>
      </form>
    </div>""" for s in specialties_list)
    or '<p class="pf-muted" style="margin-bottom:4px">No specialties yet — add one below.</p>')

    body = f"""
    <style>
    .pf-wrap{{max-width:660px;margin:0 auto;padding:32px 20px 80px;animation:fadeUp .45s ease both}}
    .pf-wrap h1{{font-family:'Clash Display',sans-serif;font-size:28px;font-weight:700;color:var(--text-primary);margin:0 0 28px}}
    .pf-wrap .card{{margin-bottom:20px}}
    .pf-wrap h2{{font-family:'Clash Display',sans-serif;font-size:18px;font-weight:700;color:var(--text-primary);margin:0 0 16px}}
    .pf-wrap h3{{font-size:15px;font-weight:600;color:var(--text-primary);margin:0 0 14px}}
    .pf-divider{{height:1px;background:var(--glass-border);margin:20px 0}}
    .pf-hint{{font-weight:400;color:var(--text-tertiary);font-size:12px}}
    .pf-muted{{font-size:13px;color:var(--text-tertiary);margin:0}}
    .pf-sub{{font-size:13px;color:var(--text-secondary);margin:0 0 18px;line-height:1.6}}
    .pf-textarea{{width:100%;font-family:'JetBrains Mono',monospace;font-size:12px;line-height:1.6;border:1.5px solid var(--glass-border);border-radius:8px;padding:12px;resize:vertical;background:rgba(255,255,255,.12);color:var(--text-primary);box-sizing:border-box}}
    .pf-textarea:focus{{outline:none;border-color:rgba(26,92,52,.5);box-shadow:0 0 0 3px rgba(26,92,52,.1)}}
    .pf-board-label{{display:block;font-size:11px;font-weight:700;letter-spacing:.04em;color:var(--text-primary);margin:18px 0 8px;text-transform:uppercase}}
    .pf-spec-row{{padding:12px 0;border-bottom:1px solid var(--glass-border)}}
    .pf-spec-row:last-of-type{{border-bottom:none;padding-bottom:0}}
    .pf-spec-top{{display:flex;justify-content:space-between;align-items:center;gap:10px;flex-wrap:wrap}}
    .pf-spec-name{{font-size:15px;font-weight:700;color:var(--text-primary)}}
    .pf-spec-actions{{display:flex;align-items:center;gap:10px;flex-shrink:0}}
    .pf-spec-keywords{{font-size:12px;color:var(--text-tertiary);margin-top:4px}}
    .spec-cv-ok{{color:#4ade80;font-weight:500;font-size:13px}}
    .spec-cv-missing{{color:var(--text-tertiary);font-style:italic;font-size:13px}}
    .btn-upload-cv{{display:inline-flex;align-items:center;gap:6px;background:rgba(26,92,52,.18);color:var(--green-mid);border:1.5px solid rgba(26,92,52,.3);padding:6px 14px;border-radius:8px;cursor:pointer;font-size:13px;font-weight:600;transition:all .15s ease;white-space:nowrap}}
    .btn-upload-cv:hover{{background:rgba(26,92,52,.3)}}
    .btn-link-sm{{font-size:12px;color:var(--text-tertiary);cursor:pointer;text-decoration:underline;text-underline-offset:2px;white-space:nowrap;background:none;border:none;padding:0}}
    .btn-link-sm:hover{{color:var(--text-primary)}}
    .pf-spec-profile{{padding:16px 0}}
    .pf-spec-profile+.pf-spec-profile{{border-top:1px solid var(--glass-border)}}
    .pf-spec-profile-label{{font-size:13px;font-weight:700;color:var(--text-primary);margin-bottom:8px;text-transform:uppercase;letter-spacing:.04em}}
    .pf-add-spec-row{{display:flex;gap:10px;align-items:center}}
    .pf-add-spec-row input[type=text]{{flex:1;min-width:0}}
    .pf-del-btn{{background:none;color:var(--text-tertiary);border:none;padding:0;cursor:pointer;font-size:12px;text-decoration:underline;text-underline-offset:2px}}
    .pf-del-btn:hover{{color:#f87171}}
    .pf-bl-row{{display:flex;justify-content:space-between;align-items:center;padding:10px 14px;border-radius:9px;background:rgba(255,255,255,.1);border:1.5px solid rgba(255,255,255,.18);margin-bottom:8px}}
    .pf-bl-name{{font-size:14px;color:var(--text-primary);font-weight:500}}
    .pf-bl-remove{{background:none;border:none;cursor:pointer;color:#7a0e0e;font-size:13px;padding:4px 8px;font-weight:500;opacity:.75;transition:opacity .12s ease}}
    .pf-bl-remove:hover{{opacity:1}}
    .pf-add-bl{{display:flex;gap:10px;margin-bottom:18px}}
    .pf-add-bl input{{flex:1}}
    </style>
    <div class="pf-wrap">
      {alert}
      <h1>Settings</h1>

      <div class="card">
        <h2>Your details</h2>
        <form method="POST">
          <input type="hidden" name="action" value="profile">
          <div class="row">
            <div class="field"><label class="label">First name</label>
              <input type="text" name="first_name" value="{u['first_name']}" required></div>
            <div class="field"><label class="label">Last name</label>
              <input type="text" name="last_name" value="{u['last_name']}" required></div>
          </div>
          <div class="field"><label class="label">Phone</label>
            <input type="tel" name="phone" value="{u['phone']}" placeholder="+447700900000"></div>
          <div class="field"><label class="label">Location (for profile)</label>
            <input type="text" name="location" value="{u['location']}"></div>
          <div class="pf-divider"></div>
          <h3>Job search preferences</h3>
          <div class="field"><label class="label">Job titles / keywords <span class="pf-hint">(comma-separated)</span></label>
            <input type="text" name="keywords" value="{u['keywords']}"></div>
          <div class="row">
            <div class="field"><label class="label">Search location</label>
              <input type="text" name="search_location" value="{u['search_location']}"></div>
            <div class="field"><label class="label">Minimum match score (1–10)</label>
              <input type="number" name="threshold" value="{u['threshold']}" min="1" max="10"></div>
          </div>
          <button type="submit" class="btn btn-primary">Save profile</button>
        </form>
      </div>

      <div class="card" id="specialties">
        <h2>Specialties</h2>
        {spec_cards_html}
        <div class="pf-divider"></div>
        <h3 style="margin-bottom:12px">Add a specialty</h3>
        <form method="POST" enctype="multipart/form-data">
          <input type="hidden" name="action" value="add_specialty">
          <div class="pf-add-spec-row">
            <input type="text" name="specialty_name" placeholder="e.g. AML Compliance Officer" required>
            <label class="btn-upload-cv" id="cv-upload-label">Upload CV
              <input type="file" name="specialty_cv" accept=".docx" required style="display:none" id="spec-cv-input"
                onchange="handleSpecCvChange(this)">
            </label>
          </div>
          <span id="cv-chosen" style="font-size:12px;color:var(--text-tertiary);margin-top:4px;display:block">No file chosen</span>
          <script>
          function handleSpecCvChange(input) {{
            var chosen = document.getElementById('cv-chosen');
            var kwField = document.querySelector('input[name=spec_keywords]');
            if (!input.files[0]) return;
            chosen.textContent = input.files[0].name;
            chosen.style.color = 'var(--green-mid)';
            kwField.placeholder = 'Extracting from CV…';
            kwField.value = '';
            var fd = new FormData();
            fd.append('specialty_cv', input.files[0]);
            fetch('/profile/extract-cv-keywords', {{method:'POST', body:fd}})
              .then(function(r){{return r.json();}})
              .then(function(d){{
                kwField.placeholder = '';
                if (d.keywords) kwField.value = d.keywords;
              }})
              .catch(function(){{ kwField.placeholder = 'Could not extract — enter manually'; }});
          }}
          </script>
          <div class="field" style="margin-top:10px;margin-bottom:0">
            <label class="label">Keywords <span class="pf-hint">(auto-populated from CV — edit if needed)</span></label>
            <input type="text" name="spec_keywords" placeholder="Upload a CV above to auto-populate">
          </div>
          <button type="submit" class="btn btn-primary" style="margin-top:12px">Save</button>
        </form>
      </div>

      <div class="card" id="profile-content">
        <h2>Profile content</h2>
        <p class="pf-sub">Extracted from each specialty CV automatically. Used for scoring, cover letters, and ATS tailoring. Edit if anything needs correcting.</p>
        {spec_profiles_html}
      </div>

      <div class="card">
        <h2>Email &amp; job board credentials</h2>
        <div class="info-tip">
          <b>Gmail App Password <span style="font-weight:400">(optional)</span>:</b>
          If you want reports emailed to <strong>{u['email']}</strong> after each run,
          go to <a href="https://myaccount.google.com/apppasswords" target="_blank" style="color:var(--green-mid)">myaccount.google.com/apppasswords</a>
          &#8594; create one called "ApplyExpress" &#8594; paste the 16-char code below.
        </div>
        <form method="POST">
          <input type="hidden" name="action" value="credentials">
          <div class="field"><label class="label">Gmail App Password <span class="pf-hint">(optional — leave blank to skip email reports)</span></label>
            <input type="password" name="smtp_password" value="{u['smtp_password']}" placeholder="xxxx xxxx xxxx xxxx"></div>
          <div class="pf-divider"></div>
          <p class="pf-sub">Enter your login for each job board you want to auto-apply on. Leave blank to skip that board.</p>
          <span class="pf-board-label">LinkedIn</span>
          <div class="row">
            <div class="field"><label class="label">Email</label>
              <input type="email" name="linkedin_email" value="{u['linkedin_email']}" placeholder="your@email.com"></div>
            <div class="field"><label class="label">Password</label>
              <input type="password" name="linkedin_pass" value="{u['linkedin_pass']}" placeholder="Leave blank to skip"></div>
          </div>
          <span class="pf-board-label">Reed.co.uk</span>
          <div class="row">
            <div class="field"><label class="label">Email</label>
              <input type="email" name="reed_email" value="{u['reed_email']}" placeholder="your@email.com"></div>
            <div class="field"><label class="label">Password</label>
              <input type="password" name="reed_pass" value="{u['reed_pass']}" placeholder="Leave blank to skip"></div>
          </div>
          <span class="pf-board-label">Indeed</span>
          <div class="row">
            <div class="field"><label class="label">Email</label>
              <input type="email" name="indeed_email" value="{u['indeed_email']}" placeholder="your@email.com"></div>
            <div class="field"><label class="label">Password</label>
              <input type="password" name="indeed_pass" value="{u['indeed_pass']}" placeholder="Leave blank to skip"></div>
          </div>
          <button type="submit" class="btn btn-primary">Save credentials</button>
        </form>
      </div>

      <div class="card" id="sessions">
        <h2>Job board sessions</h2>
        <p class="pf-sub">Applying needs a logged-in browser session per job board. Sessions expire every
          few weeks — when one does, applications on that board fail until it's refreshed.</p>
        <div class="pf-bl-row">
          <span class="pf-bl-name">Reed</span>
          <span id="sess-reed-age" class="pf-muted">checking&#8230;</span>
          <button type="button" class="btn btn-primary" id="sess-reed-btn"
                  onclick="refreshReed()" style="white-space:nowrap">Refresh now</button>
        </div>
        <div class="pf-bl-row">
          <span class="pf-bl-name">LinkedIn</span>
          <span id="sess-linkedin-age" class="pf-muted">checking&#8230;</span>
          <span class="pf-hint">refresh via desktop Chrome extension</span>
        </div>
        <div class="pf-bl-row">
          <span class="pf-bl-name">Indeed</span>
          <span id="sess-indeed-age" class="pf-muted">checking&#8230;</span>
          <span class="pf-hint">refresh via desktop Chrome extension</span>
        </div>
        <p id="sess-msg" class="pf-muted" style="margin-top:10px"></p>
        <script>
        async function sessAge(p){{
          try{{
            const j = await (await fetch('/profile/refresh-session-status?platform='+p)).json();
            const el = document.getElementById('sess-'+p+'-age');
            if(j.session_age_days === undefined){{ el.textContent = 'no session uploaded'; el.style.color='#dc2626'; }}
            else{{
              const d = j.session_age_days;
              el.textContent = d < 1 ? 'updated today' : 'updated ' + Math.round(d) + ' day' + (Math.round(d)===1?'':'s') + ' ago';
              el.style.color = d > 21 ? '#dc2626' : (d > 10 ? '#d97706' : '#16a34a');
            }}
          }}catch(e){{}}
        }}
        ['reed','linkedin','indeed'].forEach(sessAge);
        async function refreshReed(){{
          const btn = document.getElementById('sess-reed-btn'), msg = document.getElementById('sess-msg');
          btn.disabled = true; msg.textContent = 'Logging in to Reed from the server (up to 2 minutes)…';
          try{{
            const r = await fetch('/profile/refresh-session', {{method:'POST',
              headers:{{'Content-Type':'application/x-www-form-urlencoded'}}, body:'platform=reed'}});
            const j = await r.json();
            if(!j.ok){{ msg.textContent = j.reason || 'Could not start refresh.'; btn.disabled = false; return; }}
            let tries = 0;
            const t = setInterval(async () => {{
              tries++;
              const s = await (await fetch('/profile/refresh-session-status?platform=reed')).json();
              if(s.state === 'ok'){{ clearInterval(t); msg.textContent = 'Reed session refreshed ✓'; btn.disabled = false; sessAge('reed'); }}
              else if(s.state === 'captcha' || s.state === 'failed'){{
                clearInterval(t); msg.textContent = (s.state === 'captcha' ? '' : 'Failed: ') + (s.detail || 'Login failed.'); btn.disabled = false;
              }}
              else if(tries > 50){{ clearInterval(t); msg.textContent = 'Still running — check back in a minute.'; btn.disabled = false; }}
            }}, 3000);
          }}catch(e){{ msg.textContent = 'Error: ' + e; btn.disabled = false; }}
        }}
        </script>
      </div>

      <div class="card" id="blacklist">
        <h2>Company blacklist</h2>
        <p class="pf-sub">Jobs from these companies will be skipped automatically.</p>
        <form method="POST" class="pf-add-bl">
          <input type="hidden" name="action" value="blacklist_add">
          <input type="text" name="company" placeholder="Company name e.g. Acme Corp" required>
          <button type="submit" class="btn btn-primary" style="white-space:nowrap">Add &#8594;</button>
        </form>
        {blacklist_html}
      </div>
    </div>"""
    return _page("Settings", body, user=u)


# ── Application history ───────────────────────────────────────────────────────

@app.route("/history")
@login_required
def history():
    u      = current_user()
    db     = get_db()
    q      = request.args.get("q", "").strip()
    status = request.args.get("status", "")   # 'applied' | 'failed' | ''
    page   = max(1, int(request.args.get("page", 1)))
    per    = 25

    conditions = ["user_id=?"]
    params     = [u["id"]]
    if q:
        conditions.append("(lower(title) LIKE ? OR lower(company) LIKE ?)")
        params += [f"%{q.lower()}%", f"%{q.lower()}%"]
    if status in ("applied", "failed"):
        conditions.append("status=?")
        params.append(status)

    where = " AND ".join(conditions)
    total = db.execute(f"SELECT COUNT(*) FROM applications WHERE {where}", params).fetchone()[0]
    rows  = db.execute(
        f"SELECT * FROM applications WHERE {where} ORDER BY applied_at DESC LIMIT ? OFFSET ?",
        params + [per, (page - 1) * per]
    ).fetchall()

    n_applied  = db.execute("SELECT COUNT(*) FROM applications WHERE user_id=? AND status='applied'",        (u["id"],)).fetchone()[0]
    n_failed   = db.execute("SELECT COUNT(*) FROM applications WHERE user_id=? AND status='failed'",         (u["id"],)).fetchone()[0]
    n_sponsor  = db.execute("SELECT COUNT(*) FROM applications WHERE user_id=? AND status='sponsor_review'", (u["id"],)).fetchone()[0]

    # Pagination
    total_pages = max(1, (total + per - 1) // per)
    def page_link(p, label):
        qs = f"?q={q}&status={status}&page={p}"
        cls = "pg-link pg-active" if p == page else "pg-link"
        return f'<a href="/history{qs}" class="{cls}">{label}</a>'

    pager = ""
    if total_pages > 1:
        parts = [page_link(1, "«")]
        for p in range(max(1, page - 2), min(total_pages + 1, page + 3)):
            parts.append(page_link(p, str(p)))
        parts.append(page_link(total_pages, "»"))
        pager = f'<div class="hist-pager">{"".join(parts)}</div>'

    table_rows = ""
    for r in rows:
        badge = ('<span class="badge-green">Applied</span>'
                 if r["status"] == "applied" else
                 '<span style="background:#7c3aed;color:#fff;padding:2px 8px;border-radius:4px;font-size:12px;font-weight:600">&#127919; Sponsor — Apply Manually</span>'
                 if r["status"] == "sponsor_review" else
                 '<span class="badge-red">Failed</span>')
        url_html = (f'<a href="{r["url"]}" target="_blank" class="hist-link" title="{r["url"]}">'
                    f'{r["url"][:45]}{"…" if len(r["url"]) > 45 else ""}</a>' if r["url"] else '<span class="hist-dash">—</span>')
        notes_html = f'<div class="hist-notes">{r["notes"][:60]}</div>' if r["notes"] else ""
        table_rows += f"""<tr class="hist-row">
          <td class="hist-td">
            <div class="hist-title">{r["title"] or "—"}</div>
            <div class="hist-company">{r["company"] or "—"}</div>
            {notes_html}
          </td>
          <td class="hist-td hist-url-td">{url_html}</td>
          <td class="hist-td">{badge}</td>
          <td class="hist-td hist-date">{(r["applied_at"] or "")[:16]}</td>
        </tr>"""

    filter_qs = lambda s: f'/history?q={q}&status={s}'
    active_tab = lambda s: "hist-tab hist-tab-active" if status == s else "hist-tab"

    empty_state = """<div class="hist-empty">
      <svg width="40" height="40" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" style="opacity:.4;margin-bottom:10px">
        <path d="M9 5H7a2 2 0 00-2 2v12a2 2 0 002 2h10a2 2 0 002-2V7a2 2 0 00-2-2h-2M9 5a2 2 0 002 2h2a2 2 0 002-2M9 5a2 2 0 012-2h2a2 2 0 012 2"/>
      </svg>
      <p>No applications found.</p>
    </div>"""

    table_html = f"""<div class="hist-table-wrap">
      <table class="hist-table">
        <thead><tr class="hist-thead-row">
          <th class="hist-th">Role / Company</th>
          <th class="hist-th">Link</th>
          <th class="hist-th">Status</th>
          <th class="hist-th">Date</th>
        </tr></thead>
        <tbody>{table_rows}</tbody>
      </table>
    </div>""" if rows else empty_state

    body = f"""
    <style>
    .hist-wrap{{max-width:960px;margin:0 auto;padding:32px 20px 60px;animation:fadeUp .45s ease both}}
    .hist-header{{display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:16px;margin-bottom:28px}}
    .hist-title-group h1{{font-family:'Clash Display',sans-serif;font-size:28px;font-weight:700;color:var(--text-primary);margin:0 0 4px}}
    .hist-title-group p{{margin:0;color:var(--text-secondary);font-size:14px}}
    .hist-stats{{display:flex;gap:12px;flex-wrap:wrap}}
    .hist-stat{{background:var(--glass);backdrop-filter:blur(16px) saturate(1.4);border:1.5px solid var(--glass-border);border-radius:14px;padding:14px 20px;min-width:90px;text-align:center;box-shadow:0 4px 16px rgba(0,0,0,.06),inset 0 1px 0 rgba(255,255,255,.5)}}
    .hist-stat-num{{font-family:'Clash Display',sans-serif;font-size:26px;font-weight:700;color:var(--text-primary);line-height:1}}
    .hist-stat-num.green{{color:var(--green-mid)}}
    .hist-stat-num.red{{color:#7a0e0e}}
    .hist-stat-label{{font-size:11px;font-weight:600;letter-spacing:.06em;text-transform:uppercase;color:var(--text-secondary);margin-top:4px}}
    .hist-toolbar{{background:var(--glass);backdrop-filter:blur(20px) saturate(1.5);border:1.5px solid var(--glass-border);border-radius:18px 18px 0 0;padding:16px 20px;display:flex;gap:12px;flex-wrap:wrap;align-items:center;box-shadow:inset 0 1px 0 rgba(255,255,255,.55)}}
    .hist-search-form{{flex:1;min-width:200px;display:flex;gap:8px}}
    .hist-search-form input[type=text]{{flex:1}}
    .hist-filters{{display:flex;gap:6px;align-items:center}}
    .hist-tab{{padding:6px 14px;border-radius:20px;font-size:13px;font-weight:500;text-decoration:none;color:var(--text-secondary);transition:all .15s ease;border:1.5px solid transparent}}
    .hist-tab:hover{{background:rgba(255,255,255,.18);color:var(--text-primary);border-color:var(--glass-border)}}
    .hist-tab-active{{background:rgba(12,39,24,.18);color:var(--green);border-color:rgba(12,39,24,.28);font-weight:700}}
    .hist-card{{background:var(--glass);backdrop-filter:blur(20px) saturate(1.5);border:1.5px solid var(--glass-border);border-top:none;border-radius:0 0 18px 18px;overflow:hidden;box-shadow:0 8px 32px rgba(0,0,0,.08)}}
    .hist-table-wrap{{overflow-x:auto}}
    .hist-table{{width:100%;border-collapse:collapse}}
    .hist-thead-row{{background:rgba(255,255,255,.18);border-bottom:1.5px solid var(--glass-border)}}
    .hist-th{{padding:10px 16px;font-size:11px;font-weight:700;letter-spacing:.06em;text-transform:uppercase;color:var(--text-secondary);text-align:left;white-space:nowrap}}
    .hist-row{{border-bottom:1px solid rgba(255,255,255,.18);transition:background .12s ease}}
    .hist-row:last-child{{border-bottom:none}}
    .hist-row:hover{{background:rgba(255,255,255,.12)}}
    .hist-td{{padding:13px 16px;vertical-align:top}}
    .hist-title{{font-weight:600;font-size:14px;color:var(--text-primary);margin-bottom:2px}}
    .hist-company{{font-size:13px;color:var(--text-secondary)}}
    .hist-notes{{font-size:11px;color:var(--text-tertiary);margin-top:3px;font-style:italic}}
    .hist-url-td{{max-width:180px}}
    .hist-link{{font-size:12px;color:var(--green-mid);text-decoration:none;word-break:break-all;opacity:.85}}
    .hist-link:hover{{opacity:1;text-decoration:underline}}
    .hist-dash{{color:var(--text-tertiary)}}
    .hist-date{{font-size:12px;color:var(--text-tertiary);font-family:'JetBrains Mono',monospace;white-space:nowrap}}
    .hist-empty{{padding:60px 20px;text-align:center;color:var(--text-secondary)}}
    .hist-empty p{{margin:0;font-size:15px}}
    .hist-pager{{display:flex;gap:6px;justify-content:center;padding:18px 16px;border-top:1px solid rgba(255,255,255,.18)}}
    .pg-link{{padding:6px 12px;border-radius:10px;font-size:13px;font-weight:500;text-decoration:none;color:var(--text-secondary);background:rgba(255,255,255,.14);border:1.5px solid rgba(255,255,255,.22);transition:all .15s ease}}
    .pg-link:hover{{background:rgba(255,255,255,.26);color:var(--text-primary)}}
    .pg-active{{background:rgba(12,39,24,.2);color:var(--green);border-color:rgba(12,39,24,.3);font-weight:700}}
    .hist-back{{margin-top:18px}}
    .hist-back a{{color:var(--green-mid);font-size:14px;text-decoration:none;font-weight:500;opacity:.8}}
    .hist-back a:hover{{opacity:1}}
    </style>
    <div class="hist-wrap">
      <div class="hist-header">
        <div class="hist-title-group">
          <h1>Application History</h1>
          <p>All roles applied by your agent</p>
        </div>
        <div class="hist-stats">
          <div class="hist-stat">
            <div class="hist-stat-num green">{n_applied}</div>
            <div class="hist-stat-label">Applied</div>
          </div>
          <div class="hist-stat">
            <div class="hist-stat-num red">{n_failed}</div>
            <div class="hist-stat-label">Failed</div>
          </div>
          <div class="hist-stat">
            <div class="hist-stat-num">{n_applied + n_failed}</div>
            <div class="hist-stat-label">Total</div>
          </div>
        </div>
      </div>

      <div class="hist-toolbar">
        <form method="GET" action="/history" class="hist-search-form">
          <input type="text" name="q" value="{q}" placeholder="Search by role or company…">
          <input type="hidden" name="status" value="{status}">
          <button type="submit" class="btn btn-primary">Search</button>
        </form>
        <div class="hist-filters">
          <a href="{filter_qs('')}" class="{active_tab('')}">All</a>
          <a href="{filter_qs('applied')}" class="{active_tab('applied')}">Applied ({n_applied})</a>
          <a href="{filter_qs('sponsor_review')}" class="{active_tab('sponsor_review')}" style="{'color:#7c3aed;font-weight:700' if status=='sponsor_review' else ''}">&#127919; Visa Sponsor ({n_sponsor})</a>
          <a href="{filter_qs('failed')}" class="{active_tab('failed')}">Failed ({n_failed})</a>
        </div>
      </div>

      <div class="hist-card">
        {table_html}
        {pager}
      </div>

      <div class="hist-back"><a href="/dashboard">← Back to Dashboard</a></div>
    </div>"""
    return _page("History", body, user=u)


# ── Pipeline trigger ──────────────────────────────────────────────────────────

def _run_pipeline_bg(user_id: int, run_db_id: int):
    """Background thread: fetch → score → tailor → cover letter → email report."""
    import sqlite3 as _sq3

    conn = _sq3.connect(str(DB_PATH), check_same_thread=False)
    conn.row_factory = _sq3.Row

    def update(status, **kw):
        sets = ", ".join(f"{k}=?" for k in kw)
        vals = list(kw.values()) + [run_db_id]
        conn.execute(f"UPDATE runs SET {sets} WHERE id=?", vals)
        conn.commit()

    try:
        u = dict(conn.execute("SELECT * FROM users WHERE id=?", (user_id,)).fetchone())
        user_dir = _user_dir(user_id)
        tmp_dir  = user_dir / ".tmp"

        # Write a temporary .env-like env for the tools
        env_patch = {
            "CANDIDATE_FIRST_NAME": u["first_name"],
            "CANDIDATE_LAST_NAME":  u["last_name"],
            "CANDIDATE_EMAIL":      u["email"],
            "CANDIDATE_PHONE":      u["phone"],
            "CV_PATH":              str(user_dir / "cv.docx"),
            "TMP_DIR":              str(tmp_dir),
            "CANDIDATE_PROFILE_PATH": str(user_dir / "candidate_profile.md"),
            "JOB_SEARCH_KEYWORDS":  u["keywords"],
            "JOB_SEARCH_LOCATION":  u["search_location"],
            "SCORE_THRESHOLD":      str(u["threshold"]),
            "EMAIL_SUBJECT_PREFIX": u["email_subject"] or "Job Application",
            "SMTP_EMAIL":           u["email"],
            "SMTP_PASSWORD":        _dec(u["smtp_password"]),
            "SMTP_TO":              u["email"],
            "LINKEDIN_EMAIL":       u["linkedin_email"] or u["email"],
            "LINKEDIN_PASSWORD":    _dec(u["linkedin_pass"]),
            "REED_EMAIL":           u["reed_email"] or u["email"],
            "REED_PASSWORD":        _dec(u["reed_pass"]),
            "INDEED_EMAIL":         u["indeed_email"] or u["email"],
            "INDEED_PASSWORD":      _dec(u["indeed_pass"]),
        }
        # Patch environment for this thread
        os.environ.update({k: v or "" for k, v in env_patch.items()})

        # Write candidate_profile.md if missing
        profile_path = user_dir / "candidate_profile.md"
        if not profile_path.exists():
            profile_path.write_text(f"""## Personal Info
- Name: {u['first_name']} {u['last_name']}
- Email: {u['email']}
- Phone: {u['phone'] or ''}
- Location: {u['location'] or 'United Kingdom'}

## Target Roles
- Job titles: {u['keywords']}
- Work arrangement: Hybrid
- Seniority level:
- Salary expectations:
- Right to work in UK: Yes

## CV Summary
(Auto-generated on first run.)

## Career Goals
- Seeking: {u['keywords']}
- Location: {u['search_location']}
""")

        run_id = conn.execute("SELECT run_id FROM runs WHERE id=?", (run_db_id,)).fetchone()["run_id"]

        # ── 1. Fetch ──────────────────────────────────────────────────────────
        update(status="running", jobs_found=0)
        from tools.fetch_jobs import fetch_all_jobs
        jobs_raw = fetch_all_jobs(
            u["keywords"],
            u["search_location"],
            int(os.getenv("JOB_SEARCH_COUNT", "15"))
        )
        (tmp_dir / "jobs_raw.json").write_text(json.dumps(jobs_raw, indent=2))
        update(status="running", jobs_found=len(jobs_raw))

        # ── 2. Score ──────────────────────────────────────────────────────────
        from tools.score_job import score_job
        profile_text = profile_path.read_text()
        threshold    = u["threshold"]
        scored = []
        for job in jobs_raw:
            try:
                result = score_job(job, profile_text)
                job["score"] = result.get("score", 0)
                job["score_reason"] = result.get("reason", "")
            except Exception as e:
                job["score"] = 0
                job["score_reason"] = str(e)
            scored.append(job)
            time.sleep(2)  # Groq free tier: ~30 req/min — 2s gap keeps well within limit

        (tmp_dir / "jobs_scored.json").write_text(json.dumps(scored, indent=2))
        qualifying = [j for j in scored if j["score"] >= threshold]
        # Cap at 5 jobs to keep runtime under 15 minutes on Groq free tier
        qualifying = sorted(qualifying, key=lambda j: j["score"], reverse=True)[:5]

        if not qualifying:
            update(status="completed", finished_at=datetime.utcnow().isoformat(),
                   jobs_applied=0, report_json=json.dumps({"jobs": [], "run_id": run_id}))
            conn.close()
            return

        # ── 3. Tailor CV + cover letter ───────────────────────────────────────
        from tools.tailor_cv_docx    import tailor_cv_docx
        from tools.generate_cv_pdf   import generate_cv_pdf
        from tools.generate_cover_letter import generate_cover_letter

        def _slug(t, n=22):
            return re.sub(r"[^\w]", "_", t.lower())[:n].strip("_")

        report = {"run_id": run_id, "keywords": u["keywords"],
                  "location": u["search_location"], "threshold": threshold, "jobs": []}

        for job in qualifying:
            title   = job.get("title", "Unknown")
            company = job.get("company", "Unknown")
            desc    = job.get("description", "")

            record = {
                "title": title, "company": company,
                "url": job.get("url",""), "source": job.get("source",""),
                "score": job.get("score", 0),
                "cv_docx": None, "cv_pdf": None, "cover_letter_path": None,
                "status": "needs_review", "ats": "unknown", "notes": "",
            }

            cv_out = str(tmp_dir / f"cv_{_slug(company)}_{_slug(title)}.docx")
            try:
                tailor_cv_docx(title, company, desc, cv_out)
                record["cv_docx"] = cv_out
            except Exception as e:
                record["cv_docx"] = str(user_dir / "cv.docx")
                record["notes"] += f"CV tailor error: {e}. "

            if record["cv_docx"]:
                pdf_out = record["cv_docx"].replace(".docx", ".pdf")
                try:
                    generate_cv_pdf(record["cv_docx"], pdf_out)
                    record["cv_pdf"] = pdf_out
                except Exception as e:
                    record["notes"] += f"PDF error: {e}. "

            time.sleep(2)  # gap between tailor and cover letter LLM calls
            cl_out = str(tmp_dir / f"cl_{_slug(company)}_{_slug(title)}.txt")
            try:
                cover_letter = generate_cover_letter(job, profile_text)
                Path(cl_out).write_text(cover_letter)
                record["cover_letter_path"] = cl_out
            except Exception as e:
                record["notes"] += f"Cover letter error: {e}. "

            report["jobs"].append(record)
            time.sleep(2)  # gap between jobs

        # ── 4. Save report + email ────────────────────────────────────────────
        report_path = tmp_dir / f"application_report_{run_id}.json"
        report_path.write_text(json.dumps(report, indent=2))

        if _dec(u["smtp_password"]):
            try:
                from tools.send_application_report import send_report
                send_report(report)
                log.info(f"Report emailed to {u['email']}")
            except Exception as e:
                log.warning(f"Email failed: {e}")

        applied = len([j for j in report["jobs"] if j["status"] != "needs_review"])
        update(status="completed", finished_at=datetime.utcnow().isoformat(),
               jobs_applied=len(qualifying), report_json=report_path.read_text())
        log.info(f"Pipeline complete: {len(report['jobs'])} jobs tailored")

        # Send completion email via Sendgrid
        n = len(report["jobs"])
        dashboard_url = os.getenv("APP_URL", "http://localhost:5001") + f"/run/{run_db_id}"
        job_rows = "".join(
            f'<tr><td style="padding:8px 12px;border-bottom:1px solid #e2e8f0">{j.get("title","")}</td>'
            f'<td style="padding:8px 12px;border-bottom:1px solid #e2e8f0;color:#64748b">{j.get("company","")}</td>'
            f'<td style="padding:8px 12px;border-bottom:1px solid #e2e8f0;color:#3b82f6">{j.get("score","")}/10</td></tr>'
            for j in report["jobs"]
        )
        _send_email(
            u["email"],
            f"ApplyExpress: pipeline complete — {n} job{'s' if n != 1 else ''} tailored",
            f"""<div style="font-family:system-ui,sans-serif;max-width:600px;margin:0 auto;padding:32px">
            <h2 style="color:#1e293b">Your pipeline just ran ✓</h2>
            <p style="color:#64748b;margin:12px 0 24px">
              {n} tailored CV{'s' if n != 1 else ''} and cover letter{'s' if n != 1 else ''} are ready.
              The Chrome extension will submit them automatically.
            </p>
            <table style="width:100%;border-collapse:collapse;border:1px solid #e2e8f0;border-radius:8px;overflow:hidden">
              <thead><tr style="background:#f8fafc">
                <th style="padding:10px 12px;text-align:left;font-size:13px;color:#475569">Role</th>
                <th style="padding:10px 12px;text-align:left;font-size:13px;color:#475569">Company</th>
                <th style="padding:10px 12px;text-align:left;font-size:13px;color:#475569">Score</th>
              </tr></thead>
              <tbody>{job_rows}</tbody>
            </table>
            <p style="margin-top:24px">
              <a href="{dashboard_url}" style="background:#3b82f6;color:#fff;padding:12px 24px;
                border-radius:8px;text-decoration:none;font-weight:600">View full results →</a>
            </p>
            <p style="font-size:12px;color:#94a3b8;margin-top:32px">ApplyExpress · <a href="{os.getenv('APP_URL','')}/unsubscribe">Unsubscribe</a></p>
            </div>""",
            smtp_password=_dec(u.get("smtp_password", "")),
        )

    except Exception as e:
        log.error(f"Pipeline failed for user {user_id}: {e}")
        try:
            update(status="failed", finished_at=datetime.utcnow().isoformat())
            _send_email(
                u["email"],
                "ApplyExpress: pipeline encountered an error",
                f"<p>Your pipeline run failed with error: <code>{e}</code></p>"
                f"<p>We'll retry automatically. If this keeps happening, email support@applyexpress.io</p>",
                smtp_password=_dec(u.get("smtp_password", "")),
            )
        except Exception:
            pass
    finally:
        conn.close()


@app.route("/run", methods=["POST"])
@login_required
@paid_required
def trigger_run():
    u            = current_user()
    db           = get_db()
    specialty_id = request.form.get("specialty_id","").strip() or None

    # Determine which CV to check
    if specialty_id:
        spec     = db.execute("SELECT * FROM specialties WHERE id=? AND user_id=?",
                              (specialty_id, u["id"])).fetchone()
        cv_check = _specialty_dir(u["id"], spec["slug"]) / "cv.docx" if spec else None
        if not cv_check or not cv_check.exists():
            return redirect(url_for("profile") + "#specialties")
    else:
        if not (_user_dir(u["id"]) / "cv.docx").exists():
            return redirect(url_for("profile") + "#cv")

    run_id = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    row = db.execute("""
        INSERT INTO runs (user_id, run_id, started_at, status, specialty_id)
        VALUES (?,?,?,?,?) RETURNING id
    """, (u["id"], run_id, datetime.utcnow().isoformat(), "running", specialty_id)).fetchone()
    db.commit()
    run_db_id = row["id"]

    import subprocess
    worker = Path(__file__).parent / "tools" / "run_pipeline_worker.py"
    subprocess.Popen(
        [sys.executable, str(worker), str(u["id"]), str(run_db_id),
         str(DB_PATH), str(ROOT), str(specialty_id or "")],
        start_new_session=True,
        stdout=open(str(DATA_DIR / f"pipeline_{run_db_id}.log"), "w"),
        stderr=subprocess.STDOUT,
    )

    return redirect(url_for("dashboard") + "?running=1")


# ── Chrome Extension API ──────────────────────────────────────────────────────

@app.route("/api/health")
def api_health():
    return jsonify({"status": "ok", "service": "ApplyExpress SaaS"})


@app.route("/api/run-status")
@login_required
def api_run_status():
    """Lightweight poll endpoint — returns current run status for the dashboard."""
    u  = current_user()
    db = get_db()
    r  = db.execute(
        "SELECT id, status, started_at, jobs_found, jobs_applied FROM runs WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (u["id"],)
    ).fetchone()
    if not r:
        return jsonify({"status": "none"})
    elapsed = ""
    if r["started_at"]:
        try:
            delta = datetime.utcnow() - datetime.fromisoformat(r["started_at"])
            mins  = int(delta.total_seconds() // 60)
            secs  = int(delta.total_seconds() % 60)
            elapsed = f"{mins}m {secs}s"
        except Exception:
            pass
    return jsonify({
        "status":      r["status"],
        "run_id":      r["id"],
        "elapsed":     elapsed,
        "jobs_found":  r["jobs_found"],
        "jobs_applied":r["jobs_applied"],
    })


@app.route("/api/profile")
@api_key_required
def api_profile():
    u = g.api_user
    return jsonify({
        "firstName": u["first_name"], "lastName":  u["last_name"],
        "fullName":  f"{u['first_name']} {u['last_name']}".strip(),
        "email":     u["email"],      "phone":     u["phone"],
    })


@app.route("/api/cv")
@api_key_required
def api_cv():
    import base64
    u       = g.api_user
    cv_path = _user_dir(u["id"]) / "cv.docx"
    if not cv_path.exists():
        return jsonify({"error": "No CV uploaded"}), 404
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    data = base64.b64encode(cv_path.read_bytes()).decode("utf-8")
    return jsonify({"filename": cv_path.name, "mime": mime,
                    "data": f"data:{mime};base64,{data}", "size": cv_path.stat().st_size})


@app.route("/api/credentials")
@api_key_required
def api_credentials():
    u = g.api_user
    return jsonify({
        "reed":     {"email": u["reed_email"]     or u["email"], "password": _dec(u["reed_pass"])},
        "linkedin": {"email": u["linkedin_email"] or u["email"], "password": _dec(u["linkedin_pass"])},
        "indeed":   {"email": u["indeed_email"]   or u["email"], "password": _dec(u["indeed_pass"])},
    })


@app.route("/api/jobs")
@api_key_required
def api_jobs():
    u         = g.api_user
    limit     = int(request.args.get("limit", 20))
    # ?min_score=0 overrides threshold (useful for debugging / manual review)
    min_score = request.args.get("min_score")
    threshold = int(min_score) if min_score is not None else u["threshold"]
    tmp_dir   = _user_dir(u["id"]) / ".tmp"
    scored    = tmp_dir / "jobs_scored.json"

    if not scored.exists():
        return jsonify({"jobs": [], "count": 0,
                        "note": "No jobs yet. Trigger a run from your dashboard."})
    try:
        all_jobs = json.loads(scored.read_text())
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    result = []
    below_threshold = 0
    for i, job in enumerate(all_jobs[:limit]):
        score = job.get("score", 0)
        if score < threshold:
            below_threshold += 1
            continue
        # Skip jobs already successfully applied to
        result.append({
            "row": i + 2, "title": job.get("title",""),
            "company": job.get("company",""), "location": job.get("location",""),
            "url": job.get("url",""), "score": score,
            "cover_letter": job.get("cover_letter",""),
            "status": "Pending Review", "_state": "pending",
        })
    return jsonify({
        "jobs":            result,
        "count":           len(result),
        "total_scored":    len(all_jobs),
        "below_threshold": below_threshold,
        "threshold":       threshold,
    })


@app.route("/api/check_job", methods=["GET"])
@api_key_required
def api_check_job():
    """Returns {duplicate: bool, blacklisted: bool} — extension calls this before opening a tab."""
    u       = g.api_user
    db      = get_db()
    url     = request.args.get("url", "").strip()
    company = request.args.get("company", "").strip()

    duplicate = False
    if url:
        duplicate = bool(db.execute(
            "SELECT id FROM applications WHERE user_id=? AND url=? AND status='applied'",
            (u["id"], url)
        ).fetchone())

    blacklisted = False
    if company:
        blacklisted = bool(db.execute(
            "SELECT id FROM blacklist WHERE user_id=? AND lower(company)=lower(?)",
            (u["id"], company)
        ).fetchone())

    return jsonify({"duplicate": duplicate, "blacklisted": blacklisted})


@app.route("/api/update_status", methods=["POST"])
@api_key_required
def api_update_status():
    u    = g.api_user
    data = request.get_json(force=True) or {}
    db   = get_db()

    # Find most recent run for this user to link the application record
    run  = db.execute("SELECT id FROM runs WHERE user_id=? ORDER BY id DESC LIMIT 1",
                      (u["id"],)).fetchone()
    run_db_id = run["id"] if run else None

    raw_status = (data.get("status") or "").lower()
    status     = "applied" if raw_status in ("applied", "✓ applied") else "failed"
    job_url    = data.get("url", "").strip()
    company    = data.get("company", "").strip()
    title      = data.get("title", "").strip()

    # Deduplication — skip if this URL was already successfully applied to
    if job_url and status == "applied":
        already = db.execute(
            "SELECT id FROM applications WHERE user_id=? AND url=? AND status='applied'",
            (u["id"], job_url)
        ).fetchone()
        if already:
            log.info(f"[EXT] Duplicate skipped: {title} @ {company} ({job_url[:60]})")
            return jsonify({"success": True, "duplicate": True})

    # Blacklist check — don't record if company is blacklisted (extension should
    # have already skipped it, but guard here too)
    if company:
        blocked = db.execute(
            "SELECT id FROM blacklist WHERE user_id=? AND lower(company)=lower(?)",
            (u["id"], company)
        ).fetchone()
        if blocked:
            log.info(f"[EXT] Blacklisted company skipped: {company}")
            return jsonify({"success": True, "blacklisted": True})

    db.execute("""INSERT INTO applications
                  (user_id, run_db_id, title, company, url, status, notes, applied_at)
                  VALUES (?,?,?,?,?,?,?,?)""",
               (u["id"], run_db_id, title, company, job_url, status,
                data.get("notes", ""), datetime.utcnow().isoformat()))

    if run_db_id and status == "applied":
        db.execute("UPDATE runs SET jobs_applied = jobs_applied + 1 WHERE id=?", (run_db_id,))

    db.commit()
    log.info(f"[EXT] Status: {title} @ {company} → {status}")
    return jsonify({"success": True})


@app.route("/design-system")
def design_system():
    """Brand + design system review page."""
    body = """
<style>
/* ── Page-specific extras ─────────────────────── */
.ds-section{margin-bottom:72px}
.ds-section-label{font-size:11px;font-weight:600;color:var(--accent);letter-spacing:.12em;
  text-transform:uppercase;margin-bottom:10px;display:flex;align-items:center;gap:8px}
.ds-section-label::before{content:'';width:20px;height:1px;background:var(--accent)}
.ds-section-label::after{content:'';flex:1;height:1px;background:var(--border)}
.ds-grid-2{display:grid;grid-template-columns:1fr 1fr;gap:20px}
.ds-grid-3{display:grid;grid-template-columns:1fr 1fr 1fr;gap:20px}
@media(max-width:700px){.ds-grid-2,.ds-grid-3{grid-template-columns:1fr}}

/* ── Logo cards ──────────────────────────────── */
.logo-card{
  background:var(--bg-surface);border:1px solid var(--border);border-radius:20px;
  overflow:hidden;transition:border-color .2s,transform .2s;cursor:pointer;
}
.logo-card:hover{border-color:var(--border-bright);transform:translateY(-3px)}
.logo-card.selected{border-color:var(--accent);box-shadow:0 0 0 1px var(--accent),0 0 24px var(--accent-glow)}
.logo-canvas{
  height:200px;display:flex;align-items:center;justify-content:center;
  flex-direction:column;gap:16px;position:relative;overflow:hidden;
}
.logo-canvas-1{background:radial-gradient(ellipse at 50% 60%,rgba(0,200,255,.06) 0%,transparent 70%)}
.logo-canvas-2{background:radial-gradient(ellipse at 50% 60%,rgba(16,214,138,.06) 0%,transparent 70%)}
.logo-canvas-3{background:radial-gradient(ellipse at 50% 60%,rgba(245,166,35,.06) 0%,transparent 70%)}
.logo-info{padding:20px 24px;border-top:1px solid var(--border)}
.logo-slogan{font-size:13px;color:var(--text-tertiary);margin-top:6px;font-style:italic;line-height:1.5}

/* ── Pipeline animation ──────────────────────── */
.pipeline{display:flex;align-items:center;gap:0;margin:32px 0;overflow-x:auto;padding:4px}
.pipe-step{
  flex:1;min-width:120px;display:flex;flex-direction:column;align-items:center;
  gap:10px;padding:20px 12px;border-radius:16px;border:1px solid var(--border);
  background:var(--bg-surface);position:relative;transition:all .3s;
  opacity:.35;
}
.pipe-step.active{opacity:1;border-color:var(--accent);background:var(--accent-surface)}
.pipe-step.done{opacity:.7;border-color:var(--border-bright)}
.pipe-arrow{width:32px;flex-shrink:0;display:flex;align-items:center;justify-content:center;
            color:var(--text-tertiary);font-size:18px;margin:0 -1px}
.pipe-icon{font-size:28px}
.pipe-label{font-size:12px;font-weight:600;text-align:center;color:var(--text-primary)}
.pipe-sub{font-size:11px;text-align:center;color:var(--text-tertiary);line-height:1.4}
.pipe-step.active .pipe-label{color:var(--accent)}
.pipe-step.active .pipe-sub{color:var(--text-secondary)}

/* ── Colour swatches ─────────────────────────── */
.swatch-row{display:flex;align-items:center;gap:16px;padding:12px 0;border-bottom:1px solid var(--border)}
.swatch-row:last-child{border-bottom:none}
.swatch{width:40px;height:40px;border-radius:10px;flex-shrink:0;border:1px solid rgba(255,255,255,.06)}
.swatch-info{flex:1}
.swatch-name{font-family:var(--font-mono);font-size:12px;color:var(--text-primary);margin-bottom:3px}
.swatch-desc{font-size:12px;color:var(--text-tertiary)}
.swatch-hex{font-family:var(--font-mono);font-size:11px;color:var(--text-tertiary);margin-left:auto}

/* ── Placeholder animation ───────────────────── */
@keyframes blink{0%,100%{opacity:1}50%{opacity:0}}
.cursor{display:inline-block;width:2px;height:1em;background:var(--accent);
        margin-left:1px;vertical-align:middle;animation:blink 1s ease-in-out infinite}
.typed-placeholder{font-family:var(--font-mono);font-size:13px;color:var(--text-tertiary)}

/* ── Logo SVG animations ─────────────────────── */
@keyframes draw-path{from{stroke-dashoffset:200}to{stroke-dashoffset:0}}
@keyframes pulse-ring{0%,100%{r:8;opacity:1}50%{r:12;opacity:.4}}
@keyframes spin-target{from{transform:rotate(0deg)}to{transform:rotate(360deg)}}
@keyframes pulse-signal{0%,100%{opacity:.3}50%{opacity:1}}
.logo-path-anim{stroke-dasharray:200;stroke-dashoffset:200;animation:draw-path 1.2s var(--ease-spring) forwards}
.ring-pulse{animation:pulse-ring 2s ease-in-out infinite;transform-origin:28px 28px}

/* ── Annotation callout ──────────────────────── */
.callout{
  display:flex;gap:12px;align-items:flex-start;
  background:var(--bg-elevated);border-radius:var(--r);
  padding:14px 16px;margin-top:12px;border-left:3px solid var(--accent);
}
.callout-icon{font-size:16px;flex-shrink:0;margin-top:1px}
.callout p{font-size:13px;color:var(--text-secondary);line-height:1.6;margin:0}
</style>

<div class="container-lg" style="padding-top:52px;padding-bottom:80px">

  <!-- ── Page header ───────────────────────────────────────────────── -->
  <div style="margin-bottom:56px">
    <div style="font-size:11px;font-weight:600;color:var(--accent);letter-spacing:.12em;text-transform:uppercase;margin-bottom:12px">Brand &amp; Design Review</div>
    <h1 style="font-size:38px;margin-bottom:12px">ApplyExpress<br><span style="color:var(--text-secondary);font-size:24px;font-weight:500">Brand + System Guide</span></h1>
    <p style="color:var(--text-secondary);font-size:16px;max-width:540px;line-height:1.7">
      Three logo options to choose from. Animated product demo. Every colour, font, and component explained in plain English.
    </p>
    <div style="display:flex;gap:12px;margin-top:24px;flex-wrap:wrap">
      <a href="#logos" class="btn btn-primary btn-sm">Choose a logo →</a>
      <a href="#pipeline" class="btn btn-outline btn-sm">See how it works</a>
      <a href="#components" class="btn btn-outline btn-sm">Components</a>
    </div>
  </div>

  <!-- ════════════════════════════════════════════════════════════════ -->
  <!-- SECTION 1: Logo Options -->
  <!-- ════════════════════════════════════════════════════════════════ -->
  <div class="ds-section" id="logos">
    <div class="ds-section-label">01 · Logo Options</div>
    <h2 style="margin-bottom:8px">Choose Your Brand Identity</h2>
    <p style="color:var(--text-secondary);font-size:14px;margin-bottom:28px">Three distinct directions. Click one to select it. Each includes the logo mark, wordmark, and slogan.</p>

    <div class="ds-grid-3">

      <!-- Logo 1: THE ARROW (speed + direction) -->
      <div class="logo-card" id="logo-opt-1" onclick="selectLogo(1)">
        <div class="logo-canvas logo-canvas-1">
          <svg viewBox="0 0 64 64" width="72" height="72" xmlns="http://www.w3.org/2000/svg">
            <defs>
              <linearGradient id="arrow-grad" x1="0" y1="0" x2="0" y2="1">
                <stop offset="0%" stop-color="#00c8ff"/>
                <stop offset="100%" stop-color="#0095cc"/>
              </linearGradient>
              <filter id="arrow-glow">
                <feGaussianBlur stdDeviation="2" result="blur"/>
                <feMerge><feMergeNode in="blur"/><feMergeNode in="SourceGraphic"/></feMerge>
              </filter>
            </defs>
            <!-- A left leg -->
            <path class="logo-path-anim" d="M32 8 L10 52" stroke="url(#arrow-grad)" stroke-width="5" fill="none" stroke-linecap="round" filter="url(#arrow-glow)"/>
            <!-- A right leg -->
            <path class="logo-path-anim" d="M32 8 L54 52" stroke="url(#arrow-grad)" stroke-width="5" fill="none" stroke-linecap="round" filter="url(#arrow-glow)" style="animation-delay:.2s"/>
            <!-- Crossbar -->
            <path class="logo-path-anim" d="M19 38 L45 38" stroke="#00c8ff" stroke-width="4" fill="none" stroke-linecap="round" style="animation-delay:.4s"/>
            <!-- Apex glow dot -->
            <circle cx="32" cy="8" r="4" fill="#00c8ff" opacity="0" style="animation:fade-up .4s .6s forwards">
              <animate attributeName="opacity" values="0;1" dur="0.4s" begin="0.6s" fill="freeze"/>
              <animate attributeName="r" values="4;6;4" dur="2s" begin="1s" repeatCount="indefinite"/>
            </circle>
          </svg>
          <div style="font-family:var(--font-display);font-size:22px;font-weight:700;letter-spacing:-.02em;color:var(--text-primary)">
            Apply<span style="color:var(--accent)">Express</span>
          </div>
        </div>
        <div class="logo-info">
          <div style="font-size:13px;font-weight:600;color:var(--text-primary);margin-bottom:4px">Option A · "The Arrow"</div>
          <div class="logo-slogan">"Land your next role on autopilot."</div>
          <div style="display:flex;gap:6px;margin-top:14px">
            <div style="width:20px;height:20px;border-radius:50%;background:#080c12;border:1px solid var(--border)"></div>
            <div style="width:20px;height:20px;border-radius:50%;background:#00c8ff"></div>
            <div style="width:20px;height:20px;border-radius:50%;background:#e8f1ff"></div>
          </div>
          <div style="margin-top:14px;font-size:11px;color:var(--text-tertiary)">Clean · Fast · Directional · Works at any size</div>
        </div>
      </div>

      <!-- Logo 2: THE TARGET (precision + control) -->
      <div class="logo-card" id="logo-opt-2" onclick="selectLogo(2)">
        <div class="logo-canvas logo-canvas-2">
          <svg viewBox="0 0 64 64" width="72" height="72" xmlns="http://www.w3.org/2000/svg">
            <defs>
              <filter id="target-glow">
                <feGaussianBlur stdDeviation="1.5" result="blur"/>
                <feMerge><feMergeNode in="blur"/><feMergeNode in="SourceGraphic"/></feMerge>
              </filter>
            </defs>
            <!-- Outer ring (dim) -->
            <circle cx="32" cy="32" r="28" stroke="#1e2d42" stroke-width="1.5" fill="none"/>
            <!-- Middle ring -->
            <circle cx="32" cy="32" r="19" stroke="#2d4060" stroke-width="1.5" fill="none"/>
            <!-- Inner ring (accent, animated) -->
            <circle cx="32" cy="32" r="10" stroke="#10d68a" stroke-width="2" fill="none" filter="url(#target-glow)">
              <animate attributeName="r" values="10;12;10" dur="2.5s" repeatCount="indefinite"/>
              <animate attributeName="opacity" values="1;.6;1" dur="2.5s" repeatCount="indefinite"/>
            </circle>
            <!-- Center dot -->
            <circle cx="32" cy="32" r="3.5" fill="#10d68a" filter="url(#target-glow)"/>
            <!-- Crosshairs -->
            <line x1="32" y1="2" x2="32" y2="20" stroke="#526278" stroke-width="1.5" stroke-linecap="round"/>
            <line x1="32" y1="44" x2="32" y2="62" stroke="#526278" stroke-width="1.5" stroke-linecap="round"/>
            <line x1="2" y1="32" x2="20" y2="32" stroke="#526278" stroke-width="1.5" stroke-linecap="round"/>
            <line x1="44" y1="32" x2="62" y2="32" stroke="#526278" stroke-width="1.5" stroke-linecap="round"/>
          </svg>
          <div style="font-family:var(--font-display);font-size:22px;font-weight:700;letter-spacing:-.02em;color:var(--text-primary)">
            APPLY<span style="color:#10d68a">EXPRESS</span>
          </div>
        </div>
        <div class="logo-info">
          <div style="font-size:13px;font-weight:600;color:var(--text-primary);margin-bottom:4px">Option B · "The Target"</div>
          <div class="logo-slogan">"Your unfair advantage in every job market."</div>
          <div style="display:flex;gap:6px;margin-top:14px">
            <div style="width:20px;height:20px;border-radius:50%;background:#080c12;border:1px solid var(--border)"></div>
            <div style="width:20px;height:20px;border-radius:50%;background:#10d68a"></div>
            <div style="width:20px;height:20px;border-radius:50%;background:#e8f1ff"></div>
          </div>
          <div style="margin-top:14px;font-size:11px;color:var(--text-tertiary)">Precise · Professional · Technical · Trustworthy</div>
        </div>
      </div>

      <!-- Logo 3: THE SIGNAL (automation + flow) -->
      <div class="logo-card" id="logo-opt-3" onclick="selectLogo(3)">
        <div class="logo-canvas logo-canvas-3">
          <svg viewBox="0 0 96 48" width="108" height="54" xmlns="http://www.w3.org/2000/svg">
            <defs>
              <linearGradient id="sig-grad" x1="0" y1="0" x2="1" y2="0">
                <stop offset="0%" stop-color="#1e2d42"/>
                <stop offset="35%" stop-color="#f5a623"/>
                <stop offset="65%" stop-color="#f5a623"/>
                <stop offset="100%" stop-color="#1e2d42"/>
              </linearGradient>
              <filter id="sig-glow">
                <feGaussianBlur stdDeviation="2" result="blur"/>
                <feMerge><feMergeNode in="blur"/><feMergeNode in="SourceGraphic"/></feMerge>
              </filter>
            </defs>
            <!-- Signal baseline + spike (EKG-style) -->
            <polyline points="0,24 22,24 28,24 32,6 37,40 42,18 47,24 55,24 96,24"
              stroke="url(#sig-grad)" stroke-width="2.5" fill="none"
              stroke-linecap="round" stroke-linejoin="round" filter="url(#sig-glow)">
              <animate attributeName="stroke-dasharray" values="0,300;300,0" dur="1.8s" begin="0s" fill="freeze"/>
              <animate attributeName="stroke-dashoffset" values="0" dur="1.8s" begin="0s" fill="freeze"/>
            </polyline>
            <!-- Glowing dot at peak -->
            <circle cx="32" cy="6" r="3" fill="#f5a623" filter="url(#sig-glow)">
              <animate attributeName="opacity" values="0;1" dur=".3s" begin="0.9s" fill="freeze"/>
              <animate attributeName="r" values="3;5;3" dur="1.8s" begin="1.2s" repeatCount="indefinite"/>
            </circle>
          </svg>
          <div style="font-family:var(--font-display);font-size:22px;font-weight:700;letter-spacing:-.02em;color:var(--text-primary)">
            Apply<span style="color:#f5a623">Express</span>
          </div>
        </div>
        <div class="logo-info">
          <div style="font-size:13px;font-weight:600;color:var(--text-primary);margin-bottom:4px">Option C · "The Signal"</div>
          <div class="logo-slogan">"From search to offer. Completely automated."</div>
          <div style="display:flex;gap:6px;margin-top:14px">
            <div style="width:20px;height:20px;border-radius:50%;background:#080c12;border:1px solid var(--border)"></div>
            <div style="width:20px;height:20px;border-radius:50%;background:#f5a623"></div>
            <div style="width:20px;height:20px;border-radius:50%;background:#e8f1ff"></div>
          </div>
          <div style="margin-top:14px;font-size:11px;color:var(--text-tertiary)">Dynamic · Energetic · Signal/Momentum · Bold</div>
        </div>
      </div>

    </div><!-- /ds-grid-3 -->

    <!-- Selection feedback -->
    <div id="logo-selected-msg" style="display:none;margin-top:20px" class="alert alert-success">
      ✓ Logo selected — noted. This can be applied to all pages once you confirm.
    </div>
  </div><!-- /logos -->


  <!-- ════════════════════════════════════════════════════════════════ -->
  <!-- SECTION 2: How the System Works (animated) -->
  <!-- ════════════════════════════════════════════════════════════════ -->
  <div class="ds-section" id="pipeline">
    <div class="ds-section-label">02 · Product Walkthrough</div>
    <h2 style="margin-bottom:8px">How ApplyExpress Works</h2>
    <p style="color:var(--text-secondary);font-size:14px;margin-bottom:28px">Watch the pipeline run — each step lights up in sequence, showing exactly what happens when you click "Run now".</p>

    <div class="pipeline" id="pipeline-demo">
      <div class="pipe-step" data-step="0" id="ps0">
        <div class="pipe-icon">📄</div>
        <div class="pipe-label">Upload CV</div>
        <div class="pipe-sub">Your master CV.<br>Set it once.</div>
      </div>
      <div class="pipe-arrow">→</div>
      <div class="pipe-step" data-step="1" id="ps1">
        <div class="pipe-icon">🔍</div>
        <div class="pipe-label">Search Jobs</div>
        <div class="pipe-sub">Reed, LinkedIn,<br>Indeed &amp; more</div>
      </div>
      <div class="pipe-arrow">→</div>
      <div class="pipe-step" data-step="2" id="ps2">
        <div class="pipe-icon">🧠</div>
        <div class="pipe-label">AI Scores</div>
        <div class="pipe-sub">Each job scored<br>1–10 for fit</div>
      </div>
      <div class="pipe-arrow">→</div>
      <div class="pipe-step" data-step="3" id="ps3">
        <div class="pipe-icon">✍️</div>
        <div class="pipe-label">Tailor CV</div>
        <div class="pipe-sub">Unique CV +<br>cover letter</div>
      </div>
      <div class="pipe-arrow">→</div>
      <div class="pipe-step" data-step="4" id="ps4">
        <div class="pipe-icon">⚡</div>
        <div class="pipe-label">Auto Apply</div>
        <div class="pipe-sub">Extension submits<br>through your browser</div>
      </div>
      <div class="pipe-arrow">→</div>
      <div class="pipe-step" data-step="5" id="ps5">
        <div class="pipe-icon">🎯</div>
        <div class="pipe-label">Interview!</div>
        <div class="pipe-sub">Dashboard shows<br>every result</div>
      </div>
    </div>

    <!-- Progress bar under the pipeline -->
    <div style="background:var(--bg-elevated);border-radius:4px;height:3px;overflow:hidden;margin-top:4px">
      <div id="pipeline-bar" style="height:100%;background:var(--accent);width:0%;transition:width .6s ease;border-radius:4px"></div>
    </div>
    <div style="display:flex;justify-content:space-between;margin-top:8px">
      <span id="pipeline-status" style="font-size:12px;color:var(--text-tertiary);font-family:var(--font-mono)">Starting pipeline…</span>
      <button class="btn btn-ghost btn-sm" onclick="restartPipeline()">↺ Replay</button>
    </div>

    <div class="callout" style="margin-top:24px">
      <span class="callout-icon">💡</span>
      <p>Steps 1–4 run on the server automatically, 3× per day. Step 5 needs your Chrome browser open — the extension applies using your own logged-in session, keeping everything above board.</p>
    </div>
  </div>


  <!-- ════════════════════════════════════════════════════════════════ -->
  <!-- SECTION 3: Colour System -->
  <!-- ════════════════════════════════════════════════════════════════ -->
  <div class="ds-section" id="colours">
    <div class="ds-section-label">03 · Colour System</div>
    <h2 style="margin-bottom:8px">Colours &amp; What They Mean</h2>
    <p style="color:var(--text-secondary);font-size:14px;margin-bottom:24px">Each colour has a specific job. Nothing is decorative — every choice communicates meaning.</p>

    <div class="ds-grid-2">
      <div class="card" style="padding:20px">
        <div style="font-size:12px;font-weight:600;color:var(--text-tertiary);text-transform:uppercase;letter-spacing:.06em;margin-bottom:14px">Backgrounds (darkest → lightest)</div>
        <div class="swatch-row"><div class="swatch" style="background:#080c12"></div><div class="swatch-info"><div class="swatch-name">--bg-base</div><div class="swatch-desc">Page background — the deepest layer</div></div><div class="swatch-hex">#080c12</div></div>
        <div class="swatch-row"><div class="swatch" style="background:#0e1520"></div><div class="swatch-info"><div class="swatch-name">--bg-surface</div><div class="swatch-desc">Cards and panels sit here</div></div><div class="swatch-hex">#0e1520</div></div>
        <div class="swatch-row"><div class="swatch" style="background:#141d2e"></div><div class="swatch-info"><div class="swatch-name">--bg-elevated</div><div class="swatch-desc">Hover states, active rows, inputs</div></div><div class="swatch-hex">#141d2e</div></div>
        <div class="swatch-row"><div class="swatch" style="background:#1a2640"></div><div class="swatch-info"><div class="swatch-name">--bg-overlay</div><div class="swatch-desc">Modals and dropdowns float here</div></div><div class="swatch-hex">#1a2640</div></div>
      </div>
      <div class="card" style="padding:20px">
        <div style="font-size:12px;font-weight:600;color:var(--text-tertiary);text-transform:uppercase;letter-spacing:.06em;margin-bottom:14px">Semantic Colours (signal meaning)</div>
        <div class="swatch-row"><div class="swatch" style="background:#00c8ff;box-shadow:0 0 12px rgba(0,200,255,.4)"></div><div class="swatch-info"><div class="swatch-name">--accent (Electric Cyan)</div><div class="swatch-desc">Primary action, links, focus, running status</div></div></div>
        <div class="swatch-row"><div class="swatch" style="background:#10d68a;box-shadow:0 0 12px rgba(16,214,138,.3)"></div><div class="swatch-info"><div class="swatch-name">--success (Green)</div><div class="swatch-desc">Applied, completed, paid plan, positive</div></div></div>
        <div class="swatch-row"><div class="swatch" style="background:#f5a623;box-shadow:0 0 12px rgba(245,166,35,.3)"></div><div class="swatch-info"><div class="swatch-name">--warning (Amber)</div><div class="swatch-desc">Trial period, attention needed, caution</div></div></div>
        <div class="swatch-row"><div class="swatch" style="background:#ff4d6a;box-shadow:0 0 12px rgba(255,77,106,.3)"></div><div class="swatch-info"><div class="swatch-name">--error (Red)</div><div class="swatch-desc">Failed application, connection error, delete</div></div></div>
      </div>
    </div>
    <div class="callout">
      <span class="callout-icon">🎨</span>
      <p><strong style="color:var(--text-primary)">Why dark?</strong> This product works for you in the background — it's a control room, not a greeting card. Dark backgrounds reduce eye strain during long sessions and make the cyan accent pop with maximum contrast and impact.</p>
    </div>
  </div>


  <!-- ════════════════════════════════════════════════════════════════ -->
  <!-- SECTION 4: Typography -->
  <!-- ════════════════════════════════════════════════════════════════ -->
  <div class="ds-section" id="type">
    <div class="ds-section-label">04 · Typography</div>
    <h2 style="margin-bottom:8px">Three Fonts, Three Jobs</h2>
    <p style="color:var(--text-secondary);font-size:14px;margin-bottom:24px">Each font is chosen for a specific purpose — no generic defaults.</p>

    <div class="card" style="margin-bottom:20px">
      <div style="font-family:var(--font-display);font-size:42px;font-weight:700;letter-spacing:-.02em;color:var(--text-primary);line-height:1.05;margin-bottom:4px">Clash Display</div>
      <div style="font-size:12px;color:var(--text-tertiary);font-family:var(--font-mono);margin-bottom:20px">Used for: Page titles, section headings, hero text, logo wordmark</div>
      <div style="font-family:var(--font-display);font-size:28px;font-weight:600;color:var(--text-secondary);margin-bottom:4px">Your Dashboard. Your Applications.</div>
      <div style="font-family:var(--font-display);font-size:20px;font-weight:500;color:var(--text-tertiary)">Heading Three Size · Subheadings</div>
    </div>

    <div class="card" style="margin-bottom:20px">
      <div style="font-family:var(--font-body);font-size:36px;font-weight:600;color:var(--text-primary);line-height:1.2;margin-bottom:4px">DM Sans</div>
      <div style="font-size:12px;color:var(--text-tertiary);font-family:var(--font-mono);margin-bottom:20px">Used for: Body text, labels, buttons, navigation, forms, everything readable</div>
      <div style="font-family:var(--font-body);font-size:16px;color:var(--text-primary);margin-bottom:8px">ApplyExpress searches Reed, LinkedIn, Indeed and more — scores each job against your profile, writes a tailored CV and cover letter, then applies automatically.</div>
      <div style="font-family:var(--font-body);font-size:13px;color:var(--text-secondary)">Labels · Navigation · Buttons · Form text · 13–14px is the primary reading size</div>
    </div>

    <div class="card">
      <div style="font-family:var(--font-mono);font-size:36px;font-weight:500;color:var(--accent);line-height:1.2;margin-bottom:4px">JetBrains Mono</div>
      <div style="font-size:12px;color:var(--text-tertiary);font-family:var(--font-mono);margin-bottom:20px">Used for: API keys, run IDs, timestamps, match scores, stat numbers</div>
      <div style="font-family:var(--font-mono);font-size:14px;color:var(--text-secondary);line-height:2">
        API Key: aa-Xk9p2mQ4vR7tNwBp3dJsLcFe<br>
        Run ID: 20260414_143022 · Score: 8.4/10<br>
        Applications: <span style="color:var(--success)">142 applied</span> · <span style="color:var(--error)">12 failed</span>
      </div>
    </div>
  </div>


  <!-- ════════════════════════════════════════════════════════════════ -->
  <!-- SECTION 5: Animated Placeholders Demo -->
  <!-- ════════════════════════════════════════════════════════════════ -->
  <div class="ds-section" id="placeholders">
    <div class="ds-section-label">05 · Placeholder Examples</div>
    <h2 style="margin-bottom:8px">Smart Placeholders</h2>
    <p style="color:var(--text-secondary);font-size:14px;margin-bottom:24px">Placeholders guide users without cluttering the UI. Watch the signup form fill itself in below.</p>

    <!-- Animated signup demo -->
    <div class="card" style="max-width:480px;margin-bottom:32px">
      <div style="font-size:13px;font-weight:600;color:var(--accent);margin-bottom:18px;display:flex;align-items:center;gap:8px">
        <span style="width:8px;height:8px;border-radius:50%;background:var(--accent);display:inline-block"></span>
        Live demo — watch the form fill itself
      </div>
      <div class="row" style="margin-bottom:14px">
        <div class="field">
          <label class="label">First name</label>
          <input type="text" id="demo-first" placeholder="e.g. Sarah">
        </div>
        <div class="field">
          <label class="label">Last name</label>
          <input type="text" id="demo-last" placeholder="e.g. Chen">
        </div>
      </div>
      <div class="field" style="margin-bottom:14px">
        <label class="label">Email address</label>
        <input type="email" id="demo-email" placeholder="your@email.com">
      </div>
      <div class="field" style="margin-bottom:14px">
        <label class="label">Job titles you're applying for</label>
        <input type="text" id="demo-jobs" placeholder="e.g. AML Compliance Analyst, MLRO, Risk Manager">
      </div>
      <div class="field" style="margin-bottom:14px">
        <label class="label">Search location</label>
        <input type="text" id="demo-loc" placeholder="London, UK">
      </div>
      <div style="display:flex;gap:10px;align-items:center;margin-bottom:18px">
        <input type="checkbox" id="demo-check" style="width:16px;height:16px">
        <label for="demo-check" style="font-size:12px;color:var(--text-tertiary);line-height:1.5">
          I authorise ApplyExpress to submit job applications on my behalf
        </label>
      </div>
      <button class="btn btn-primary btn-full" id="demo-btn" disabled style="opacity:.4">Start free trial →</button>
      <p style="font-size:11px;color:var(--text-tertiary);text-align:center;margin-top:8px">14 days free · No credit card required</p>
    </div>

    <div class="callout">
      <span class="callout-icon">✍️</span>
      <p>Good placeholders show a <em>real example</em> of what to type — not just the field name repeated. "AML Compliance Analyst, MLRO" is better than "Enter job titles". Users understand immediately without reading help text.</p>
    </div>
  </div>


  <!-- ════════════════════════════════════════════════════════════════ -->
  <!-- SECTION 6: Components Reference -->
  <!-- ════════════════════════════════════════════════════════════════ -->
  <div class="ds-section" id="components">
    <div class="ds-section-label">06 · Components</div>
    <h2 style="margin-bottom:8px">Every Building Block</h2>
    <p style="color:var(--text-secondary);font-size:14px;margin-bottom:24px">All components with a label explaining when each is used across the product.</p>

    <!-- Buttons -->
    <div class="card" style="margin-bottom:20px">
      <h3 style="margin-bottom:4px">Buttons</h3>
      <p style="font-size:12px;color:var(--text-tertiary);margin-bottom:16px">Primary = main action. Outline = secondary. Ghost = least important. Danger = destructive (delete, cancel).</p>
      <div style="display:flex;flex-wrap:wrap;gap:10px;align-items:center">
        <button class="btn btn-primary">Run pipeline</button>
        <button class="btn btn-success">✓ Applied</button>
        <button class="btn btn-outline">Download CV</button>
        <button class="btn btn-danger">Delete specialty</button>
        <button class="btn btn-ghost">← Back</button>
        <button class="btn btn-primary btn-sm">Save</button>
        <button class="btn btn-primary" disabled style="opacity:.35;cursor:not-allowed;transform:none">Processing…</button>
      </div>
    </div>

    <!-- Badges -->
    <div class="card" style="margin-bottom:20px">
      <h3 style="margin-bottom:4px">Status Badges</h3>
      <p style="font-size:12px;color:var(--text-tertiary);margin-bottom:16px">Show the current state of any run, application, or account at a glance.</p>
      <div style="display:flex;flex-wrap:wrap;gap:10px;align-items:center">
        <span class="badge badge-green">Applied</span>
        <span class="badge badge-green">Completed</span>
        <span class="badge badge-red">Failed</span>
        <span class="badge badge-red">Interrupted</span>
        <span class="badge badge-blue">Running</span>
        <span class="badge badge-blue">Active</span>
        <span class="badge badge-yellow">Trial — 8 days left</span>
        <span class="badge badge-gray">Pending</span>
      </div>
    </div>

    <!-- Match Scores -->
    <div class="card" style="margin-bottom:20px">
      <h3 style="margin-bottom:4px">Match Score Chips</h3>
      <p style="font-size:12px;color:var(--text-tertiary);margin-bottom:16px">AI scores each job 1–10. Green = strong match, amber = acceptable, red = below threshold (skipped).</p>
      <div style="display:flex;gap:28px;align-items:flex-end">
        <div style="text-align:center"><div class="score-chip score-high" style="font-size:36px;font-family:var(--font-mono)">9.1</div><div style="font-size:11px;color:var(--success);margin-top:6px;letter-spacing:.06em">STRONG FIT</div></div>
        <div style="text-align:center"><div class="score-chip score-mid" style="font-size:36px;font-family:var(--font-mono)">6.8</div><div style="font-size:11px;color:var(--warning);margin-top:6px;letter-spacing:.06em">GOOD FIT</div></div>
        <div style="text-align:center"><div class="score-chip score-low" style="font-size:36px;font-family:var(--font-mono)">3.4</div><div style="font-size:11px;color:var(--error);margin-top:6px;letter-spacing:.06em">BELOW THRESHOLD</div></div>
      </div>
    </div>

    <!-- Alerts -->
    <div class="card" style="margin-bottom:20px">
      <h3 style="margin-bottom:4px">Alert Banners</h3>
      <p style="font-size:12px;color:var(--text-tertiary);margin-bottom:16px">Used at the top of pages to communicate important state changes.</p>
      <div class="alert alert-success" style="margin-bottom:10px">✓ Pipeline completed — 4 CVs tailored and 3 applications submitted.</div>
      <div class="alert alert-error" style="margin-bottom:10px">✕ LinkedIn login failed — please update your credentials in Settings.</div>
      <div class="alert alert-info">· Your free trial ends 28 April 2026. Subscribe to keep the pipeline running.</div>
    </div>

    <!-- Stat Cards -->
    <div class="card" style="margin-bottom:20px">
      <h3 style="margin-bottom:4px">Stat Cards</h3>
      <p style="font-size:12px;color:var(--text-tertiary);margin-bottom:16px">Large numbers count up when they appear on screen (scroll down past them and back to see it again).</p>
      <div class="stat-row">
        <div class="stat"><div class="stat-num">247</div><div class="stat-label">Applications sent</div></div>
        <div class="stat"><div class="stat-num">63</div><div class="stat-label">Pipeline runs</div></div>
        <div class="stat"><div class="stat-num">14</div><div class="stat-label">Interviews</div></div>
      </div>
    </div>

    <!-- Toast demo -->
    <div class="card" style="margin-bottom:20px">
      <h3 style="margin-bottom:4px">Toast Notifications</h3>
      <p style="font-size:12px;color:var(--text-tertiary);margin-bottom:16px">Bottom-right pop-ups that appear after actions — auto-dismiss after 3.5s. Click to preview:</p>
      <div style="display:flex;flex-wrap:wrap;gap:10px">
        <button class="btn btn-outline btn-sm" onclick="showToast('3 applications submitted successfully','success')">✓ Success</button>
        <button class="btn btn-outline btn-sm" onclick="showToast('Connection failed — check server URL','error')">✕ Error</button>
        <button class="btn btn-outline btn-sm" onclick="showToast('Trial ends in 3 days — consider upgrading','warning')">⚠ Warning</button>
        <button class="btn btn-outline btn-sm" onclick="showToast('Extension connected to server','info')">· Info</button>
      </div>
    </div>

    <!-- Skeletons -->
    <div class="card" style="margin-bottom:20px">
      <h3 style="margin-bottom:4px">Loading Skeletons</h3>
      <p style="font-size:12px;color:var(--text-tertiary);margin-bottom:16px">Show the exact shape of content while it loads — no generic spinners on content areas.</p>
      <div style="display:flex;gap:14px;align-items:center;margin-bottom:14px">
        <div class="skeleton" style="width:44px;height:44px;border-radius:50%;flex-shrink:0"></div>
        <div style="flex:1"><div class="skeleton" style="height:14px;width:55%;margin-bottom:8px"></div><div class="skeleton" style="height:11px;width:35%"></div></div>
      </div>
      <div class="skeleton" style="height:12px;width:100%;margin-bottom:7px"></div>
      <div class="skeleton" style="height:12px;width:82%;margin-bottom:7px"></div>
      <div class="skeleton" style="height:12px;width:65%"></div>
    </div>

  </div><!-- /components -->

  <!-- Bottom nav -->
  <div style="text-align:center;padding:24px 0 48px;border-top:1px solid var(--border)">
    <p style="color:var(--text-tertiary);font-size:13px;margin-bottom:16px">Design system reviewed. Ready to apply to all pages.</p>
    <a href="/dashboard" class="btn btn-outline">← Dashboard</a>
    &nbsp;
    <a href="/" class="btn btn-outline">Landing page →</a>
  </div>

</div><!-- /container-lg -->

<script>
/* ── Logo selector ───────────────── */
function selectLogo(n) {
  [1,2,3].forEach(function(i) {
    var c = document.getElementById('logo-opt-'+i);
    c.classList.toggle('selected', i===n);
  });
  var msg = document.getElementById('logo-selected-msg');
  msg.style.display = 'flex';
  msg.style.alignItems = 'center';
  msg.style.gap = '10px';
  msg.textContent = '✓ Logo Option ' + ['A','B','C'][n-1] + ' selected — "' +
    ['"Land your next role on autopilot."',
     '"Your unfair advantage in every job market."',
     '"From search to offer. Completely automated."'][n-1];
}

/* ── Pipeline animation ──────────── */
var pipeSteps = [0,1,2,3,4,5];
var pipeStatus = ['Uploading your CV…', 'Searching Reed, LinkedIn, Indeed…',
  'AI scoring each job against your profile…', 'Tailoring CV and cover letter…',
  'Extension submitting applications…', '🎯 Done! Check your dashboard.'];
var pipeTimer, pipeIdx = 0;

function runPipeline() {
  pipeTimer = setInterval(function() {
    var steps = document.querySelectorAll('.pipe-step');
    var bar   = document.getElementById('pipeline-bar');
    var label = document.getElementById('pipeline-status');
    steps.forEach(function(s, i) {
      s.classList.toggle('active', i === pipeIdx);
      s.classList.toggle('done',   i < pipeIdx);
    });
    if (bar) bar.style.width = Math.round((pipeIdx / (pipeSteps.length - 1)) * 100) + '%';
    if (label) label.textContent = pipeStatus[pipeIdx] || '';
    pipeIdx++;
    if (pipeIdx > pipeSteps.length) {
      clearInterval(pipeTimer);
      pipeIdx = pipeSteps.length;
    }
  }, 1200);
}

function restartPipeline() {
  clearInterval(pipeTimer);
  pipeIdx = 0;
  var steps = document.querySelectorAll('.pipe-step');
  steps.forEach(function(s) { s.classList.remove('active','done'); });
  var bar = document.getElementById('pipeline-bar');
  if (bar) bar.style.width = '0%';
  setTimeout(runPipeline, 200);
}

/* Start on scroll into view */
var pipeObserver = new IntersectionObserver(function(entries) {
  if (entries[0].isIntersecting) { runPipeline(); pipeObserver.disconnect(); }
}, {threshold: 0.3});
var pipeEl = document.getElementById('pipeline-demo');
if (pipeEl) pipeObserver.observe(pipeEl);

/* ── Animated signup demo ────────── */
var demoData = [
  {id:'demo-first', val:'Sarah'},
  {id:'demo-last',  val:'Chen'},
  {id:'demo-email', val:'sarah.chen@gmail.com'},
  {id:'demo-jobs',  val:'AML Compliance Analyst, MLRO, Risk Manager'},
  {id:'demo-loc',   val:'London, UK'},
];
var demoIdx = 0, demoCharIdx = 0, demoTimer;

function typeNextChar() {
  if (demoIdx >= demoData.length) {
    var chk = document.getElementById('demo-check');
    var btn = document.getElementById('demo-btn');
    if (chk) { chk.checked = true; }
    if (btn) { btn.disabled = false; btn.style.opacity = '1'; }
    return;
  }
  var item = demoData[demoIdx];
  var el   = document.getElementById(item.id);
  if (!el) { demoIdx++; typeNextChar(); return; }
  if (demoCharIdx === 0 && demoIdx > 0) {
    /* brief pause between fields */
    setTimeout(typeNextChar, 300);
    demoCharIdx = -1; return;
  }
  if (demoCharIdx < 0) demoCharIdx = 0;
  if (demoCharIdx < item.val.length) {
    el.value = item.val.slice(0, demoCharIdx + 1);
    demoCharIdx++;
    demoTimer = setTimeout(typeNextChar, 55 + Math.random() * 40);
  } else {
    demoIdx++;
    demoCharIdx = 0;
    setTimeout(typeNextChar, 500);
  }
}

/* Start demo when it scrolls into view */
var demoObserver = new IntersectionObserver(function(entries) {
  if (entries[0].isIntersecting) { setTimeout(typeNextChar, 800); demoObserver.disconnect(); }
}, {threshold: 0.4});
var demoEl = document.getElementById('demo-first');
if (demoEl) demoObserver.observe(demoEl);
</script>"""
    return _page("Brand + Design System", body)


# ── Dashboard Preview (sample only — real /dashboard unchanged) ───────────────
@app.route("/dashboard-preview")
def dashboard_preview():
    body = """
<style>
/* ── Page-level overrides for preview ── */
.preview-nav{
  position:fixed;top:0;left:0;right:0;z-index:200;
  background:rgba(8,12,18,.92);backdrop-filter:blur(12px);
  border-bottom:1px solid #1e2d42;
  display:flex;align-items:center;justify-content:space-between;
  padding:0 32px;height:56px;
}
.preview-nav-brand{
  display:flex;align-items:center;gap:10px;
  font-family:var(--font-display);font-size:18px;font-weight:700;
  color:var(--text-primary);text-decoration:none;letter-spacing:-.01em;
}
.preview-nav-links{display:flex;gap:4px;align-items:center}
.preview-nav-link{
  padding:6px 14px;border-radius:8px;font-size:13px;font-weight:500;
  color:var(--text-secondary);text-decoration:none;transition:all .15s;
}
.preview-nav-link:hover{color:var(--text-primary);background:var(--bg-elevated)}
.preview-nav-link.active{color:var(--accent);background:var(--accent-surface)}
.preview-nav-right{display:flex;align-items:center;gap:10px}
.preview-avatar{
  width:32px;height:32px;border-radius:50%;
  background:linear-gradient(135deg,var(--accent),var(--accent-dim));
  display:flex;align-items:center;justify-content:center;
  font-size:12px;font-weight:700;color:#080c12;
  font-family:var(--font-body);cursor:pointer;
}

/* ── Dashboard layout ── */
.dash-wrap{
  max-width:1100px;margin:0 auto;padding:88px 24px 60px;
  animation:fade-up .4s cubic-bezier(.16,1,.3,1) both;
}
@keyframes fade-up{from{opacity:0;transform:translateY(16px)}to{opacity:1;transform:none}}

/* Hero row */
.dash-hero{
  display:flex;align-items:flex-start;justify-content:space-between;
  flex-wrap:wrap;gap:20px;margin-bottom:32px;
}
.dash-greeting{font-family:var(--font-display);font-size:32px;font-weight:700;
  letter-spacing:-.01em;color:var(--text-primary);margin:0 0 6px}
.dash-subline{font-size:13px;color:var(--text-secondary);display:flex;
  align-items:center;gap:8px;margin:0}
.dash-status-dot{width:7px;height:7px;border-radius:50%;background:var(--success);
  box-shadow:0 0 0 2px rgba(16,214,138,.25);animation:pulse-dot 2.5s ease infinite}
@keyframes pulse-dot{0%,100%{box-shadow:0 0 0 2px rgba(16,214,138,.25)}
  50%{box-shadow:0 0 0 5px rgba(16,214,138,.08)}}

/* Run form */
.run-form{display:flex;gap:10px;align-items:center;flex-wrap:wrap}
.run-select{
  padding:10px 16px;background:var(--bg-surface);border:1px solid var(--border);
  border-radius:10px;color:var(--text-primary);font-family:var(--font-body);
  font-size:13px;cursor:pointer;transition:border-color .15s;
}
.run-select:hover,.run-select:focus{border-color:var(--border-bright);outline:none}
.btn-run{
  display:inline-flex;align-items:center;gap:8px;
  padding:11px 22px;background:var(--accent);color:#080c12;
  border:none;border-radius:10px;font-family:var(--font-body);font-size:14px;
  font-weight:700;cursor:pointer;transition:all .12s ease;
}
.btn-run:hover{background:var(--accent-dim);transform:scale(1.02)}
.btn-run:active{transform:scale(.97)}
.run-icon{font-size:15px;line-height:1}

/* Stat cards */
.stat-grid{
  display:grid;grid-template-columns:repeat(3,1fr);gap:16px;
  margin-bottom:28px;
}
.stat-card{
  background:var(--bg-surface);border:1px solid var(--border);
  border-radius:14px;padding:24px 24px 20px;transition:all .15s;
  position:relative;overflow:hidden;
}
.stat-card::before{
  content:'';position:absolute;inset:0;
  background:linear-gradient(135deg,var(--accent-surface),transparent 60%);
  opacity:0;transition:opacity .2s;
}
.stat-card:hover{border-color:var(--border-bright);transform:translateY(-2px)}
.stat-card:hover::before{opacity:1}
.stat-icon{font-size:18px;margin-bottom:14px;display:block}
.stat-value{
  font-family:var(--font-display);font-size:42px;font-weight:700;
  color:var(--text-primary);letter-spacing:-.02em;line-height:1;margin-bottom:6px;
}
.stat-label{font-size:12px;color:var(--text-tertiary);text-transform:uppercase;
  letter-spacing:.06em;font-weight:500}
.stat-change{font-size:12px;color:var(--success);margin-top:10px;font-weight:600}

/* Section header */
.section-hd{
  display:flex;align-items:center;justify-content:space-between;
  margin-bottom:16px;
}
.section-title{font-family:var(--font-display);font-size:16px;font-weight:600;
  color:var(--text-primary);letter-spacing:-.01em}
.section-link{font-size:12px;color:var(--accent);text-decoration:none;
  font-weight:600;letter-spacing:.02em}
.section-link:hover{color:var(--accent-dim)}

/* Runs table */
.runs-card{
  background:var(--bg-surface);border:1px solid var(--border);
  border-radius:14px;overflow:hidden;margin-bottom:24px;
}
.runs-table{width:100%;border-collapse:collapse}
.runs-table th{
  padding:13px 20px;background:var(--bg-elevated);
  font-size:11px;font-weight:600;color:var(--text-tertiary);
  text-transform:uppercase;letter-spacing:.08em;
  text-align:left;border-bottom:1px solid var(--border);
}
.runs-table td{
  padding:14px 20px;border-bottom:1px solid rgba(30,45,66,.6);
  font-size:13px;color:var(--text-primary);
}
.runs-table tr:last-child td{border-bottom:none}
.runs-table tr:hover td{background:var(--bg-elevated)}
.run-date{font-family:var(--font-mono);font-size:12px;color:var(--text-secondary)}
.run-score-bar{
  display:flex;align-items:center;gap:10px;
}
.score-bar-track{
  flex:1;max-width:80px;height:4px;background:var(--bg-overlay);
  border-radius:2px;overflow:hidden;
}
.score-bar-fill{height:100%;border-radius:2px;
  background:linear-gradient(90deg,var(--accent-dim),var(--accent));
  transition:width .8s cubic-bezier(.16,1,.3,1)}
.score-label{font-family:var(--font-mono);font-size:12px;color:var(--accent);
  font-weight:600;min-width:26px}
.view-btn{
  font-size:12px;color:var(--accent);text-decoration:none;
  font-weight:600;opacity:.8;transition:opacity .15s;
}
.view-btn:hover{opacity:1}

/* Extension card */
.ext-card{
  background:var(--bg-surface);border:1px solid var(--border);
  border-radius:14px;padding:32px;
  display:flex;align-items:center;gap:32px;flex-wrap:wrap;
}
.ext-icon-wrap{
  width:64px;height:64px;border-radius:16px;
  background:var(--accent-surface);border:1px solid rgba(0,200,255,.15);
  display:flex;align-items:center;justify-content:center;
  font-size:28px;flex-shrink:0;
}
.ext-body{flex:1;min-width:200px}
.ext-title{font-family:var(--font-display);font-size:18px;font-weight:700;
  color:var(--text-primary);margin:0 0 8px}
.ext-desc{font-size:13px;color:var(--text-secondary);margin:0 0 20px;
  max-width:480px;line-height:1.6}
.btn-connect{
  display:inline-flex;align-items:center;gap:8px;
  padding:11px 24px;background:var(--bg-elevated);
  border:1px solid var(--border-bright);border-radius:10px;
  color:var(--text-primary);font-family:var(--font-body);font-size:14px;
  font-weight:600;cursor:pointer;text-decoration:none;transition:all .15s;
}
.btn-connect:hover{border-color:var(--accent);color:var(--accent)}

/* Logo B inline */
.logo-b{display:flex;align-items:center;justify-content:center;
  width:28px;height:28px}

/* Preview banner */
.preview-banner{
  position:fixed;bottom:0;left:0;right:0;z-index:300;
  background:#0e1520;border-top:1px solid var(--border-bright);
  padding:12px 24px;display:flex;align-items:center;justify-content:center;
  gap:16px;
}
.preview-banner-text{font-size:13px;color:var(--text-secondary)}
.preview-banner-text strong{color:var(--accent)}

@media(max-width:680px){
  .stat-grid{grid-template-columns:1fr}
  .dash-hero{flex-direction:column}
  .ext-card{flex-direction:column;gap:20px}
}
</style>

<!-- ── Nav ──────────────────────────────────────────────────────────────── -->
<nav class="preview-nav">
  <a href="/dashboard-preview" class="preview-nav-brand">
    <!-- Logo B — The Target -->
    <svg class="logo-b" viewBox="0 0 28 28" fill="none" xmlns="http://www.w3.org/2000/svg">
      <defs>
        <radialGradient id="lg-b-glow" cx="50%" cy="50%" r="50%">
          <stop offset="0%" stop-color="#00c8ff" stop-opacity=".25"/>
          <stop offset="100%" stop-color="#00c8ff" stop-opacity="0"/>
        </radialGradient>
      </defs>
      <!-- Outer glow -->
      <circle cx="14" cy="14" r="13" fill="url(#lg-b-glow)"/>
      <!-- Ring 3 -->
      <circle cx="14" cy="14" r="12" stroke="#1e2d42" stroke-width="1.2"/>
      <!-- Ring 2 — pulsing -->
      <circle cx="14" cy="14" r="8" stroke="#00c8ff" stroke-width="1.4" stroke-opacity=".5">
        <animate attributeName="r" values="8;9.5;8" dur="3s" repeatCount="indefinite"/>
        <animate attributeName="stroke-opacity" values=".5;.15;.5" dur="3s" repeatCount="indefinite"/>
      </circle>
      <!-- Ring 1 — inner pulsing -->
      <circle cx="14" cy="14" r="4.5" stroke="#10d68a" stroke-width="1.6">
        <animate attributeName="r" values="4.5;5.5;4.5" dur="2.4s" repeatCount="indefinite"/>
        <animate attributeName="stroke-opacity" values="1;.4;1" dur="2.4s" repeatCount="indefinite"/>
      </circle>
      <!-- Centre dot -->
      <circle cx="14" cy="14" r="2" fill="#00c8ff">
        <animate attributeName="opacity" values="1;.5;1" dur="1.8s" repeatCount="indefinite"/>
      </circle>
    </svg>
    ApplyExpress
  </a>

  <div class="preview-nav-links">
    <a class="preview-nav-link active" href="#">Dashboard</a>
    <a class="preview-nav-link" href="#">Runs</a>
    <a class="preview-nav-link" href="#">History</a>
    <a class="preview-nav-link" href="#">Settings</a>
  </div>

  <div class="preview-nav-right">
    <span style="font-size:12px;color:var(--text-tertiary);font-family:var(--font-mono)">14-day trial · 11 days left</span>
    <div class="preview-avatar">SC</div>
  </div>
</nav>

<!-- ── Dashboard body ────────────────────────────────────────────────────── -->
<div class="dash-wrap">

  <!-- Hero row -->
  <div class="dash-hero">
    <div>
      <h1 class="dash-greeting">Hi, Sarah &#x1F44B;</h1>
      <p class="dash-subline">
        <span class="dash-status-dot"></span>
        <span class="badge badge-green" style="font-size:12px">CV uploaded</span>
        &nbsp;
        <span style="color:var(--text-tertiary)">Last run: 2 hours ago</span>
      </p>
    </div>
    <div class="run-form">
      <select class="run-select">
        <option>— Default CV —</option>
        <option>AML Compliance</option>
        <option>Risk Manager</option>
      </select>
      <button class="btn-run" onclick="window.showToast('Pipeline started — the extension will begin applying jobs shortly.','success')">
        <span class="run-icon">&#9654;</span> Run now
      </button>
    </div>
  </div>

  <!-- Stat cards -->
  <div class="stat-grid">
    <div class="stat-card">
      <span class="stat-icon">&#127919;</span>
      <div class="stat-value" data-target="142" data-count>0</div>
      <div class="stat-label">Applications sent</div>
      <div class="stat-change">&#8593; 23 this week</div>
    </div>
    <div class="stat-card">
      <span class="stat-icon">&#128269;</span>
      <div class="stat-value" data-target="38" data-count>0</div>
      <div class="stat-label">Pipeline runs</div>
      <div class="stat-change">&#8593; 4 this week</div>
    </div>
    <div class="stat-card">
      <span class="stat-icon">&#9889;</span>
      <div class="stat-value" data-target="9" data-count>0</div>
      <div class="stat-label">Interviews</div>
      <div class="stat-change" style="color:var(--accent)">&#8593; 2 this month</div>
    </div>
  </div>

  <!-- Recent runs -->
  <div class="section-hd">
    <span class="section-title">Recent runs</span>
    <a class="section-link" href="#">View all &rarr;</a>
  </div>
  <div class="runs-card">
    <table class="runs-table">
      <thead><tr>
        <th>Date &amp; time</th>
        <th>Status</th>
        <th>Found</th>
        <th>Applied</th>
        <th>Avg score</th>
        <th></th>
      </tr></thead>
      <tbody>
        <tr>
          <td class="run-date">2026-04-14 &nbsp;09:12</td>
          <td><span class="badge badge-green">Completed</span></td>
          <td style="color:var(--text-secondary)">24</td>
          <td style="color:var(--text-primary);font-weight:600">18</td>
          <td>
            <div class="run-score-bar">
              <div class="score-bar-track"><div class="score-bar-fill" style="width:87%"></div></div>
              <span class="score-label">8.7</span>
            </div>
          </td>
          <td><a href="#" class="view-btn">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="run-date">2026-04-13 &nbsp;21:44</td>
          <td><span class="badge badge-green">Completed</span></td>
          <td style="color:var(--text-secondary)">31</td>
          <td style="color:var(--text-primary);font-weight:600">22</td>
          <td>
            <div class="run-score-bar">
              <div class="score-bar-track"><div class="score-bar-fill" style="width:74%"></div></div>
              <span class="score-label">7.4</span>
            </div>
          </td>
          <td><a href="#" class="view-btn">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="run-date">2026-04-12 &nbsp;14:08</td>
          <td><span class="badge badge-red">Failed</span></td>
          <td style="color:var(--text-secondary)">7</td>
          <td style="color:var(--text-primary);font-weight:600">0</td>
          <td>
            <div class="run-score-bar">
              <div class="score-bar-track"><div class="score-bar-fill" style="width:0%"></div></div>
              <span class="score-label" style="color:var(--error)">—</span>
            </div>
          </td>
          <td><a href="#" class="view-btn">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="run-date">2026-04-11 &nbsp;09:00</td>
          <td><span class="badge badge-green">Completed</span></td>
          <td style="color:var(--text-secondary)">19</td>
          <td style="color:var(--text-primary);font-weight:600">14</td>
          <td>
            <div class="run-score-bar">
              <div class="score-bar-track"><div class="score-bar-fill" style="width:81%"></div></div>
              <span class="score-label">8.1</span>
            </div>
          </td>
          <td><a href="#" class="view-btn">View &rarr;</a></td>
        </tr>
      </tbody>
    </table>
  </div>

  <!-- Extension connect card -->
  <div class="ext-card">
    <div class="ext-icon-wrap">&#128279;</div>
    <div class="ext-body">
      <h2 class="ext-title">Connect your Chrome extension</h2>
      <p class="ext-desc">
        One click to sync the extension with your account. No copying API keys.
        Once connected, press <strong style="color:var(--text-primary)">Start</strong>
        in the popup to begin submitting applications automatically.
      </p>
      <a href="#" class="btn-connect"
         onclick="window.showToast('Connection link copied — open in Chrome to activate.','info');return false">
        Connect extension &rarr;
      </a>
    </div>
    <div style="text-align:right;flex-shrink:0">
      <div style="font-size:11px;color:var(--text-tertiary);text-transform:uppercase;
        letter-spacing:.06em;margin-bottom:8px">Extension status</div>
      <span class="badge badge-green" style="font-size:12px">&#8226; Connected</span>
    </div>
  </div>

</div>

<!-- ── Preview banner ────────────────────────────────────────────────────── -->
<div class="preview-banner">
  <span class="preview-banner-text">
    &#128065; This is a <strong>preview only</strong> — your live dashboard is unchanged at
    <a href="/dashboard" style="color:var(--accent);text-decoration:none">/dashboard</a>
  </span>
  <a href="/design-system" style="font-size:12px;color:var(--text-tertiary);text-decoration:none">
    &larr; Back to design system
  </a>
</div>

<script>
/* Count-up animation on stat cards */
(function(){
  var counters = document.querySelectorAll('[data-count]');
  var started = false;
  function runCounters(){
    if(started) return; started = true;
    counters.forEach(function(el){
      var target = parseInt(el.dataset.target, 10);
      var duration = 900;
      var steps = 40;
      var step = 0;
      var iv = setInterval(function(){
        step++;
        var pct = step / steps;
        var eased = 1 - Math.pow(1 - pct, 3);
        el.textContent = Math.round(eased * target);
        if(step >= steps){ el.textContent = target; clearInterval(iv); }
      }, duration / steps);
    });
  }
  if('IntersectionObserver' in window){
    var obs = new IntersectionObserver(function(entries){
      if(entries[0].isIntersecting) runCounters();
    },{threshold:.3});
    var grid = document.querySelector('.stat-grid');
    if(grid) obs.observe(grid);
  } else { setTimeout(runCounters, 400); }

  /* prefers-reduced-motion */
  if(window.matchMedia('(prefers-reduced-motion: reduce)').matches){
    counters.forEach(function(el){ el.textContent = el.dataset.target; });
  }
})();
</script>"""

    return _page("Dashboard — Preview", body)


# ── Shared dashboard preview template ────────────────────────────────────────
def _dash_preview_html(scheme_name, tokens, logo_svg, nav_links_html, banner_label, banner_bg, banner_border, banner_text_color, banner_accent_color):
    """Renders a dashboard preview with a given colour scheme injected as CSS vars."""
    return f"""
<style>
:root{{
  --bg-base:{tokens['bg_base']};
  --bg-surface:{tokens['bg_surface']};
  --bg-elevated:{tokens['bg_elevated']};
  --bg-overlay:{tokens['bg_overlay']};
  --accent:{tokens['accent']};
  --accent-dim:{tokens['accent_dim']};
  --accent-glow:{tokens['accent_glow']};
  --accent-surface:{tokens['accent_surface']};
  --text-primary:{tokens['text_primary']};
  --text-secondary:{tokens['text_secondary']};
  --text-tertiary:{tokens['text_tertiary']};
  --border:{tokens['border']};
  --border-bright:{tokens['border_bright']};
  --success:{tokens.get('success','#10d68a')};
  --warning:{tokens.get('warning','#f5a623')};
  --error:{tokens.get('error','#ff4d6a')};
}}
body{{background:var(--bg-base);color:var(--text-primary);
  font-family:var(--font-body);margin:0;padding:0}}

.pv-nav{{
  position:fixed;top:0;left:0;right:0;z-index:200;
  background:{tokens['nav_bg']};backdrop-filter:blur(14px);-webkit-backdrop-filter:blur(14px);
  border-bottom:1px solid var(--border);
  display:flex;align-items:center;justify-content:space-between;
  padding:0 32px;height:56px;
}}
.pv-brand{{display:flex;align-items:center;gap:10px;
  font-family:var(--font-display);font-size:18px;font-weight:700;
  color:var(--text-primary);text-decoration:none;letter-spacing:-.01em}}
.pv-links{{display:flex;gap:4px}}
.pv-link{{padding:6px 14px;border-radius:8px;font-size:13px;font-weight:500;
  color:var(--text-secondary);text-decoration:none;transition:all .15s}}
.pv-link:hover{{color:var(--text-primary);background:var(--bg-elevated)}}
.pv-link.on{{color:var(--accent);background:var(--accent-surface)}}
.pv-right{{display:flex;align-items:center;gap:10px}}
.pv-trial{{font-size:12px;color:var(--text-tertiary);font-family:var(--font-mono)}}
.pv-avatar{{width:32px;height:32px;border-radius:50%;
  background:linear-gradient(135deg,var(--accent),var(--accent-dim));
  display:flex;align-items:center;justify-content:center;
  font-size:12px;font-weight:700;color:{tokens['avatar_text']};
  font-family:var(--font-body)}}

.dw{{max-width:1100px;margin:0 auto;padding:80px 24px 80px;
  animation:fu .4s cubic-bezier(.16,1,.3,1) both}}
@keyframes fu{{from{{opacity:0;transform:translateY(14px)}}to{{opacity:1;transform:none}}}}

.dh{{display:flex;align-items:flex-start;justify-content:space-between;
  flex-wrap:wrap;gap:20px;margin-bottom:32px}}
.dg{{font-family:var(--font-display);font-size:32px;font-weight:700;
  letter-spacing:-.01em;color:var(--text-primary);margin:0 0 8px}}
.ds{{font-size:13px;color:var(--text-secondary);display:flex;align-items:center;gap:8px;margin:0}}
.dot{{width:7px;height:7px;border-radius:50%;background:var(--success);
  animation:pd 2.5s ease infinite}}
@keyframes pd{{0%,100%{{box-shadow:0 0 0 2px {tokens['success_glow']}}}
  50%{{box-shadow:0 0 0 5px {tokens['success_glow_far']}}}}}

.rf{{display:flex;gap:10px;align-items:center;flex-wrap:wrap}}
.rs{{padding:10px 16px;background:var(--bg-surface);border:1px solid var(--border);
  border-radius:10px;color:var(--text-primary);font-family:var(--font-body);
  font-size:13px;cursor:pointer;transition:border-color .15s}}
.rs:hover,.rs:focus{{border-color:var(--border-bright);outline:none}}
.rbtn{{display:inline-flex;align-items:center;gap:8px;
  padding:11px 22px;background:var(--accent);color:{tokens['btn_text']};
  border:none;border-radius:10px;font-family:var(--font-body);font-size:14px;
  font-weight:700;cursor:pointer;transition:all .12s ease;letter-spacing:.01em}}
.rbtn:hover{{background:var(--accent-dim);transform:scale(1.02)}}
.rbtn:active{{transform:scale(.97)}}

.sg{{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-bottom:28px}}
.sc{{background:var(--bg-surface);border:1px solid var(--border);
  border-radius:14px;padding:24px;transition:all .15s;position:relative;overflow:hidden}}
.sc::before{{content:'';position:absolute;inset:0;
  background:linear-gradient(135deg,var(--accent-surface),transparent 60%);
  opacity:0;transition:opacity .2s}}
.sc:hover{{border-color:var(--border-bright);transform:translateY(-2px)}}
.sc:hover::before{{opacity:1}}
.si{{font-size:18px;margin-bottom:14px;display:block}}
.sv{{font-family:var(--font-display);font-size:42px;font-weight:700;
  color:var(--text-primary);letter-spacing:-.02em;line-height:1;margin-bottom:6px}}
.sl{{font-size:11px;color:var(--text-tertiary);text-transform:uppercase;
  letter-spacing:.07em;font-weight:500}}
.sc-{{font-size:12px;color:var(--success);margin-top:10px;font-weight:600}}

.sh{{display:flex;align-items:center;justify-content:space-between;margin-bottom:14px}}
.st{{font-family:var(--font-display);font-size:16px;font-weight:600;
  color:var(--text-primary);letter-spacing:-.01em}}
.sa{{font-size:12px;color:var(--accent);text-decoration:none;font-weight:600}}
.sa:hover{{color:var(--accent-dim)}}

.rc{{background:var(--bg-surface);border:1px solid var(--border);
  border-radius:14px;overflow:hidden;margin-bottom:24px}}
.rt{{width:100%;border-collapse:collapse}}
.rt th{{padding:13px 20px;background:var(--bg-elevated);
  font-size:11px;font-weight:600;color:var(--text-tertiary);
  text-transform:uppercase;letter-spacing:.08em;text-align:left;
  border-bottom:1px solid var(--border)}}
.rt td{{padding:14px 20px;border-bottom:1px solid {tokens['row_border']};
  font-size:13px;color:var(--text-primary)}}
.rt tr:last-child td{{border-bottom:none}}
.rt tr:hover td{{background:var(--bg-elevated)}}
.rd{{font-family:var(--font-mono);font-size:12px;color:var(--text-secondary)}}
.sb-wrap{{display:flex;align-items:center;gap:10px}}
.sb-track{{flex:1;max-width:80px;height:4px;background:var(--bg-overlay);
  border-radius:2px;overflow:hidden}}
.sb-fill{{height:100%;border-radius:2px;
  background:linear-gradient(90deg,var(--accent-dim),var(--accent));
  transition:width .8s cubic-bezier(.16,1,.3,1)}}
.sb-val{{font-family:var(--font-mono);font-size:12px;color:var(--accent);
  font-weight:600;min-width:26px}}
.vb{{font-size:12px;color:var(--accent);text-decoration:none;
  font-weight:600;opacity:.8;transition:opacity .15s}}
.vb:hover{{opacity:1}}

.ec{{background:var(--bg-surface);border:1px solid var(--border);
  border-radius:14px;padding:32px;display:flex;align-items:center;
  gap:32px;flex-wrap:wrap}}
.ei{{width:64px;height:64px;border-radius:16px;background:var(--accent-surface);
  border:1px solid {tokens['ext_icon_border']};display:flex;align-items:center;
  justify-content:center;font-size:28px;flex-shrink:0}}
.eb{{flex:1;min-width:200px}}
.et{{font-family:var(--font-display);font-size:18px;font-weight:700;
  color:var(--text-primary);margin:0 0 8px}}
.ed{{font-size:13px;color:var(--text-secondary);margin:0 0 20px;
  max-width:480px;line-height:1.6}}
.cbtn{{display:inline-flex;align-items:center;gap:8px;
  padding:11px 24px;background:var(--bg-elevated);
  border:1px solid var(--border-bright);border-radius:10px;
  color:var(--text-primary);font-family:var(--font-body);font-size:14px;
  font-weight:600;cursor:pointer;text-decoration:none;transition:all .15s}}
.cbtn:hover{{border-color:var(--accent);color:var(--accent)}}

/* scheme switcher banner */
.scheme-bar{{
  position:fixed;bottom:0;left:0;right:0;z-index:300;
  background:{banner_bg};border-top:1px solid {banner_border};
  padding:12px 24px;display:flex;align-items:center;
  justify-content:center;gap:24px;flex-wrap:wrap;
}}
.scheme-label{{font-size:13px;color:{banner_text_color};
  font-weight:700;letter-spacing:.02em}}
.scheme-links{{display:flex;gap:8px;align-items:center}}
.scheme-link{{padding:6px 14px;border-radius:8px;font-size:12px;
  font-weight:600;text-decoration:none;border:1px solid {banner_border};
  color:{banner_text_color};transition:all .15s}}
.scheme-link:hover,.scheme-link.current{{
  background:{banner_accent_color};color:{tokens['btn_text']};border-color:{banner_accent_color}}}
.scheme-back{{font-size:12px;color:{banner_text_color};opacity:.5;
  text-decoration:none;margin-left:16px}}
.scheme-back:hover{{opacity:1}}

@media(max-width:680px){{
  .sg{{grid-template-columns:1fr}}
  .dh{{flex-direction:column}}
  .ec{{flex-direction:column;gap:20px}}
}}
</style>

<!-- Nav -->
<nav class="pv-nav">
  <a href="#" class="pv-brand">
    {logo_svg}
    ApplyExpress
  </a>
  <div class="pv-links">
    <a class="pv-link on" href="#">Dashboard</a>
    <a class="pv-link" href="#">Runs</a>
    <a class="pv-link" href="#">History</a>
    <a class="pv-link" href="#">Settings</a>
  </div>
  <div class="pv-right">
    <span class="pv-trial">14-day trial &middot; 11 days left</span>
    <div class="pv-avatar">SC</div>
  </div>
</nav>

<div class="dw">
  <!-- Hero -->
  <div class="dh">
    <div>
      <h1 class="dg">Hi, Sarah &#x1F44B;</h1>
      <p class="ds">
        <span class="dot"></span>
        <span class="badge badge-green" style="font-size:12px">CV uploaded</span>
        &nbsp;
        <span style="color:var(--text-tertiary)">Last run: 2 hours ago</span>
      </p>
    </div>
    <div class="rf">
      <select class="rs">
        <option>&#8212; Default CV &#8212;</option>
        <option>AML Compliance</option>
        <option>Risk Manager</option>
      </select>
      <button class="rbtn"
        onclick="window.showToast('Pipeline started &mdash; the extension will begin applying shortly.','success')">
        &#9654;&nbsp; Run now
      </button>
    </div>
  </div>

  <!-- Stats -->
  <div class="sg">
    <div class="sc">
      <span class="si">&#127919;</span>
      <div class="sv" data-target="142" data-count>0</div>
      <div class="sl">Applications sent</div>
      <div class="sc-">&#8593; 23 this week</div>
    </div>
    <div class="sc">
      <span class="si">&#128269;</span>
      <div class="sv" data-target="38" data-count>0</div>
      <div class="sl">Pipeline runs</div>
      <div class="sc-">&#8593; 4 this week</div>
    </div>
    <div class="sc">
      <span class="si">&#9889;</span>
      <div class="sv" data-target="9" data-count>0</div>
      <div class="sl">Interviews</div>
      <div class="sc-" style="color:var(--accent)">&#8593; 2 this month</div>
    </div>
  </div>

  <!-- Runs table -->
  <div class="sh">
    <span class="st">Recent runs</span>
    <a class="sa" href="#">View all &rarr;</a>
  </div>
  <div class="rc">
    <table class="rt">
      <thead><tr>
        <th>Date &amp; time</th><th>Status</th>
        <th>Found</th><th>Applied</th><th>Avg score</th><th></th>
      </tr></thead>
      <tbody>
        <tr>
          <td class="rd">2026-04-14 &nbsp;09:12</td>
          <td><span class="badge badge-green">Completed</span></td>
          <td style="color:var(--text-secondary)">24</td>
          <td style="font-weight:600">18</td>
          <td><div class="sb-wrap"><div class="sb-track"><div class="sb-fill" style="width:87%"></div></div><span class="sb-val">8.7</span></div></td>
          <td><a href="#" class="vb">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="rd">2026-04-13 &nbsp;21:44</td>
          <td><span class="badge badge-green">Completed</span></td>
          <td style="color:var(--text-secondary)">31</td>
          <td style="font-weight:600">22</td>
          <td><div class="sb-wrap"><div class="sb-track"><div class="sb-fill" style="width:74%"></div></div><span class="sb-val">7.4</span></div></td>
          <td><a href="#" class="vb">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="rd">2026-04-12 &nbsp;14:08</td>
          <td><span class="badge badge-red">Failed</span></td>
          <td style="color:var(--text-secondary)">7</td>
          <td style="font-weight:600">0</td>
          <td><div class="sb-wrap"><div class="sb-track"><div class="sb-fill" style="width:0%"></div></div><span class="sb-val" style="color:var(--error)">&#8212;</span></div></td>
          <td><a href="#" class="vb">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="rd">2026-04-11 &nbsp;09:00</td>
          <td><span class="badge badge-green">Completed</span></td>
          <td style="color:var(--text-secondary)">19</td>
          <td style="font-weight:600">14</td>
          <td><div class="sb-wrap"><div class="sb-track"><div class="sb-fill" style="width:81%"></div></div><span class="sb-val">8.1</span></div></td>
          <td><a href="#" class="vb">View &rarr;</a></td>
        </tr>
      </tbody>
    </table>
  </div>

  <!-- Extension card -->
  <div class="ec">
    <div class="ei">&#128279;</div>
    <div class="eb">
      <h2 class="et">Connect your Chrome extension</h2>
      <p class="ed">One click to sync the extension with your account. No copying API keys.
        Once connected, press <strong style="color:var(--text-primary)">Start</strong>
        in the popup to begin submitting applications automatically.</p>
      <a href="#" class="cbtn"
        onclick="window.showToast('Connection link copied &mdash; open in Chrome to activate.','info');return false">
        Connect extension &rarr;
      </a>
    </div>
    <div style="text-align:right;flex-shrink:0">
      <div style="font-size:11px;color:var(--text-tertiary);text-transform:uppercase;
        letter-spacing:.06em;margin-bottom:8px">Extension status</div>
      <span class="badge badge-green" style="font-size:12px">&#8226; Connected</span>
    </div>
  </div>
</div>

<!-- Scheme switcher -->
<div class="scheme-bar">
  <span class="scheme-label">&#127912; Colour scheme: {scheme_name}</span>
  <div class="scheme-links">
    <a href="/dashboard-preview-gold" class="scheme-link {'current' if scheme_name == 'Obsidian + Gold' else ''}">&#127775; Gold</a>
    <a href="/dashboard-preview-violet" class="scheme-link {'current' if scheme_name == 'Indigo + Violet' else ''}">&#128420; Violet</a>
    <a href="/dashboard-preview-platinum" class="scheme-link {'current' if scheme_name == 'Black + Platinum' else ''}">&#9898; Platinum</a>
    <a href="/dashboard-preview-cyan" class="scheme-link {'current' if scheme_name == 'Original Cyan' else ''}">&#128302; Original</a>
  </div>
  <a href="/design-system" class="scheme-back">&larr; Design system</a>
</div>

<script>
(function(){{
  var counters=document.querySelectorAll('[data-count]');
  var started=false;
  function run(){{
    if(started)return;started=true;
    counters.forEach(function(el){{
      var t=parseInt(el.dataset.target,10),steps=40,step=0;
      var iv=setInterval(function(){{
        step++;var p=step/steps;var e=1-Math.pow(1-p,3);
        el.textContent=Math.round(e*t);
        if(step>=steps){{el.textContent=t;clearInterval(iv);}}
      }},900/steps);
    }});
  }}
  if('IntersectionObserver' in window){{
    var obs=new IntersectionObserver(function(e){{if(e[0].isIntersecting)run();}},{{threshold:.3}});
    var g=document.querySelector('.sg');if(g)obs.observe(g);
  }}else{{setTimeout(run,400);}}
  if(window.matchMedia('(prefers-reduced-motion:reduce)').matches)
    counters.forEach(function(el){{el.textContent=el.dataset.target;}});
}})();
</script>"""


GOLD_TOKENS = dict(
    bg_base='#080602', bg_surface='#0f0c08', bg_elevated='#16120a', bg_overlay='#1e180d',
    accent='#d4a017', accent_dim='#b8880f', accent_glow='rgba(212,160,23,.14)',
    accent_surface='rgba(212,160,23,.07)',
    text_primary='#f2ead8', text_secondary='#a89070', text_tertiary='#6b5a3e',
    border='#251e10', border_bright='#3a2e18',
    success='#4ade80', warning='#f59e0b', error='#f87171',
    success_glow='rgba(74,222,128,.25)', success_glow_far='rgba(74,222,128,.06)',
    nav_bg='rgba(8,6,2,.92)', avatar_text='#080602',
    btn_text='#080602', row_border='rgba(37,30,16,.7)',
    ext_icon_border='rgba(212,160,23,.2)',
)

GOLD_LOGO = """<svg width="28" height="28" viewBox="0 0 28 28" fill="none" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <radialGradient id="gl-g" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#d4a017" stop-opacity=".3"/>
      <stop offset="100%" stop-color="#d4a017" stop-opacity="0"/>
    </radialGradient>
  </defs>
  <circle cx="14" cy="14" r="13" fill="url(#gl-g)"/>
  <circle cx="14" cy="14" r="12" stroke="#3a2e18" stroke-width="1.2"/>
  <circle cx="14" cy="14" r="8" stroke="#d4a017" stroke-width="1.4" stroke-opacity=".5">
    <animate attributeName="r" values="8;9.5;8" dur="3s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values=".5;.15;.5" dur="3s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="4.5" stroke="#f2c94c" stroke-width="1.6">
    <animate attributeName="r" values="4.5;5.5;4.5" dur="2.4s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values="1;.4;1" dur="2.4s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="2" fill="#d4a017">
    <animate attributeName="opacity" values="1;.5;1" dur="1.8s" repeatCount="indefinite"/>
  </circle>
</svg>"""

VIOLET_TOKENS = dict(
    bg_base='#06040f', bg_surface='#0d0a1a', bg_elevated='#130f24', bg_overlay='#1a152e',
    accent='#a855f7', accent_dim='#7c3aed', accent_glow='rgba(168,85,247,.14)',
    accent_surface='rgba(168,85,247,.07)',
    text_primary='#ede9fe', text_secondary='#9585c4', text_tertiary='#5c4d7a',
    border='#1e1535', border_bright='#2e2050',
    success='#34d399', warning='#fbbf24', error='#f87171',
    success_glow='rgba(52,211,153,.25)', success_glow_far='rgba(52,211,153,.06)',
    nav_bg='rgba(6,4,15,.92)', avatar_text='#06040f',
    btn_text='#06040f', row_border='rgba(30,21,53,.7)',
    ext_icon_border='rgba(168,85,247,.2)',
)

VIOLET_LOGO = """<svg width="28" height="28" viewBox="0 0 28 28" fill="none" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <radialGradient id="vl-g" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#a855f7" stop-opacity=".3"/>
      <stop offset="100%" stop-color="#a855f7" stop-opacity="0"/>
    </radialGradient>
  </defs>
  <circle cx="14" cy="14" r="13" fill="url(#vl-g)"/>
  <circle cx="14" cy="14" r="12" stroke="#2e2050" stroke-width="1.2"/>
  <circle cx="14" cy="14" r="8" stroke="#a855f7" stroke-width="1.4" stroke-opacity=".5">
    <animate attributeName="r" values="8;9.5;8" dur="3s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values=".5;.15;.5" dur="3s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="4.5" stroke="#c084fc" stroke-width="1.6">
    <animate attributeName="r" values="4.5;5.5;4.5" dur="2.4s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values="1;.4;1" dur="2.4s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="2" fill="#a855f7">
    <animate attributeName="opacity" values="1;.5;1" dur="1.8s" repeatCount="indefinite"/>
  </circle>
</svg>"""

PLATINUM_TOKENS = dict(
    bg_base='#000000', bg_surface='#0a0a0a', bg_elevated='#111111', bg_overlay='#1a1a1a',
    accent='#e2e2e2', accent_dim='#c0c0c0', accent_glow='rgba(226,226,226,.1)',
    accent_surface='rgba(226,226,226,.04)',
    text_primary='#f5f5f5', text_secondary='#7a7a7a', text_tertiary='#444444',
    border='#1c1c1c', border_bright='#2e2e2e',
    success='#6ee7b7', warning='#fcd34d', error='#fca5a5',
    success_glow='rgba(110,231,183,.2)', success_glow_far='rgba(110,231,183,.05)',
    nav_bg='rgba(0,0,0,.95)', avatar_text='#000000',
    btn_text='#000000', row_border='rgba(28,28,28,.8)',
    ext_icon_border='rgba(226,226,226,.12)',
)

PLATINUM_LOGO = """<svg width="28" height="28" viewBox="0 0 28 28" fill="none" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <radialGradient id="pl-g" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#e2e2e2" stop-opacity=".2"/>
      <stop offset="100%" stop-color="#e2e2e2" stop-opacity="0"/>
    </radialGradient>
  </defs>
  <circle cx="14" cy="14" r="13" fill="url(#pl-g)"/>
  <circle cx="14" cy="14" r="12" stroke="#2e2e2e" stroke-width="1.2"/>
  <circle cx="14" cy="14" r="8" stroke="#c0c0c0" stroke-width="1.4" stroke-opacity=".4">
    <animate attributeName="r" values="8;9.5;8" dur="3s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values=".4;.1;.4" dur="3s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="4.5" stroke="#e2e2e2" stroke-width="1.6">
    <animate attributeName="r" values="4.5;5.5;4.5" dur="2.4s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values="1;.3;1" dur="2.4s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="2" fill="#e2e2e2">
    <animate attributeName="opacity" values="1;.4;1" dur="1.8s" repeatCount="indefinite"/>
  </circle>
</svg>"""

CYAN_TOKENS = dict(
    bg_base='#080c12', bg_surface='#0e1520', bg_elevated='#141d2e', bg_overlay='#1a2640',
    accent='#00c8ff', accent_dim='#0095cc', accent_glow='rgba(0,200,255,.12)',
    accent_surface='rgba(0,200,255,.06)',
    text_primary='#e8f1ff', text_secondary='#8ea4c8', text_tertiary='#526278',
    border='#1e2d42', border_bright='#2d4060',
    success='#10d68a', warning='#f5a623', error='#ff4d6a',
    success_glow='rgba(16,214,138,.25)', success_glow_far='rgba(16,214,138,.06)',
    nav_bg='rgba(8,12,18,.92)', avatar_text='#080c12',
    btn_text='#080c12', row_border='rgba(30,45,66,.6)',
    ext_icon_border='rgba(0,200,255,.15)',
)

CYAN_LOGO = """<svg width="28" height="28" viewBox="0 0 28 28" fill="none" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <radialGradient id="cy-g" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#00c8ff" stop-opacity=".25"/>
      <stop offset="100%" stop-color="#00c8ff" stop-opacity="0"/>
    </radialGradient>
  </defs>
  <circle cx="14" cy="14" r="13" fill="url(#cy-g)"/>
  <circle cx="14" cy="14" r="12" stroke="#1e2d42" stroke-width="1.2"/>
  <circle cx="14" cy="14" r="8" stroke="#00c8ff" stroke-width="1.4" stroke-opacity=".5">
    <animate attributeName="r" values="8;9.5;8" dur="3s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values=".5;.15;.5" dur="3s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="4.5" stroke="#10d68a" stroke-width="1.6">
    <animate attributeName="r" values="4.5;5.5;4.5" dur="2.4s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values="1;.4;1" dur="2.4s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="2" fill="#00c8ff">
    <animate attributeName="opacity" values="1;.5;1" dur="1.8s" repeatCount="indefinite"/>
  </circle>
</svg>"""


@app.route("/dashboard-preview-gold")
def dashboard_preview_gold():
    body = _dash_preview_html(
        scheme_name='Obsidian + Gold', tokens=GOLD_TOKENS,
        logo_svg=GOLD_LOGO,
        nav_links_html='', banner_label='Gold',
        banner_bg='#0f0c08', banner_border='#251e10',
        banner_text_color='#a89070', banner_accent_color='#d4a017',
    )
    return _page("Dashboard — Gold Preview", body)


@app.route("/dashboard-preview-violet")
def dashboard_preview_violet():
    body = _dash_preview_html(
        scheme_name='Indigo + Violet', tokens=VIOLET_TOKENS,
        logo_svg=VIOLET_LOGO,
        nav_links_html='', banner_label='Violet',
        banner_bg='#0d0a1a', banner_border='#1e1535',
        banner_text_color='#9585c4', banner_accent_color='#a855f7',
    )
    return _page("Dashboard — Violet Preview", body)


@app.route("/dashboard-preview-platinum")
def dashboard_preview_platinum():
    body = _dash_preview_html(
        scheme_name='Black + Platinum', tokens=PLATINUM_TOKENS,
        logo_svg=PLATINUM_LOGO,
        nav_links_html='', banner_label='Platinum',
        banner_bg='#0a0a0a', banner_border='#1c1c1c',
        banner_text_color='#7a7a7a', banner_accent_color='#e2e2e2',
    )
    return _page("Dashboard — Platinum Preview", body)


@app.route("/dashboard-preview-cyan")
def dashboard_preview_cyan_alt():
    body = _dash_preview_html(
        scheme_name='Original Cyan', tokens=CYAN_TOKENS,
        logo_svg=CYAN_LOGO,
        nav_links_html='', banner_label='Cyan',
        banner_bg='#0e1520', banner_border='#1e2d42',
        banner_text_color='#8ea4c8', banner_accent_color='#00c8ff',
    )
    return _page("Dashboard — Cyan Preview", body)


# ── Bright / light-mode colour scheme previews ────────────────────────────────
def _bright_preview_html(scheme_name, css_vars, logo_svg, badge_overrides=""):
    """Renders a full light-mode dashboard preview."""
    return f"""
<style>
/* Hide the dark _page() chrome — we render our own */
#main-nav, footer {{ display:none !important }}
body {{ background:{css_vars['bg_base']};margin:0;padding:0;
  font-family:'DM Sans',system-ui,sans-serif;color:{css_vars['text_primary']} }}

/* ── CSS vars override ── */
:root {{
  --bg-base:{css_vars['bg_base']};
  --bg-surface:{css_vars['bg_surface']};
  --bg-elevated:{css_vars['bg_elevated']};
  --bg-overlay:{css_vars['bg_overlay']};
  --accent:{css_vars['accent']};
  --accent-dim:{css_vars['accent_dim']};
  --accent-glow:{css_vars['accent_glow']};
  --accent-surface:{css_vars['accent_surface']};
  --text-primary:{css_vars['text_primary']};
  --text-secondary:{css_vars['text_secondary']};
  --text-tertiary:{css_vars['text_tertiary']};
  --border:{css_vars['border']};
  --border-bright:{css_vars['border_bright']};
  --success:#059669;
  --warning:#d97706;
  --error:#dc2626;
  --font-display:'Clash Display',sans-serif;
  --font-body:'DM Sans',sans-serif;
  --font-mono:'JetBrains Mono',monospace;
}}

/* ── Badge overrides for light mode ── */
.badge-green {{ background:#dcfce7;color:#15803d;border:none }}
.badge-red   {{ background:#fee2e2;color:#dc2626;border:none }}
.badge-blue  {{ background:#dbeafe;color:#1d4ed8;border:none }}
.badge-yellow{{ background:#fef9c3;color:#a16207;border:none }}
.badge-gray  {{ background:#f1f5f9;color:#475569;border:none }}
.badge       {{ font-weight:600;font-size:11px;padding:3px 10px;border-radius:20px;
                display:inline-block;letter-spacing:.03em }}
{badge_overrides}

/* ── Nav ── */
.bv-nav {{
  position:fixed;top:0;left:0;right:0;z-index:200;height:60px;
  background:{css_vars['nav_bg']};border-bottom:1px solid {css_vars['border']};
  display:flex;align-items:center;justify-content:space-between;padding:0 32px;
  box-shadow:0 1px 3px {css_vars['shadow']};
}}
.bv-brand {{ display:flex;align-items:center;gap:10px;text-decoration:none;
  font-family:var(--font-display);font-size:18px;font-weight:700;
  color:{css_vars['text_primary']};letter-spacing:-.01em }}
.bv-links {{ display:flex;gap:4px }}
.bv-link {{ padding:7px 14px;border-radius:8px;font-size:13px;font-weight:500;
  color:{css_vars['text_secondary']};text-decoration:none;transition:all .15s }}
.bv-link:hover {{ color:{css_vars['text_primary']};background:{css_vars['bg_elevated']} }}
.bv-link.on {{ color:{css_vars['accent']};background:{css_vars['accent_surface']};font-weight:600 }}
.bv-right {{ display:flex;align-items:center;gap:12px }}
.bv-trial {{ font-size:12px;color:{css_vars['text_tertiary']};font-family:var(--font-mono) }}
.bv-avatar {{ width:34px;height:34px;border-radius:50%;
  background:linear-gradient(135deg,{css_vars['accent']},{css_vars['accent_dim']});
  display:flex;align-items:center;justify-content:center;
  font-size:12px;font-weight:700;color:{css_vars['avatar_text']};
  font-family:var(--font-body) }}

/* ── Wrapper ── */
.bw {{ max-width:1100px;margin:0 auto;padding:84px 24px 80px;
  animation:bfu .4s cubic-bezier(.16,1,.3,1) both }}
@keyframes bfu {{ from{{opacity:0;transform:translateY(14px)}}to{{opacity:1;transform:none}} }}

/* ── Hero ── */
.bh {{ display:flex;align-items:flex-start;justify-content:space-between;
  flex-wrap:wrap;gap:20px;margin-bottom:32px }}
.bg {{ font-family:var(--font-display);font-size:32px;font-weight:700;
  letter-spacing:-.01em;color:{css_vars['text_primary']};margin:0 0 8px }}
.bsub {{ font-size:13px;color:{css_vars['text_secondary']};display:flex;
  align-items:center;gap:8px;margin:0 }}
.bdot {{ width:7px;height:7px;border-radius:50%;background:#059669;
  box-shadow:0 0 0 2px rgba(5,150,105,.2);animation:bpd 2.5s ease infinite }}
@keyframes bpd {{ 0%,100%{{box-shadow:0 0 0 2px rgba(5,150,105,.2)}}
  50%{{box-shadow:0 0 0 5px rgba(5,150,105,.07)}} }}

/* ── Run form ── */
.brf {{ display:flex;gap:10px;align-items:center;flex-wrap:wrap }}
.brs {{ padding:10px 16px;background:{css_vars['bg_surface']};
  border:1.5px solid {css_vars['border']};border-radius:10px;
  color:{css_vars['text_primary']};font-family:var(--font-body);font-size:13px;
  cursor:pointer;transition:border-color .15s }}
.brs:hover,.brs:focus {{ border-color:{css_vars['accent']};outline:none }}
.brbtn {{ display:inline-flex;align-items:center;gap:8px;
  padding:11px 22px;background:{css_vars['accent']};color:{css_vars['avatar_text']};
  border:none;border-radius:10px;font-family:var(--font-body);font-size:14px;
  font-weight:700;cursor:pointer;transition:all .12s ease;letter-spacing:.01em;
  box-shadow:0 2px 12px {css_vars['btn_shadow']} }}
.brbtn:hover {{ background:{css_vars['accent_dim']};transform:scale(1.02) }}
.brbtn:active {{ transform:scale(.97) }}

/* ── Stat grid ── */
.bsg {{ display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-bottom:28px }}
.bsc {{ background:{css_vars['bg_surface']};border:1.5px solid {css_vars['border']};
  border-radius:14px;padding:24px;transition:all .2s;position:relative;overflow:hidden }}
.bsc:hover {{ border-color:{css_vars['accent']};transform:translateY(-2px);
  box-shadow:0 8px 24px {css_vars['card_shadow']} }}
.bsi {{ font-size:20px;margin-bottom:14px;display:block }}
.bsv {{ font-family:var(--font-display);font-size:42px;font-weight:700;
  color:{css_vars['text_primary']};letter-spacing:-.02em;line-height:1;margin-bottom:6px }}
.bsl {{ font-size:11px;color:{css_vars['text_tertiary']};text-transform:uppercase;
  letter-spacing:.07em;font-weight:500 }}
.bsc- {{ font-size:12px;color:#059669;margin-top:10px;font-weight:600 }}
.bsc-a {{ font-size:12px;color:{css_vars['accent']};margin-top:10px;font-weight:600 }}

/* ── Section header ── */
.bsh {{ display:flex;align-items:center;justify-content:space-between;margin-bottom:14px }}
.bst {{ font-family:var(--font-display);font-size:16px;font-weight:600;
  color:{css_vars['text_primary']};letter-spacing:-.01em }}
.bsa {{ font-size:12px;color:{css_vars['accent']};text-decoration:none;font-weight:600 }}

/* ── Runs card ── */
.brc {{ background:{css_vars['bg_surface']};border:1.5px solid {css_vars['border']};
  border-radius:14px;overflow:hidden;margin-bottom:24px;
  box-shadow:0 1px 4px {css_vars['shadow']} }}
.brt {{ width:100%;border-collapse:collapse }}
.brt th {{ padding:13px 20px;background:{css_vars['bg_elevated']};
  font-size:11px;font-weight:600;color:{css_vars['text_tertiary']};
  text-transform:uppercase;letter-spacing:.08em;text-align:left;
  border-bottom:1.5px solid {css_vars['border']} }}
.brt td {{ padding:14px 20px;border-bottom:1px solid {css_vars['border']};
  font-size:13px;color:{css_vars['text_primary']} }}
.brt tr:last-child td {{ border-bottom:none }}
.brt tr:hover td {{ background:{css_vars['bg_elevated']} }}
.brd {{ font-family:var(--font-mono);font-size:12px;color:{css_vars['text_secondary']} }}
.bsb-wrap {{ display:flex;align-items:center;gap:10px }}
.bsb-track {{ flex:1;max-width:80px;height:5px;background:{css_vars['bg_overlay']};
  border-radius:3px;overflow:hidden }}
.bsb-fill {{ height:100%;border-radius:3px;
  background:linear-gradient(90deg,{css_vars['accent_dim']},{css_vars['accent']});
  transition:width .8s cubic-bezier(.16,1,.3,1) }}
.bsb-val {{ font-family:var(--font-mono);font-size:12px;color:{css_vars['accent']};
  font-weight:700;min-width:26px }}
.bvb {{ font-size:12px;color:{css_vars['accent']};text-decoration:none;
  font-weight:600;opacity:.8;transition:opacity .15s }}
.bvb:hover {{ opacity:1 }}

/* ── Extension card ── */
.bec {{ background:{css_vars['bg_surface']};border:1.5px solid {css_vars['border']};
  border-radius:14px;padding:32px;display:flex;align-items:center;
  gap:32px;flex-wrap:wrap;box-shadow:0 1px 4px {css_vars['shadow']} }}
.bei {{ width:64px;height:64px;border-radius:16px;background:{css_vars['accent_surface']};
  border:1.5px solid {css_vars['accent_border']};display:flex;align-items:center;
  justify-content:center;font-size:28px;flex-shrink:0 }}
.beb {{ flex:1;min-width:200px }}
.bet {{ font-family:var(--font-display);font-size:18px;font-weight:700;
  color:{css_vars['text_primary']};margin:0 0 8px }}
.bed {{ font-size:13px;color:{css_vars['text_secondary']};margin:0 0 20px;
  max-width:480px;line-height:1.6 }}
.bcbtn {{ display:inline-flex;align-items:center;gap:8px;
  padding:11px 24px;background:{css_vars['bg_elevated']};
  border:1.5px solid {css_vars['border_bright']};border-radius:10px;
  color:{css_vars['text_primary']};font-family:var(--font-body);font-size:14px;
  font-weight:600;cursor:pointer;text-decoration:none;transition:all .15s }}
.bcbtn:hover {{ border-color:{css_vars['accent']};color:{css_vars['accent']} }}

/* ── Scheme bar ── */
.bbar {{
  position:fixed;bottom:0;left:0;right:0;z-index:300;
  background:{css_vars['nav_bg']};border-top:1.5px solid {css_vars['border']};
  padding:11px 24px;display:flex;align-items:center;
  justify-content:center;gap:20px;flex-wrap:wrap;
  box-shadow:0 -1px 4px {css_vars['shadow']};
}}
.bbar-label {{ font-size:13px;color:{css_vars['text_secondary']};font-weight:700 }}
.bbar-links {{ display:flex;gap:6px }}
.bbar-link {{ padding:6px 14px;border-radius:8px;font-size:12px;font-weight:600;
  text-decoration:none;border:1.5px solid {css_vars['border']};
  color:{css_vars['text_secondary']};transition:all .15s;background:transparent }}
.bbar-link:hover,.bbar-link.cur {{
  background:{css_vars['accent']};color:{css_vars['avatar_text']};
  border-color:{css_vars['accent']} }}
.bbar-back {{ font-size:12px;color:{css_vars['text_tertiary']};
  text-decoration:none;margin-left:8px }}
.bbar-back:hover {{ color:{css_vars['text_secondary']} }}

@media(max-width:680px){{
  .bsg{{grid-template-columns:1fr}}
  .bh{{flex-direction:column}}
  .bec{{flex-direction:column;gap:20px}}
  .bv-nav{{padding:0 16px}}
  .bv-links{{display:none}}
}}
</style>

<!-- Nav -->
<nav class="bv-nav">
  <a href="#" class="bv-brand">{logo_svg} ApplyExpress</a>
  <div class="bv-links">
    <a class="bv-link on" href="#">Dashboard</a>
    <a class="bv-link" href="#">Runs</a>
    <a class="bv-link" href="#">History</a>
    <a class="bv-link" href="#">Settings</a>
  </div>
  <div class="bv-right">
    <span class="bv-trial">14-day trial &middot; 11 days left</span>
    <div class="bv-avatar">SC</div>
  </div>
</nav>

<div class="bw">

  <!-- Hero -->
  <div class="bh">
    <div>
      <h1 class="bg">Hi, Sarah &#x1F44B;</h1>
      <p class="bsub">
        <span class="bdot"></span>
        <span class="badge badge-green">CV uploaded</span>
        &nbsp;
        <span style="color:var(--text-tertiary)">Last run: 2 hours ago</span>
      </p>
    </div>
    <div class="brf">
      <select class="brs">
        <option>&#8212; Default CV &#8212;</option>
        <option>AML Compliance</option>
        <option>Risk Manager</option>
      </select>
      <button class="brbtn"
        onclick="window.showToast('Pipeline started &mdash; the extension will begin applying shortly.','success')">
        &#9654;&nbsp; Run now
      </button>
    </div>
  </div>

  <!-- Stats -->
  <div class="bsg">
    <div class="bsc">
      <span class="bsi">&#127919;</span>
      <div class="bsv" data-target="142" data-bcount>0</div>
      <div class="bsl">Applications sent</div>
      <div class="bsc-">&#8593; 23 this week</div>
    </div>
    <div class="bsc">
      <span class="bsi">&#128269;</span>
      <div class="bsv" data-target="38" data-bcount>0</div>
      <div class="bsl">Pipeline runs</div>
      <div class="bsc-">&#8593; 4 this week</div>
    </div>
    <div class="bsc">
      <span class="bsi">&#9889;</span>
      <div class="bsv" data-target="9" data-bcount>0</div>
      <div class="bsl">Interviews</div>
      <div class="bsc-a">&#8593; 2 this month</div>
    </div>
  </div>

  <!-- Runs table -->
  <div class="bsh">
    <span class="bst">Recent runs</span>
    <a class="bsa" href="#">View all &rarr;</a>
  </div>
  <div class="brc">
    <table class="brt">
      <thead><tr>
        <th>Date &amp; time</th><th>Status</th>
        <th>Found</th><th>Applied</th><th>Avg score</th><th></th>
      </tr></thead>
      <tbody>
        <tr>
          <td class="brd">2026-04-14&nbsp; 09:12</td>
          <td><span class="badge badge-green">Completed</span></td>
          <td style="color:var(--text-secondary)">24</td>
          <td style="font-weight:700">18</td>
          <td><div class="bsb-wrap"><div class="bsb-track"><div class="bsb-fill" style="width:87%"></div></div><span class="bsb-val">8.7</span></div></td>
          <td><a href="#" class="bvb">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="brd">2026-04-13&nbsp; 21:44</td>
          <td><span class="badge badge-green">Completed</span></td>
          <td style="color:var(--text-secondary)">31</td>
          <td style="font-weight:700">22</td>
          <td><div class="bsb-wrap"><div class="bsb-track"><div class="bsb-fill" style="width:74%"></div></div><span class="bsb-val">7.4</span></div></td>
          <td><a href="#" class="bvb">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="brd">2026-04-12&nbsp; 14:08</td>
          <td><span class="badge badge-red">Failed</span></td>
          <td style="color:var(--text-secondary)">7</td>
          <td style="font-weight:700">0</td>
          <td><div class="bsb-wrap"><div class="bsb-track"><div class="bsb-fill" style="width:0%"></div></div><span class="bsb-val" style="color:#dc2626">&#8212;</span></div></td>
          <td><a href="#" class="bvb">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="brd">2026-04-11&nbsp; 09:00</td>
          <td><span class="badge badge-green">Completed</span></td>
          <td style="color:var(--text-secondary)">19</td>
          <td style="font-weight:700">14</td>
          <td><div class="bsb-wrap"><div class="bsb-track"><div class="bsb-fill" style="width:81%"></div></div><span class="bsb-val">8.1</span></div></td>
          <td><a href="#" class="bvb">View &rarr;</a></td>
        </tr>
      </tbody>
    </table>
  </div>

  <!-- Extension card -->
  <div class="bec">
    <div class="bei">&#128279;</div>
    <div class="beb">
      <h2 class="bet">Connect your Chrome extension</h2>
      <p class="bed">One click to sync the extension with your account. No copying API keys.
        Once connected, press <strong>Start</strong> in the popup to begin submitting
        applications automatically.</p>
      <a href="#" class="bcbtn"
        onclick="window.showToast('Connection link copied &mdash; open in Chrome to activate.','info');return false">
        Connect extension &rarr;
      </a>
    </div>
    <div style="text-align:right;flex-shrink:0">
      <div style="font-size:11px;color:var(--text-tertiary);text-transform:uppercase;
        letter-spacing:.06em;margin-bottom:8px">Extension status</div>
      <span class="badge badge-green">&#8226; Connected</span>
    </div>
  </div>

</div>

<!-- Scheme switcher -->
<div class="bbar">
  <span class="bbar-label">&#127759; Bright scheme: {scheme_name}</span>
  <div class="bbar-links">
    <a href="/dashboard-preview-cobalt"
      class="bbar-link {'cur' if scheme_name == 'Brilliant White + Cobalt' else ''}">&#128274; Cobalt</a>
    <a href="/dashboard-preview-emerald"
      class="bbar-link {'cur' if scheme_name == 'Warm Paper + Emerald' else ''}">&#127807; Emerald</a>
    <a href="/dashboard-preview-rose"
      class="bbar-link {'cur' if scheme_name == 'Cloud + Rose' else ''}">&#127774; Rose</a>
    <a href="/dashboard-preview-gold"
      class="bbar-link">&#9680; Dark options</a>
  </div>
  <a href="/design-system" class="bbar-back">&larr; Design system</a>
</div>

<script>
(function(){{
  var counters=document.querySelectorAll('[data-bcount]');
  var started=false;
  function run(){{
    if(started)return;started=true;
    counters.forEach(function(el){{
      var t=parseInt(el.dataset.target,10),steps=40,step=0;
      var iv=setInterval(function(){{
        step++;var p=step/steps;var e=1-Math.pow(1-p,3);
        el.textContent=Math.round(e*t);
        if(step>=steps){{el.textContent=t;clearInterval(iv);}}
      }},900/steps);
    }});
  }}
  if('IntersectionObserver' in window){{
    var obs=new IntersectionObserver(function(e){{if(e[0].isIntersecting)run();}},{{threshold:.3}});
    var g=document.querySelector('.bsg');if(g)obs.observe(g);
  }}else{{setTimeout(run,400);}}
  if(window.matchMedia('(prefers-reduced-motion:reduce)').matches)
    counters.forEach(function(el){{el.textContent=el.dataset.target;}});
}})();
</script>"""


# ── Bright Scheme 1: Brilliant White + Cobalt Blue ───────────────────────────
COBALT_VARS = dict(
    bg_base='#ffffff', bg_surface='#f8faff', bg_elevated='#eef2ff', bg_overlay='#e0e8ff',
    accent='#2563eb', accent_dim='#1d4ed8', accent_glow='rgba(37,99,235,.15)',
    accent_surface='rgba(37,99,235,.07)', accent_border='rgba(37,99,235,.2)',
    text_primary='#0f172a', text_secondary='#475569', text_tertiary='#94a3b8',
    border='#e2e8f0', border_bright='#bfdbfe',
    nav_bg='rgba(255,255,255,.95)',
    avatar_text='#ffffff', btn_shadow='rgba(37,99,235,.35)',
    card_shadow='rgba(37,99,235,.08)', shadow='rgba(15,23,42,.06)',
)

COBALT_LOGO = """<svg width="28" height="28" viewBox="0 0 28 28" fill="none">
  <defs>
    <radialGradient id="cb-g" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#2563eb" stop-opacity=".15"/>
      <stop offset="100%" stop-color="#2563eb" stop-opacity="0"/>
    </radialGradient>
  </defs>
  <circle cx="14" cy="14" r="13" fill="url(#cb-g)"/>
  <circle cx="14" cy="14" r="12" stroke="#bfdbfe" stroke-width="1.2"/>
  <circle cx="14" cy="14" r="8" stroke="#2563eb" stroke-width="1.4" stroke-opacity=".5">
    <animate attributeName="r" values="8;9.5;8" dur="3s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values=".5;.15;.5" dur="3s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="4.5" stroke="#3b82f6" stroke-width="1.6">
    <animate attributeName="r" values="4.5;5.5;4.5" dur="2.4s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values="1;.4;1" dur="2.4s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="2" fill="#2563eb">
    <animate attributeName="opacity" values="1;.5;1" dur="1.8s" repeatCount="indefinite"/>
  </circle>
</svg>"""

# ── Bright Scheme 2: Warm Paper + Emerald ────────────────────────────────────
EMERALD_VARS = dict(
    bg_base='#fdfaf5', bg_surface='#f6f1e8', bg_elevated='#ede7d8', bg_overlay='#e3dbc8',
    accent='#059669', accent_dim='#047857', accent_glow='rgba(5,150,105,.15)',
    accent_surface='rgba(5,150,105,.08)', accent_border='rgba(5,150,105,.22)',
    text_primary='#1a1209', text_secondary='#5c4f3a', text_tertiary='#9c8870',
    border='#dfd6c4', border_bright='#6ee7b7',
    nav_bg='rgba(253,250,245,.97)',
    avatar_text='#ffffff', btn_shadow='rgba(5,150,105,.3)',
    card_shadow='rgba(5,150,105,.07)', shadow='rgba(26,18,9,.06)',
)

EMERALD_LOGO = """<svg width="28" height="28" viewBox="0 0 28 28" fill="none">
  <defs>
    <radialGradient id="em-g" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#059669" stop-opacity=".18"/>
      <stop offset="100%" stop-color="#059669" stop-opacity="0"/>
    </radialGradient>
  </defs>
  <circle cx="14" cy="14" r="13" fill="url(#em-g)"/>
  <circle cx="14" cy="14" r="12" stroke="#dfd6c4" stroke-width="1.2"/>
  <circle cx="14" cy="14" r="8" stroke="#059669" stroke-width="1.4" stroke-opacity=".5">
    <animate attributeName="r" values="8;9.5;8" dur="3s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values=".5;.15;.5" dur="3s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="4.5" stroke="#10b981" stroke-width="1.6">
    <animate attributeName="r" values="4.5;5.5;4.5" dur="2.4s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values="1;.4;1" dur="2.4s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="2" fill="#059669">
    <animate attributeName="opacity" values="1;.5;1" dur="1.8s" repeatCount="indefinite"/>
  </circle>
</svg>"""

# ── Bright Scheme 3: Cloud White + Rose ──────────────────────────────────────
ROSE_VARS = dict(
    bg_base='#fff8f9', bg_surface='#fff0f2', bg_elevated='#ffe4e8', bg_overlay='#fecdd3',
    accent='#e11d48', accent_dim='#be123c', accent_glow='rgba(225,29,72,.15)',
    accent_surface='rgba(225,29,72,.07)', accent_border='rgba(225,29,72,.2)',
    text_primary='#1a0510', text_secondary='#6b3040', text_tertiary='#b07080',
    border='#fecdd3', border_bright='#fda4af',
    nav_bg='rgba(255,248,249,.97)',
    avatar_text='#ffffff', btn_shadow='rgba(225,29,72,.3)',
    card_shadow='rgba(225,29,72,.07)', shadow='rgba(26,5,16,.05)',
)

ROSE_LOGO = """<svg width="28" height="28" viewBox="0 0 28 28" fill="none">
  <defs>
    <radialGradient id="rs-g" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#e11d48" stop-opacity=".18"/>
      <stop offset="100%" stop-color="#e11d48" stop-opacity="0"/>
    </radialGradient>
  </defs>
  <circle cx="14" cy="14" r="13" fill="url(#rs-g)"/>
  <circle cx="14" cy="14" r="12" stroke="#fecdd3" stroke-width="1.2"/>
  <circle cx="14" cy="14" r="8" stroke="#e11d48" stroke-width="1.4" stroke-opacity=".5">
    <animate attributeName="r" values="8;9.5;8" dur="3s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values=".5;.15;.5" dur="3s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="4.5" stroke="#fb7185" stroke-width="1.6">
    <animate attributeName="r" values="4.5;5.5;4.5" dur="2.4s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values="1;.4;1" dur="2.4s" repeatCount="indefinite"/>
  </circle>
  <circle cx="14" cy="14" r="2" fill="#e11d48">
    <animate attributeName="opacity" values="1;.5;1" dur="1.8s" repeatCount="indefinite"/>
  </circle>
</svg>"""


@app.route("/dashboard-preview-cobalt")
def dashboard_preview_cobalt():
    return _page("Dashboard — Cobalt Preview",
        _bright_preview_html('Brilliant White + Cobalt', COBALT_VARS, COBALT_LOGO))

@app.route("/dashboard-preview-emerald")
def dashboard_preview_emerald():
    return _page("Dashboard — Emerald Preview",
        _bright_preview_html('Warm Paper + Emerald', EMERALD_VARS, EMERALD_LOGO))

@app.route("/dashboard-preview-rose")
def dashboard_preview_rose():
    return _page("Dashboard — Rose Preview",
        _bright_preview_html('Cloud + Rose', ROSE_VARS, ROSE_LOGO))


# ── Poppy colour schemes (from uploaded flower photo) ─────────────────────────
# Colours extracted: vivid orange #e8621a, amber #f5a84e, warm cream #fdf8f2,
# deep poppy dark #120804, dusty sage-teal #7a9e8a

POPPY_LOGO = """<svg width="28" height="28" viewBox="0 0 28 28" fill="none">
  <defs>
    <radialGradient id="pp-g" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#e8621a" stop-opacity=".25"/>
      <stop offset="100%" stop-color="#f5a84e" stop-opacity="0"/>
    </radialGradient>
    <radialGradient id="pp-c" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="#e8621a"/>
      <stop offset="100%" stop-color="#f5a84e"/>
    </radialGradient>
  </defs>
  <circle cx="14" cy="14" r="13" fill="url(#pp-g)"/>
  <circle cx="14" cy="14" r="12" stroke="#f5d4b0" stroke-width="1.2"/>
  <!-- outer ring: amber -->
  <circle cx="14" cy="14" r="9" stroke="#f5a84e" stroke-width="1.2" stroke-opacity=".45">
    <animate attributeName="r" values="9;10.5;9" dur="3.2s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values=".45;.1;.45" dur="3.2s" repeatCount="indefinite"/>
  </circle>
  <!-- inner ring: poppy orange -->
  <circle cx="14" cy="14" r="5" stroke="#e8621a" stroke-width="1.8">
    <animate attributeName="r" values="5;6;5" dur="2.4s" repeatCount="indefinite"/>
    <animate attributeName="stroke-opacity" values="1;.35;1" dur="2.4s" repeatCount="indefinite"/>
  </circle>
  <!-- sage teal hint ring -->
  <circle cx="14" cy="14" r="11.5" stroke="#7a9e8a" stroke-width=".8" stroke-opacity=".3">
    <animate attributeName="stroke-opacity" values=".3;.07;.3" dur="4s" repeatCount="indefinite"/>
  </circle>
  <!-- centre dot -->
  <circle cx="14" cy="14" r="2.2" fill="url(#pp-c)">
    <animate attributeName="opacity" values="1;.5;1" dur="1.8s" repeatCount="indefinite"/>
  </circle>
</svg>"""

# ── Poppy Light: Warm Cream + Vivid Orange ───────────────────────────────────
POPPY_LIGHT_VARS = dict(
    bg_base    = '#fdf8f2',   # warm white — cream from petal highlights
    bg_surface = '#fef2e4',   # soft peach-cream
    bg_elevated= '#fde5cc',   # deeper warm peach
    bg_overlay = '#fad4ae',   # muted amber tint
    accent     = '#e8621a',   # vivid poppy orange
    accent_dim = '#c94d0f',   # deeper burnt orange on hover
    accent_glow= 'rgba(232,98,26,.18)',
    accent_surface='rgba(232,98,26,.08)',
    accent_border ='rgba(232,98,26,.25)',
    text_primary  = '#120804',   # near-black warm dark (the poppy centre)
    text_secondary= '#7a4830',   # mid warm brown
    text_tertiary = '#b8866a',   # light terracotta
    border        = '#f0d0a8',   # warm sandy border
    border_bright = '#f5a84e',   # amber — petal gold
    nav_bg        = 'rgba(253,248,242,.97)',
    avatar_text   = '#ffffff',
    btn_shadow    = 'rgba(232,98,26,.32)',
    card_shadow   = 'rgba(232,98,26,.08)',
    shadow        = 'rgba(18,8,4,.07)',
)

# ── Poppy Dark: Deep Warm Night + Orange/Amber ────────────────────────────────
POPPY_DARK_VARS = dict(
    bg_base    = '#0e0604',   # deep warm black — poppy centre
    bg_surface = '#180b06',   # dark warm surface
    bg_elevated= '#211008',   # elevated warm dark
    bg_overlay = '#2c150a',   # overlay warm dark
    accent     = '#f07030',   # vivid orange — slightly softened for dark bg
    accent_dim = '#e8621a',   # poppy orange
    accent_glow= 'rgba(240,112,48,.16)',
    accent_surface='rgba(240,112,48,.08)',
    accent_border ='rgba(240,112,48,.2)',
    text_primary  = '#fdf0e0',   # warm cream — petal highlight
    text_secondary= '#c49070',   # mid warm amber
    text_tertiary = '#7a5540',   # muted terracotta
    border        = '#2c1a0e',   # warm dark border
    border_bright = '#4a2e18',   # hover border
    nav_bg        = 'rgba(14,6,4,.92)',
    avatar_text   = '#0e0604',
    btn_shadow    = 'rgba(240,112,48,.35)',
    card_shadow   = 'rgba(240,112,48,.1)',
    shadow        = 'rgba(0,0,0,.4)',
)


def _poppy_dark_html(scheme_name, tokens, logo_svg):
    """Dark-mode variant of the poppy preview — reuses _dash_preview_html structure."""
    return _dash_preview_html(
        scheme_name=scheme_name, tokens={
            **tokens,
            'success': '#4ade80', 'warning': '#fbbf24', 'error': '#f87171',
            'success_glow': 'rgba(74,222,128,.22)',
            'success_glow_far': 'rgba(74,222,128,.06)',
            'row_border': 'rgba(44,26,14,.7)',
            'ext_icon_border': 'rgba(240,112,48,.2)',
        },
        logo_svg=logo_svg,
        nav_links_html='',
        banner_label=scheme_name,
        banner_bg=tokens['bg_surface'],
        banner_border=tokens['border'],
        banner_text_color=tokens['text_secondary'],
        banner_accent_color=tokens['accent'],
    ).replace(
        'scheme-link {\'current\' if scheme_name == \'Obsidian + Gold\' else \'\'}',''
    ).replace(
        '<a href="/dashboard-preview-gold" class="scheme-link ',
        '<a href="/dashboard-preview-poppy-light" class="scheme-link '
    )


@app.route("/dashboard-preview-poppy-light")
def dashboard_preview_poppy_light():
    # Bright version uses _bright_preview_html
    extra_badge = """.badge-green{{background:#dcfce7;color:#15803d}}
.badge-red{{background:#fee2e2;color:#dc2626}}"""
    return _page("Dashboard — Poppy Light",
        _bright_preview_html('Poppy — Warm Cream + Orange', POPPY_LIGHT_VARS, POPPY_LOGO,
                             badge_overrides=extra_badge).replace(
            'href="/dashboard-preview-gold"\n      class="bbar-link">&#9680; Dark options',
            'href="/dashboard-preview-poppy-dark"\n      class="bbar-link">&#127768; Dark Poppy'
        ).replace(
            "{'cur' if scheme_name == 'Brilliant White + Cobalt' else ''}",
            "{'cur' if 'Cobalt' in 'Poppy — Warm Cream + Orange' else ''}"
        ).replace(
            "{'cur' if scheme_name == 'Warm Paper + Emerald' else ''}",
            "{'cur' if 'Emerald' in 'Poppy — Warm Cream + Orange' else ''}"
        ).replace(
            "{'cur' if scheme_name == 'Cloud + Rose' else ''}",
            "{'cur' if 'Rose' in 'Poppy — Warm Cream + Orange' else ''}"
        ).replace(
            'Bright scheme: Poppy &mdash; Warm Cream + Orange',
            '&#127799; Poppy palette &mdash; Light'
        ).replace(
            "bbar-link {'cur' if 'Cobalt'", "bbar-link "
        )
    )


@app.route("/dashboard-preview-poppy-dark")
def dashboard_preview_poppy_dark():
    t = {**POPPY_DARK_VARS,
         'btn_text':'#0e0604',
         'success':'#4ade80','warning':'#fbbf24','error':'#f87171',
         'success_glow':'rgba(74,222,128,.22)',
         'success_glow_far':'rgba(74,222,128,.06)',
         'row_border':'rgba(44,26,14,.7)',
         'ext_icon_border':'rgba(240,112,48,.2)'}
    body = _dash_preview_html(
        scheme_name='Poppy — Deep Night + Orange',
        tokens=t, logo_svg=POPPY_LOGO,
        nav_links_html='',
        banner_bg=POPPY_DARK_VARS['bg_surface'],
        banner_border=POPPY_DARK_VARS['border'],
        banner_text_color=POPPY_DARK_VARS['text_secondary'],
        banner_accent_color=POPPY_DARK_VARS['accent'],
        banner_label='Poppy Dark',
    )
    return _page("Dashboard — Poppy Dark", body)


# ── Orange-background previews ─────────────────────────────────────────────────
def _orange_bg_preview(scheme_name, text_color, text_secondary, text_tertiary,
                        card_bg, card_border, card_bg_hover,
                        nav_bg, nav_border, btn_bg, btn_text,
                        score_bar_color, switcher_other_url, switcher_other_label):
    return f"""
<style>
#main-nav, footer {{ display:none !important }}
* {{ box-sizing:border-box; margin:0; padding:0 }}
body {{
  background:#e8621a;
  font-family:'DM Sans',system-ui,sans-serif;
  color:{text_color};
  min-height:100vh;
}}

/* Nav */
.on {{ position:fixed;top:0;left:0;right:0;z-index:200;height:60px;
  background:{nav_bg};border-bottom:1px solid {nav_border};
  display:flex;align-items:center;justify-content:space-between;padding:0 32px;
  backdrop-filter:blur(10px);-webkit-backdrop-filter:blur(10px); }}
.ob {{ display:flex;align-items:center;gap:10px;text-decoration:none;
  font-family:'Clash Display',sans-serif;font-size:18px;font-weight:700;
  color:{text_color};letter-spacing:-.01em }}
.ol {{ display:flex;gap:4px }}
.olk {{ padding:7px 14px;border-radius:8px;font-size:13px;font-weight:500;
  color:{text_secondary};text-decoration:none;transition:all .15s }}
.olk:hover {{ color:{text_color};background:rgba(0,0,0,.1) }}
.olk.act {{ color:{text_color};background:rgba(0,0,0,.15);font-weight:700 }}
.or {{ display:flex;align-items:center;gap:12px }}
.ot {{ font-size:12px;color:{text_secondary};font-family:'JetBrains Mono',monospace }}
.oa {{ width:34px;height:34px;border-radius:50%;background:{btn_bg};
  display:flex;align-items:center;justify-content:center;
  font-size:12px;font-weight:700;color:{btn_text};font-family:'DM Sans',sans-serif }}

/* Wrapper */
.ow {{ max-width:1100px;margin:0 auto;padding:84px 24px 80px;
  animation:ofu .4s cubic-bezier(.16,1,.3,1) both }}
@keyframes ofu {{ from{{opacity:0;transform:translateY(14px)}}to{{opacity:1;transform:none}} }}

/* Hero */
.oh {{ display:flex;align-items:flex-start;justify-content:space-between;
  flex-wrap:wrap;gap:20px;margin-bottom:32px }}
.og {{ font-family:'Clash Display',sans-serif;font-size:36px;font-weight:700;
  letter-spacing:-.02em;color:{text_color};margin-bottom:8px }}
.osub {{ font-size:13px;color:{text_secondary};display:flex;align-items:center;gap:8px }}
.odot {{ width:8px;height:8px;border-radius:50%;background:{btn_bg};
  box-shadow:0 0 0 3px rgba(0,0,0,.1);animation:opd 2.5s ease infinite }}
@keyframes opd {{ 0%,100%{{box-shadow:0 0 0 3px rgba(0,0,0,.1)}}
  50%{{box-shadow:0 0 0 7px rgba(0,0,0,.04)}} }}

/* Badge custom (on orange bg) */
.obadge {{ display:inline-block;font-size:11px;font-weight:700;
  padding:3px 10px;border-radius:20px;letter-spacing:.03em }}
.obadge-green {{ background:rgba(0,0,0,.18);color:{text_color} }}
.obadge-red   {{ background:rgba(0,0,0,.25);color:{text_color} }}

/* Run form */
.orf {{ display:flex;gap:10px;align-items:center;flex-wrap:wrap }}
.ors {{ padding:10px 16px;background:rgba(0,0,0,.15);
  border:1.5px solid rgba(0,0,0,.2);border-radius:10px;
  color:{text_color};font-family:'DM Sans',sans-serif;font-size:13px;cursor:pointer;
  transition:border-color .15s }}
.ors:focus {{ outline:none;border-color:rgba(0,0,0,.4) }}
.orbtn {{ display:inline-flex;align-items:center;gap:8px;
  padding:11px 22px;background:{btn_bg};color:{btn_text};
  border:none;border-radius:10px;font-family:'DM Sans',sans-serif;font-size:14px;
  font-weight:700;cursor:pointer;transition:all .12s ease;letter-spacing:.01em;
  box-shadow:0 2px 12px rgba(0,0,0,.25) }}
.orbtn:hover {{ transform:scale(1.02);box-shadow:0 4px 20px rgba(0,0,0,.3) }}
.orbtn:active {{ transform:scale(.97) }}

/* Stat cards */
.osg {{ display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-bottom:28px }}
.osc {{ background:{card_bg};border:1.5px solid {card_border};
  border-radius:16px;padding:26px;transition:all .2s;position:relative;overflow:hidden }}
.osc:hover {{ background:{card_bg_hover};transform:translateY(-2px);
  box-shadow:0 10px 30px rgba(0,0,0,.15) }}
.osi {{ font-size:20px;margin-bottom:14px;display:block }}
.osv {{ font-family:'Clash Display',sans-serif;font-size:44px;font-weight:700;
  color:{text_color};letter-spacing:-.02em;line-height:1;margin-bottom:6px }}
.osl {{ font-size:11px;color:{text_secondary};text-transform:uppercase;
  letter-spacing:.08em;font-weight:600 }}
.osd {{ font-size:12px;color:{text_tertiary};margin-top:10px;font-weight:600 }}

/* Section header */
.osh {{ display:flex;align-items:center;justify-content:space-between;margin-bottom:14px }}
.ost {{ font-family:'Clash Display',sans-serif;font-size:16px;font-weight:600;
  color:{text_color};letter-spacing:-.01em }}
.osa {{ font-size:12px;color:{text_secondary};text-decoration:none;font-weight:700 }}
.osa:hover {{ color:{text_color} }}

/* Table card */
.orc {{ background:{card_bg};border:1.5px solid {card_border};
  border-radius:16px;overflow:hidden;margin-bottom:24px;
  box-shadow:0 2px 8px rgba(0,0,0,.1) }}
.ort {{ width:100%;border-collapse:collapse }}
.ort th {{ padding:13px 20px;background:{card_bg_hover};
  font-size:11px;font-weight:700;color:{text_secondary};
  text-transform:uppercase;letter-spacing:.08em;text-align:left;
  border-bottom:1.5px solid {card_border} }}
.ort td {{ padding:14px 20px;border-bottom:1px solid {card_border};
  font-size:13px;color:{text_color} }}
.ort tr:last-child td {{ border-bottom:none }}
.ort tr:hover td {{ background:{card_bg_hover} }}
.ord {{ font-family:'JetBrains Mono',monospace;font-size:12px;color:{text_secondary} }}
.osb {{ display:flex;align-items:center;gap:10px }}
.osbt {{ flex:1;max-width:80px;height:5px;background:rgba(0,0,0,.15);
  border-radius:3px;overflow:hidden }}
.osbf {{ height:100%;border-radius:3px;background:{score_bar_color};
  transition:width .8s cubic-bezier(.16,1,.3,1) }}
.osbv {{ font-family:'JetBrains Mono',monospace;font-size:12px;
  color:{btn_bg};font-weight:700;min-width:26px }}
.ovb {{ font-size:12px;color:{btn_bg};text-decoration:none;
  font-weight:700;opacity:.8;transition:opacity .15s }}
.ovb:hover {{ opacity:1 }}

/* Extension card */
.oec {{ background:{card_bg};border:1.5px solid {card_border};
  border-radius:16px;padding:32px;display:flex;align-items:center;
  gap:32px;flex-wrap:wrap;box-shadow:0 2px 8px rgba(0,0,0,.1) }}
.oei {{ width:64px;height:64px;border-radius:16px;
  background:rgba(0,0,0,.12);border:1.5px solid {card_border};
  display:flex;align-items:center;justify-content:center;font-size:28px;flex-shrink:0 }}
.oeb {{ flex:1;min-width:200px }}
.oet {{ font-family:'Clash Display',sans-serif;font-size:18px;font-weight:700;
  color:{text_color};margin-bottom:10px }}
.oed {{ font-size:13px;color:{text_secondary};margin-bottom:20px;
  max-width:480px;line-height:1.65 }}
.ocbtn {{ display:inline-flex;align-items:center;gap:8px;
  padding:11px 24px;background:rgba(0,0,0,.15);
  border:1.5px solid rgba(0,0,0,.2);border-radius:10px;
  color:{text_color};font-family:'DM Sans',sans-serif;font-size:14px;
  font-weight:600;cursor:pointer;text-decoration:none;transition:all .15s }}
.ocbtn:hover {{ background:rgba(0,0,0,.22);border-color:rgba(0,0,0,.35) }}

/* Bottom switcher */
.obar {{ position:fixed;bottom:0;left:0;right:0;z-index:300;
  background:{nav_bg};border-top:1px solid {nav_border};
  padding:11px 24px;display:flex;align-items:center;
  justify-content:center;gap:20px;flex-wrap:wrap;
  backdrop-filter:blur(10px) }}
.obar-label {{ font-size:13px;color:{text_secondary};font-weight:700 }}
.obar-links {{ display:flex;gap:6px }}
.obar-link {{ padding:6px 14px;border-radius:8px;font-size:12px;font-weight:600;
  text-decoration:none;border:1.5px solid rgba(0,0,0,.2);
  color:{text_secondary};transition:all .15s;background:transparent }}
.obar-link:hover,.obar-link.cur {{
  background:{btn_bg};color:{btn_text};border-color:{btn_bg} }}
.obar-back {{ font-size:12px;color:{text_tertiary};text-decoration:none;margin-left:8px }}
.obar-back:hover {{ color:{text_secondary} }}

@media(max-width:680px){{
  .osg{{grid-template-columns:1fr}}
  .oh{{flex-direction:column}}
  .oec{{flex-direction:column;gap:20px}}
  .on{{padding:0 16px}}
  .ol{{display:none}}
}}
</style>

<!-- Logo SVG -->
<svg width="0" height="0" style="position:absolute">
  <defs>
    <radialGradient id="op-g" cx="50%" cy="50%" r="50%">
      <stop offset="0%" stop-color="{btn_bg}" stop-opacity=".4"/>
      <stop offset="100%" stop-color="{btn_bg}" stop-opacity="0"/>
    </radialGradient>
  </defs>
</svg>

<nav class="on">
  <a href="#" class="ob">
    <svg width="28" height="28" viewBox="0 0 28 28" fill="none">
      <circle cx="14" cy="14" r="13" fill="url(#op-g)"/>
      <circle cx="14" cy="14" r="12" stroke="rgba(0,0,0,.2)" stroke-width="1.2"/>
      <circle cx="14" cy="14" r="8" stroke="{btn_bg}" stroke-width="1.4" stroke-opacity=".6">
        <animate attributeName="r" values="8;9.5;8" dur="3s" repeatCount="indefinite"/>
        <animate attributeName="stroke-opacity" values=".6;.15;.6" dur="3s" repeatCount="indefinite"/>
      </circle>
      <circle cx="14" cy="14" r="4.5" stroke="{btn_bg}" stroke-width="1.8">
        <animate attributeName="r" values="4.5;5.5;4.5" dur="2.4s" repeatCount="indefinite"/>
        <animate attributeName="stroke-opacity" values="1;.35;1" dur="2.4s" repeatCount="indefinite"/>
      </circle>
      <circle cx="14" cy="14" r="2" fill="{btn_bg}">
        <animate attributeName="opacity" values="1;.5;1" dur="1.8s" repeatCount="indefinite"/>
      </circle>
    </svg>
    ApplyExpress
  </a>
  <div class="ol">
    <a class="olk act" href="#">Dashboard</a>
    <a class="olk" href="#">Runs</a>
    <a class="olk" href="#">History</a>
    <a class="olk" href="#">Settings</a>
  </div>
  <div class="or">
    <span class="ot">14-day trial &middot; 11 days left</span>
    <div class="oa">SC</div>
  </div>
</nav>

<div class="ow">

  <!-- Hero -->
  <div class="oh">
    <div>
      <h1 class="og">Hi, Sarah &#x1F44B;</h1>
      <div class="osub">
        <span class="odot"></span>
        <span class="obadge obadge-green">CV uploaded</span>
        &nbsp;
        <span style="color:{text_tertiary}">Last run: 2 hours ago</span>
      </div>
    </div>
    <div class="orf">
      <select class="ors">
        <option>&#8212; Default CV &#8212;</option>
        <option>AML Compliance</option>
        <option>Risk Manager</option>
      </select>
      <button class="orbtn"
        onclick="window.showToast('Pipeline started &mdash; applying shortly.','success')">
        &#9654;&nbsp; Run now
      </button>
    </div>
  </div>

  <!-- Stats -->
  <div class="osg">
    <div class="osc">
      <span class="osi">&#127919;</span>
      <div class="osv" data-target="142" data-ocount>0</div>
      <div class="osl">Applications sent</div>
      <div class="osd">&#8593; 23 this week</div>
    </div>
    <div class="osc">
      <span class="osi">&#128269;</span>
      <div class="osv" data-target="38" data-ocount>0</div>
      <div class="osl">Pipeline runs</div>
      <div class="osd">&#8593; 4 this week</div>
    </div>
    <div class="osc">
      <span class="osi">&#9889;</span>
      <div class="osv" data-target="9" data-ocount>0</div>
      <div class="osl">Interviews</div>
      <div class="osd">&#8593; 2 this month</div>
    </div>
  </div>

  <!-- Runs table -->
  <div class="osh">
    <span class="ost">Recent runs</span>
    <a class="osa" href="#">View all &rarr;</a>
  </div>
  <div class="orc">
    <table class="ort">
      <thead><tr>
        <th>Date &amp; time</th><th>Status</th>
        <th>Found</th><th>Applied</th><th>Avg score</th><th></th>
      </tr></thead>
      <tbody>
        <tr>
          <td class="ord">2026-04-14&nbsp; 09:12</td>
          <td><span class="obadge obadge-green">Completed</span></td>
          <td style="opacity:.7">24</td><td style="font-weight:700">18</td>
          <td><div class="osb"><div class="osbt"><div class="osbf" style="width:87%"></div></div><span class="osbv">8.7</span></div></td>
          <td><a href="#" class="ovb">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="ord">2026-04-13&nbsp; 21:44</td>
          <td><span class="obadge obadge-green">Completed</span></td>
          <td style="opacity:.7">31</td><td style="font-weight:700">22</td>
          <td><div class="osb"><div class="osbt"><div class="osbf" style="width:74%"></div></div><span class="osbv">7.4</span></div></td>
          <td><a href="#" class="ovb">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="ord">2026-04-12&nbsp; 14:08</td>
          <td><span class="obadge obadge-red">Failed</span></td>
          <td style="opacity:.7">7</td><td style="font-weight:700">0</td>
          <td><div class="osb"><div class="osbt"><div class="osbf" style="width:0%"></div></div><span class="osbv" style="opacity:.4">&#8212;</span></div></td>
          <td><a href="#" class="ovb">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="ord">2026-04-11&nbsp; 09:00</td>
          <td><span class="obadge obadge-green">Completed</span></td>
          <td style="opacity:.7">19</td><td style="font-weight:700">14</td>
          <td><div class="osb"><div class="osbt"><div class="osbf" style="width:81%"></div></div><span class="osbv">8.1</span></div></td>
          <td><a href="#" class="ovb">View &rarr;</a></td>
        </tr>
      </tbody>
    </table>
  </div>

  <!-- Extension card -->
  <div class="oec">
    <div class="oei">&#128279;</div>
    <div class="oeb">
      <h2 class="oet">Connect your Chrome extension</h2>
      <p class="oed">One click to sync with your account. No copying API keys.
        Once connected, press <strong>Start</strong> in the popup to begin applying automatically.</p>
      <a href="#" class="ocbtn"
        onclick="window.showToast('Link copied &mdash; open in Chrome to activate.','info');return false">
        Connect extension &rarr;
      </a>
    </div>
    <div style="text-align:right;flex-shrink:0">
      <div style="font-size:11px;color:{text_tertiary};text-transform:uppercase;
        letter-spacing:.06em;margin-bottom:8px">Extension status</div>
      <span class="obadge obadge-green">&#8226; Connected</span>
    </div>
  </div>

</div>

<!-- Switcher -->
<div class="obar">
  <span class="obar-label">&#127757; Orange background: {scheme_name}</span>
  <div class="obar-links">
    <a href="/dashboard-preview-orange-dark"
      class="obar-link {'cur' if 'Dark' in scheme_name else ''}">&#9679; Dark text</a>
    <a href="/dashboard-preview-orange-green"
      class="obar-link {'cur' if 'Green' in scheme_name else ''}">&#127807; Green text</a>
    <a href="/dashboard-preview-poppy-light"
      class="obar-link">&#8592; Other previews</a>
  </div>
  <a href="/design-system" class="obar-back">&larr; Design system</a>
</div>

<script>
(function(){{
  var counters=document.querySelectorAll('[data-ocount]');
  var started=false;
  function run(){{
    if(started)return;started=true;
    counters.forEach(function(el){{
      var t=parseInt(el.dataset.target,10),steps=40,step=0;
      var iv=setInterval(function(){{
        step++;var p=step/steps;var e=1-Math.pow(1-p,3);
        el.textContent=Math.round(e*t);
        if(step>=steps){{el.textContent=t;clearInterval(iv);}}
      }},900/steps);
    }});
  }}
  if('IntersectionObserver' in window){{
    var obs=new IntersectionObserver(function(e){{if(e[0].isIntersecting)run();}},{{threshold:.3}});
    var g=document.querySelector('.osg');if(g)obs.observe(g);
  }}else{{setTimeout(run,400);}}
  if(window.matchMedia('(prefers-reduced-motion:reduce)').matches)
    counters.forEach(function(el){{el.textContent=el.dataset.target;}});
}})();
</script>"""


@app.route("/dashboard-preview-orange-dark")
def dashboard_preview_orange_dark():
    body = _orange_bg_preview(
        scheme_name          = 'Orange + Dark text',
        text_color           = '#120804',        # deep warm near-black
        text_secondary       = '#5a2a10',        # dark brown-orange
        text_tertiary        = '#8a4a20',        # mid brown
        card_bg              = 'rgba(255,255,255,.18)',   # semi-transparent white cards
        card_border          = 'rgba(255,255,255,.28)',
        card_bg_hover        = 'rgba(255,255,255,.26)',
        nav_bg               = 'rgba(232,98,26,.88)',
        nav_border           = 'rgba(0,0,0,.15)',
        btn_bg               = '#120804',        # near-black button
        btn_text             = '#fdf8f2',        # cream text on button
        score_bar_color      = '#120804',
        switcher_other_url   = '/dashboard-preview-orange-green',
        switcher_other_label = 'Green text',
    )
    return _page("Dashboard — Orange + Dark", body)


@app.route("/dashboard-preview-orange-green")
def dashboard_preview_orange_green():
    body = """
<style>
/* ── Hide _page() chrome ── */
#main-nav, footer { display:none !important }
* { box-sizing:border-box; margin:0; padding:0 }

/* ── Rich orange gradient background ── */
body {
  min-height:100vh;
  font-family:'DM Sans',system-ui,sans-serif;
  background:
    radial-gradient(ellipse at 20% 15%, rgba(255,180,110,.45) 0%, transparent 45%),
    radial-gradient(ellipse at 82% 85%, rgba(200,90,10,.35) 0%, transparent 45%),
    #ff812d;
  color:#0c2718;
}

/* ── Nav — frosted glass on orange ── */
.nav {
  position:fixed;top:0;left:0;right:0;z-index:200;height:62px;
  background:rgba(220,90,22,.55);
  backdrop-filter:blur(22px) saturate(1.8);
  -webkit-backdrop-filter:blur(22px) saturate(1.8);
  border-bottom:1px solid rgba(255,255,255,.22);
  display:flex;align-items:center;justify-content:space-between;padding:0 36px;
}
.nav-brand {
  display:flex;align-items:center;gap:11px;text-decoration:none;
  font-family:'Clash Display',sans-serif;font-size:19px;font-weight:700;
  color:#0c2718;letter-spacing:-.02em;
}
.nav-links { display:flex;gap:2px }
.nav-link {
  padding:7px 15px;border-radius:9px;font-size:13px;font-weight:500;
  color:rgba(12,39,24,.65);text-decoration:none;transition:all .15s;
}
.nav-link:hover { color:#0c2718;background:rgba(255,255,255,.18) }
.nav-link.on {
  color:#0c2718;background:rgba(255,255,255,.28);font-weight:700;
  box-shadow:inset 0 1px 0 rgba(255,255,255,.5);
}
.nav-right { display:flex;align-items:center;gap:14px }
.nav-trial { font-size:12px;color:rgba(12,39,24,.6);font-family:'JetBrains Mono',monospace }
.nav-avatar {
  width:35px;height:35px;border-radius:50%;
  background:linear-gradient(135deg,#0c2718,#1a5c34);
  display:flex;align-items:center;justify-content:center;
  font-size:12px;font-weight:800;color:#fdf8f0;
  border:2px solid rgba(255,255,255,.35);
  box-shadow:0 2px 8px rgba(0,0,0,.2);
}

/* ── Wrapper ── */
.wrap {
  max-width:1120px;margin:0 auto;padding:86px 28px 82px;
  animation:rise .45s cubic-bezier(.16,1,.3,1) both;
}
@keyframes rise { from{opacity:0;transform:translateY(18px)} to{opacity:1;transform:none} }

/* ── Hero row ── */
.hero {
  display:flex;align-items:flex-start;justify-content:space-between;
  flex-wrap:wrap;gap:22px;margin-bottom:34px;
}
.greeting {
  font-family:'Clash Display',sans-serif;font-size:38px;font-weight:700;
  letter-spacing:-.025em;color:#0c2718;margin-bottom:10px;
  text-shadow:0 1px 2px rgba(255,255,255,.15);
}
.subline { display:flex;align-items:center;gap:10px;font-size:13px;color:rgba(12,39,24,.65) }
.pulse-dot {
  width:8px;height:8px;border-radius:50%;flex-shrink:0;
  background:#0c2718;
  box-shadow:0 0 0 3px rgba(12,39,24,.2);
  animation:pd 2.5s ease infinite;
}
@keyframes pd {
  0%,100%{box-shadow:0 0 0 3px rgba(12,39,24,.2)}
  50%{box-shadow:0 0 0 8px rgba(12,39,24,.06)}
}
.chip {
  display:inline-flex;align-items:center;gap:5px;
  padding:4px 12px;border-radius:20px;font-size:11px;font-weight:700;
  letter-spacing:.04em;
}
.chip-green {
  background:rgba(12,39,24,.18);color:#0c2718;
  border:1px solid rgba(12,39,24,.2);
  backdrop-filter:blur(6px);
}

/* ── Run form ── */
.run-row { display:flex;gap:10px;align-items:center;flex-wrap:wrap }
.run-select {
  padding:11px 18px;
  background:rgba(255,255,255,.22);
  border:1.5px solid rgba(255,255,255,.38);
  border-radius:11px;color:#0c2718;
  font-family:'DM Sans',sans-serif;font-size:13px;font-weight:500;
  cursor:pointer;transition:all .15s;
  backdrop-filter:blur(10px);
  box-shadow:inset 0 1px 0 rgba(255,255,255,.4);
}
.run-select:focus { outline:none;border-color:rgba(12,39,24,.5) }
.run-btn {
  display:inline-flex;align-items:center;gap:9px;
  padding:12px 26px;
  background:linear-gradient(145deg,#1a5c34,#0c2718);
  color:#fdf8f0;border:none;border-radius:11px;
  font-family:'DM Sans',sans-serif;font-size:14px;font-weight:700;
  cursor:pointer;transition:all .15s ease;letter-spacing:.01em;
  box-shadow:0 4px 16px rgba(12,39,24,.4), inset 0 1px 0 rgba(255,255,255,.12);
}
.run-btn:hover {
  background:linear-gradient(145deg,#1f6e3e,#12301e);
  transform:translateY(-1px);
  box-shadow:0 8px 24px rgba(12,39,24,.45), inset 0 1px 0 rgba(255,255,255,.15);
}
.run-btn:active { transform:translateY(0) scale(.98) }

/* ── Glass card base ── */
.glass {
  background:rgba(255,255,255,.22);
  backdrop-filter:blur(24px) saturate(1.6);
  -webkit-backdrop-filter:blur(24px) saturate(1.6);
  border:1.5px solid rgba(255,255,255,.42);
  border-radius:18px;
  box-shadow:
    0 8px 32px rgba(0,0,0,.1),
    inset 0 1.5px 0 rgba(255,255,255,.55),
    inset 0 -1px 0 rgba(0,0,0,.04);
  transition:all .2s ease;
}
.glass:hover {
  background:rgba(255,255,255,.3);
  box-shadow:
    0 14px 48px rgba(0,0,0,.14),
    inset 0 1.5px 0 rgba(255,255,255,.65),
    inset 0 -1px 0 rgba(0,0,0,.04);
  transform:translateY(-2px);
}

/* ── Stat grid ── */
.stat-grid { display:grid;grid-template-columns:repeat(3,1fr);gap:18px;margin-bottom:30px }
.stat-card {
  padding:28px 26px 22px;
}
.stat-icon { font-size:22px;display:block;margin-bottom:16px }
.stat-num {
  font-family:'Clash Display',sans-serif;font-size:48px;font-weight:700;
  color:#0c2718;letter-spacing:-.03em;line-height:1;margin-bottom:6px;
}
.stat-label {
  font-size:11px;color:rgba(12,39,24,.55);
  text-transform:uppercase;letter-spacing:.09em;font-weight:600;
}
.stat-delta {
  display:inline-flex;align-items:center;gap:4px;
  margin-top:14px;font-size:12px;font-weight:700;
  padding:3px 10px;border-radius:20px;
  background:rgba(12,39,24,.14);color:#0c2718;
}

/* ── Section header ── */
.section-row {
  display:flex;align-items:center;justify-content:space-between;
  margin-bottom:14px;
}
.section-title {
  font-family:'Clash Display',sans-serif;font-size:17px;font-weight:700;
  color:#0c2718;letter-spacing:-.01em;
}
.section-link {
  font-size:12px;color:rgba(12,39,24,.7);font-weight:700;
  text-decoration:none;letter-spacing:.02em;
}
.section-link:hover { color:#0c2718 }

/* ── Runs table ── */
.table-card { padding:0;overflow:hidden;margin-bottom:26px }
.runs-table { width:100%;border-collapse:collapse }
.runs-table th {
  padding:14px 22px;
  background:rgba(255,255,255,.15);
  font-size:11px;font-weight:700;color:rgba(12,39,24,.55);
  text-transform:uppercase;letter-spacing:.09em;text-align:left;
  border-bottom:1.5px solid rgba(255,255,255,.3);
}
.runs-table td {
  padding:15px 22px;border-bottom:1px solid rgba(255,255,255,.18);
  font-size:13px;color:#0c2718;
}
.runs-table tr:last-child td { border-bottom:none }
.runs-table tbody tr:hover td { background:rgba(255,255,255,.15) }
.run-date { font-family:'JetBrains Mono',monospace;font-size:12px;color:rgba(12,39,24,.55) }

.status-chip {
  display:inline-block;padding:4px 11px;border-radius:20px;
  font-size:11px;font-weight:700;letter-spacing:.04em;
}
.status-ok { background:rgba(12,39,24,.18);color:#0c2718;border:1px solid rgba(12,39,24,.2) }
.status-fail { background:rgba(180,30,30,.18);color:#7a1010;border:1px solid rgba(180,30,30,.2) }

.score-wrap { display:flex;align-items:center;gap:10px }
.score-track {
  flex:1;max-width:80px;height:5px;
  background:rgba(12,39,24,.12);border-radius:3px;overflow:hidden;
}
.score-fill {
  height:100%;border-radius:3px;
  background:linear-gradient(90deg,#1a5c34,#0c2718);
  transition:width .9s cubic-bezier(.16,1,.3,1);
}
.score-val {
  font-family:'JetBrains Mono',monospace;font-size:12px;
  font-weight:700;color:#0c2718;min-width:26px;
}
.view-link {
  font-size:12px;font-weight:700;color:rgba(12,39,24,.7);
  text-decoration:none;letter-spacing:.02em;transition:color .12s;
}
.view-link:hover { color:#0c2718 }

/* ── Extension card ── */
.ext-card { padding:34px;display:flex;align-items:center;gap:32px;flex-wrap:wrap }
.ext-icon {
  width:68px;height:68px;border-radius:18px;
  background:rgba(12,39,24,.14);
  border:1.5px solid rgba(12,39,24,.18);
  display:flex;align-items:center;justify-content:center;
  font-size:30px;flex-shrink:0;
  box-shadow:inset 0 1px 0 rgba(255,255,255,.3);
}
.ext-body { flex:1;min-width:220px }
.ext-title {
  font-family:'Clash Display',sans-serif;font-size:20px;font-weight:700;
  color:#0c2718;margin-bottom:10px;letter-spacing:-.01em;
}
.ext-desc { font-size:13px;color:rgba(12,39,24,.65);margin-bottom:22px;
  max-width:480px;line-height:1.7 }
.ext-btn {
  display:inline-flex;align-items:center;gap:8px;
  padding:12px 26px;
  background:rgba(255,255,255,.28);
  border:1.5px solid rgba(255,255,255,.5);
  border-radius:11px;color:#0c2718;
  font-family:'DM Sans',sans-serif;font-size:14px;font-weight:700;
  cursor:pointer;text-decoration:none;transition:all .15s;
  box-shadow:inset 0 1px 0 rgba(255,255,255,.55);
}
.ext-btn:hover {
  background:rgba(255,255,255,.38);
  border-color:rgba(255,255,255,.65);
  transform:translateY(-1px);
}
.ext-status { text-align:right;flex-shrink:0 }
.ext-status-label {
  font-size:11px;color:rgba(12,39,24,.5);text-transform:uppercase;
  letter-spacing:.07em;margin-bottom:8px;
}
.connected-badge {
  display:inline-flex;align-items:center;gap:6px;
  padding:6px 14px;border-radius:20px;font-size:12px;font-weight:700;
  background:rgba(12,39,24,.18);color:#0c2718;
  border:1px solid rgba(12,39,24,.2);
}
.connected-dot {
  width:7px;height:7px;border-radius:50%;background:#0c2718;
  animation:pd 2.5s ease infinite;
}

/* ── Bottom banner ── */
.bottom-bar {
  position:fixed;bottom:0;left:0;right:0;z-index:300;
  background:rgba(220,90,22,.6);
  backdrop-filter:blur(20px);-webkit-backdrop-filter:blur(20px);
  border-top:1px solid rgba(255,255,255,.22);
  padding:12px 28px;display:flex;align-items:center;
  justify-content:center;gap:20px;flex-wrap:wrap;
}
.bb-label { font-size:13px;color:rgba(12,39,24,.8);font-weight:700 }
.bb-links { display:flex;gap:8px }
.bb-link {
  padding:6px 16px;border-radius:9px;font-size:12px;font-weight:700;
  text-decoration:none;border:1.5px solid rgba(255,255,255,.3);
  color:rgba(12,39,24,.75);transition:all .15s;background:rgba(255,255,255,.12);
}
.bb-link:hover,.bb-link.cur {
  background:linear-gradient(135deg,#1a5c34,#0c2718);
  color:#fdf8f0;border-color:transparent;
  box-shadow:0 2px 10px rgba(12,39,24,.3);
}
.bb-back { font-size:12px;color:rgba(12,39,24,.5);text-decoration:none;margin-left:8px }
.bb-back:hover { color:rgba(12,39,24,.75) }

@media(max-width:700px){
  .stat-grid{grid-template-columns:1fr}
  .hero{flex-direction:column}
  .ext-card{flex-direction:column;gap:20px;padding:24px}
  .nav{padding:0 16px}
  .nav-links{display:none}
  .wrap{padding:80px 16px 76px}
}
</style>

<!-- ── Nav ── -->
<nav class="nav">
  <a href="#" class="nav-brand">
    <!-- Logo B — pulsing rings in green+orange -->
    <svg width="30" height="30" viewBox="0 0 30 30" fill="none">
      <defs>
        <radialGradient id="og-gl" cx="50%" cy="50%" r="50%">
          <stop offset="0%" stop-color="#1a5c34" stop-opacity=".35"/>
          <stop offset="100%" stop-color="#0c2718" stop-opacity="0"/>
        </radialGradient>
        <linearGradient id="og-ring" x1="0%" y1="0%" x2="100%" y2="100%">
          <stop offset="0%" stop-color="#1a5c34"/>
          <stop offset="100%" stop-color="#0c2718"/>
        </linearGradient>
      </defs>
      <circle cx="15" cy="15" r="14" fill="url(#og-gl)"/>
      <circle cx="15" cy="15" r="13" stroke="rgba(12,39,24,.25)" stroke-width="1"/>
      <!-- outer ring -->
      <circle cx="15" cy="15" r="9.5" stroke="#1a5c34" stroke-width="1.3" stroke-opacity=".5">
        <animate attributeName="r" values="9.5;11;9.5" dur="3s" repeatCount="indefinite"/>
        <animate attributeName="stroke-opacity" values=".5;.12;.5" dur="3s" repeatCount="indefinite"/>
      </circle>
      <!-- mid ring -->
      <circle cx="15" cy="15" r="6" stroke="url(#og-ring)" stroke-width="1.6">
        <animate attributeName="r" values="6;7.2;6" dur="2.4s" repeatCount="indefinite"/>
        <animate attributeName="stroke-opacity" values="1;.35;1" dur="2.4s" repeatCount="indefinite"/>
      </circle>
      <!-- centre -->
      <circle cx="15" cy="15" r="2.5" fill="url(#og-ring)">
        <animate attributeName="opacity" values="1;.45;1" dur="1.8s" repeatCount="indefinite"/>
      </circle>
    </svg>
    ApplyExpress
  </a>
  <div class="nav-links">
    <a class="nav-link on" href="#">Dashboard</a>
    <a class="nav-link" href="#">Runs</a>
    <a class="nav-link" href="#">History</a>
    <a class="nav-link" href="#">Settings</a>
  </div>
  <div class="nav-right">
    <span class="nav-trial">14-day trial &middot; 11 days left</span>
    <div class="nav-avatar">SC</div>
  </div>
</nav>

<!-- ── Dashboard body ── -->
<div class="wrap">

  <!-- Hero -->
  <div class="hero">
    <div>
      <h1 class="greeting">Hi, Sarah &#x1F44B;</h1>
      <div class="subline">
        <span class="pulse-dot"></span>
        <span class="chip chip-green">&#10003; CV uploaded</span>
        <span style="color:rgba(12,39,24,.5)">Last run: 2 hours ago</span>
      </div>
    </div>
    <div class="run-row">
      <select class="run-select">
        <option>&#8212; Default CV &#8212;</option>
        <option>AML Compliance</option>
        <option>Risk Manager</option>
      </select>
      <button class="run-btn"
        onclick="window.showToast('Pipeline started &mdash; the extension will begin applying shortly.','success')">
        &#9654;&nbsp; Run now
      </button>
    </div>
  </div>

  <!-- Stats -->
  <div class="stat-grid">
    <div class="glass stat-card">
      <span class="stat-icon">&#127919;</span>
      <div class="stat-num" data-target="142" data-gcount>0</div>
      <div class="stat-label">Applications sent</div>
      <div class="stat-delta">&#8593; 23 this week</div>
    </div>
    <div class="glass stat-card">
      <span class="stat-icon">&#128269;</span>
      <div class="stat-num" data-target="38" data-gcount>0</div>
      <div class="stat-label">Pipeline runs</div>
      <div class="stat-delta">&#8593; 4 this week</div>
    </div>
    <div class="glass stat-card">
      <span class="stat-icon">&#9889;</span>
      <div class="stat-num" data-target="9" data-gcount>0</div>
      <div class="stat-label">Interviews booked</div>
      <div class="stat-delta">&#8593; 2 this month</div>
    </div>
  </div>

  <!-- Runs table -->
  <div class="section-row">
    <span class="section-title">Recent runs</span>
    <a class="section-link" href="#">View all &rarr;</a>
  </div>
  <div class="glass table-card">
    <table class="runs-table">
      <thead><tr>
        <th>Date &amp; time</th><th>Status</th>
        <th>Found</th><th>Applied</th><th>Avg score</th><th></th>
      </tr></thead>
      <tbody>
        <tr>
          <td class="run-date">2026-04-14 &nbsp;09:12</td>
          <td><span class="status-chip status-ok">Completed</span></td>
          <td style="color:rgba(12,39,24,.55)">24</td>
          <td style="font-weight:800">18</td>
          <td><div class="score-wrap"><div class="score-track"><div class="score-fill" style="width:87%"></div></div><span class="score-val">8.7</span></div></td>
          <td><a href="#" class="view-link">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="run-date">2026-04-13 &nbsp;21:44</td>
          <td><span class="status-chip status-ok">Completed</span></td>
          <td style="color:rgba(12,39,24,.55)">31</td>
          <td style="font-weight:800">22</td>
          <td><div class="score-wrap"><div class="score-track"><div class="score-fill" style="width:74%"></div></div><span class="score-val">7.4</span></div></td>
          <td><a href="#" class="view-link">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="run-date">2026-04-12 &nbsp;14:08</td>
          <td><span class="status-chip status-fail">Failed</span></td>
          <td style="color:rgba(12,39,24,.55)">7</td>
          <td style="font-weight:800">0</td>
          <td><div class="score-wrap"><div class="score-track"><div class="score-fill" style="width:0%"></div></div><span class="score-val" style="opacity:.35">&#8212;</span></div></td>
          <td><a href="#" class="view-link">View &rarr;</a></td>
        </tr>
        <tr>
          <td class="run-date">2026-04-11 &nbsp;09:00</td>
          <td><span class="status-chip status-ok">Completed</span></td>
          <td style="color:rgba(12,39,24,.55)">19</td>
          <td style="font-weight:800">14</td>
          <td><div class="score-wrap"><div class="score-track"><div class="score-fill" style="width:81%"></div></div><span class="score-val">8.1</span></div></td>
          <td><a href="#" class="view-link">View &rarr;</a></td>
        </tr>
      </tbody>
    </table>
  </div>

  <!-- Extension card -->
  <div class="glass ext-card">
    <div class="ext-icon">&#128279;</div>
    <div class="ext-body">
      <h2 class="ext-title">Connect your Chrome extension</h2>
      <p class="ext-desc">One click to sync the extension with your account. No copying API keys.
        Once connected, press <strong style="color:#0c2718">Start</strong>
        in the popup to begin submitting applications automatically.</p>
      <a href="#" class="ext-btn"
        onclick="window.showToast('Connection link copied &mdash; open in Chrome to activate.','info');return false">
        Connect extension &rarr;
      </a>
    </div>
    <div class="ext-status">
      <div class="ext-status-label">Extension status</div>
      <div class="connected-badge">
        <span class="connected-dot"></span> Connected
      </div>
    </div>
  </div>

</div>

<!-- ── Bottom bar ── -->
<div class="bottom-bar">
  <span class="bb-label">&#127801; Orange + Forest Green &mdash; Rich preview</span>
  <div class="bb-links">
    <a href="/dashboard-preview-orange-dark" class="bb-link">Dark text version</a>
    <a href="/dashboard-preview-poppy-light" class="bb-link">Other schemes</a>
    <a href="/dashboard-preview-gold" class="bb-link">Dark schemes</a>
  </div>
  <a href="/design-system" class="bb-back">&larr; Design system</a>
</div>

<script>
(function(){
  var counters = document.querySelectorAll('[data-gcount]');
  var started = false;
  function run(){
    if(started) return; started = true;
    counters.forEach(function(el){
      var t = parseInt(el.dataset.target, 10), steps = 45, step = 0;
      var iv = setInterval(function(){
        step++;
        var e = 1 - Math.pow(1 - step/steps, 3);
        el.textContent = Math.round(e * t);
        if(step >= steps){ el.textContent = t; clearInterval(iv); }
      }, 950/steps);
    });
  }
  if('IntersectionObserver' in window){
    var obs = new IntersectionObserver(function(e){ if(e[0].isIntersecting) run(); },{threshold:.25});
    var g = document.querySelector('.stat-grid'); if(g) obs.observe(g);
  } else { setTimeout(run, 500); }
  if(window.matchMedia('(prefers-reduced-motion: reduce)').matches)
    counters.forEach(function(el){ el.textContent = el.dataset.target; });
})();
</script>"""
    return _page("Dashboard — Orange + Forest Green", body)


@app.route("/privacy")
def privacy_policy():
    return """<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Privacy Policy — ApplyExpress</title>
<style>
  body{font-family:system-ui,sans-serif;max-width:720px;margin:40px auto;padding:0 24px;color:#1e293b;line-height:1.7}
  h1{color:#1d4ed8}h2{margin-top:2em;color:#1e40af}
  a{color:#1d4ed8}
</style></head><body>
<h1>Privacy Policy</h1>
<p><strong>Last updated: April 2026</strong></p>

<h2>1. Overview</h2>
<p>ApplyExpress ("we", "us") provides a Chrome extension and web platform that helps users auto-fill
job applications. This policy explains what data we collect, why, and how it is used.</p>

<h2>2. Data We Collect</h2>
<ul>
  <li><strong>Account data</strong>: email address and hashed password, stored securely in our database.</li>
  <li><strong>Profile data</strong>: name, phone number, and job preferences you enter in your profile.</li>
  <li><strong>CV / resume</strong>: the document you upload is stored on our server solely to auto-fill applications on your behalf.</li>
  <li><strong>Application activity</strong>: a log of jobs applied to, scores, and pipeline results, visible only to you.</li>
</ul>

<h2>3. Data We Do NOT Collect</h2>
<ul>
  <li>We do not track browsing history.</li>
  <li>We do not sell, share, or transfer your data to third parties.</li>
  <li>We do not use your data for advertising or creditworthiness purposes.</li>
  <li>The Chrome extension stores only your server URL and API key in <code>chrome.storage.local</code> on your device.</li>
</ul>

<h2>4. How Data Is Used</h2>
<p>Your data is used exclusively to operate the ApplyExpress service: filling in your name, email,
phone, and CV on job application forms you choose to apply to.</p>

<h2>5. Data Retention</h2>
<p>Your data is retained while your account is active. You may delete your account and all associated
data at any time by contacting us.</p>

<h2>6. Security</h2>
<p>Passwords are hashed before storage. All connections use HTTPS. CV files are stored in a
private directory not accessible to other users.</p>

<h2>7. Contact</h2>
<p>For questions about this policy, email:
<a href="mailto:support@applyexpress.io">support@applyexpress.io</a></p>
</body></html>"""


@app.route("/api/read_otp", methods=["GET"])
@api_key_required
def api_read_otp():
    """Poll the user's Gmail inbox for a recent OTP from a named sender (reed, linkedin, indeed…).
    Returns {"code": "123456"} or {"code": null} if not found yet.
    Extension polls this endpoint up to 12× at 5s intervals (60s window).
    """
    u            = g.api_user
    from_pattern = request.args.get("from", "").lower()   # e.g. "reed", "linkedin"
    smtp_pass    = u.get("smtp_password") or ""
    if not smtp_pass:
        return jsonify({"error": "No Gmail app-password configured in profile"}), 400
    if not from_pattern:
        return jsonify({"error": "Missing ?from= parameter"}), 400

    try:
        M = imaplib.IMAP4_SSL("imap.gmail.com", 993)
        M.login(u["email"], smtp_pass)
        M.select("INBOX")

        # Search emails in last 5 minutes
        since_dt   = datetime.utcnow() - timedelta(minutes=5)
        since_str  = since_dt.strftime("%d-%b-%Y")
        _, ids     = M.search(None, f'(SINCE {since_str})')
        code       = None

        for num in reversed(ids[0].split()):
            _, data = M.fetch(num, "(RFC822)")
            msg     = _email.message_from_bytes(data[0][1])
            from_hdr = (msg.get("From") or "").lower()
            if from_pattern not in from_hdr:
                continue

            # Extract plain text body
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    ct = part.get_content_type()
                    if ct == "text/plain":
                        try:
                            body += part.get_payload(decode=True).decode("utf-8", errors="ignore")
                        except Exception:
                            pass
            else:
                try:
                    body = msg.get_payload(decode=True).decode("utf-8", errors="ignore")
                except Exception:
                    pass

            # Decode quoted-printable
            body = re.sub(r'=\r?\n', '', body)
            body = re.sub(r'=([0-9A-Fa-f]{2})', lambda m: chr(int(m.group(1), 16)), body)

            # Extract 4–8 digit OTP
            m = (re.search(r'\b(\d{6})\b', body)
                 or re.search(r'(?:code|OTP|verification)[^\d]*(\d{4,8})', body, re.I)
                 or re.search(r'(\d{4,8})', body))
            if m:
                code = m.group(1)
                log.info(f"[OTP] Found code for '{from_pattern}': {code}")
                break

        M.logout()
        return jsonify({"code": code})

    except imaplib.IMAP4.error as e:
        log.warning(f"[OTP] IMAP error: {e}")
        return jsonify({"error": "IMAP error: " + str(e)}), 500
    except Exception as e:
        log.warning(f"[OTP] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/log", methods=["POST"])
@api_key_required
def api_log():
    data = request.get_json(force=True) or {}
    log.info(f"[EXT] {data.get('level','info').upper()}: {data.get('message','')}")
    return jsonify({"ok": True})


@app.route("/api/trigger_now", methods=["POST"])
@api_key_required
def api_trigger_now():
    """Fire the pipeline immediately for the authenticated user — runs default + all specialties.
    Can be called from a server-side cron job:
      curl -X POST http://localhost:8080/api/trigger_now -H "X-API-Key: <key>"
    Optional body: {"specialty_id": 3} to run a single specialty only.
    """
    u  = g.api_user
    db = get_db()
    import subprocess as _sp2

    data        = request.get_json(force=True, silent=True) or {}
    only_spec   = data.get("specialty_id") or None   # if set, run only this specialty
    worker      = Path(__file__).parent / "tools" / "run_pipeline_worker.py"
    started     = []

    seq_runner  = Path(__file__).parent / "tools" / "run_sequential_runner.py"

    def _create_run(spec_id=None):
        """Create a DB run record and return its id, or None if already running."""
        if spec_id:
            already = db.execute(
                "SELECT id FROM runs WHERE user_id=? AND specialty_id=? AND status='running'",
                (u["id"], spec_id)
            ).fetchone()
        else:
            already = db.execute(
                "SELECT id FROM runs WHERE user_id=? AND specialty_id IS NULL AND status='running'",
                (u["id"],)
            ).fetchone()
        if already:
            return None
        run_id = datetime.utcnow().strftime("%Y%m%d_%H%M%S") + (f"_s{spec_id}" if spec_id else "")
        row = db.execute(
            "INSERT INTO runs (user_id, run_id, started_at, status, specialty_id) VALUES (?,?,?,?,?) RETURNING id",
            (u["id"], run_id, datetime.utcnow().isoformat(), "running", spec_id)
        ).fetchone()
        db.commit()
        return row["id"]

    if only_spec:
        rid = _create_run(spec_id=only_spec)
        if rid is None:
            return jsonify({"ok": False, "reason": "that specialty is already running"}), 409
        started.append(rid)
        _sp2.Popen(
            [sys.executable, str(seq_runner), str(u["id"]), str(DB_PATH), str(ROOT), str(DATA_DIR),
             f"{rid}:{only_spec}"],
            start_new_session=True, stdout=_sp2.DEVNULL, stderr=_sp2.DEVNULL,
        )
    else:
        # Create all run records upfront, then run sequentially in one subprocess
        pairs = []
        for spec_id in [None] + [s["id"] for s in db.execute(
                "SELECT id FROM specialties WHERE user_id=?", (u["id"],)).fetchall()]:
            rid = _create_run(spec_id=spec_id)
            if rid:
                started.append(rid)
                pairs.append(f"{rid}:{spec_id or ''}")
        if pairs:
            _sp2.Popen(
                [sys.executable, str(seq_runner), str(u["id"]), str(DB_PATH), str(ROOT), str(DATA_DIR)]
                + pairs,
                start_new_session=True, stdout=_sp2.DEVNULL, stderr=_sp2.DEVNULL,
            )
            log.info(f"[trigger_now] Sequential runner started for {len(pairs)} run(s): {pairs}")

    if not started:
        return jsonify({"ok": False, "reason": "all pipelines already running"}), 409
    return jsonify({"ok": True, "runs_started": len(started), "run_ids": started})


@app.route("/api/save_session", methods=["POST"])
@app.route("/profile/save-session", methods=["POST"])
@api_key_required
def api_save_session():
    """Save session cookies for a job platform (reed, indeed, linkedin).
    Accepts JSON body {"platform": "reed", "cookies": [...]} or FormData with same fields.
    Called automatically by the Chrome extension every 15 minutes.
    """
    u = g.api_user
    # Accept both JSON body (manual API) and FormData (Chrome extension)
    if request.content_type and "application/json" in request.content_type:
        data = request.get_json(force=True) or {}
        platform = data.get("platform", "").strip().lower()
        cookies_raw = data.get("cookies", [])
    else:
        platform = (request.form.get("platform") or "").strip().lower()
        cookies_raw = request.form.get("cookies", "[]")

    import json as _json
    if isinstance(cookies_raw, str):
        try:
            cookies = _json.loads(cookies_raw)
        except Exception:
            return jsonify({"ok": False, "reason": "cookies must be valid JSON"}), 400
    else:
        cookies = cookies_raw

    valid_platforms = {"reed", "indeed", "linkedin"}
    if platform not in valid_platforms:
        return jsonify({"ok": False, "reason": f"platform must be one of: {', '.join(valid_platforms)}"}), 400
    if not isinstance(cookies, list) or len(cookies) == 0:
        return jsonify({"ok": False, "reason": "cookies must be a non-empty array"}), 400
    sessions_dir = Path(DATA_DIR) / "users" / str(u["id"]) / "sessions"
    sessions_dir.mkdir(parents=True, exist_ok=True)
    session_file = sessions_dir / f"{platform}.json"
    session_file.write_text(_json.dumps(cookies, indent=2))
    log.info(f"[save_session] Saved {len(cookies)} {platform} cookies for user {u['id']}")
    return jsonify({"ok": True, "platform": platform, "cookies_saved": len(cookies), "count": len(cookies)})


# ── Server-side session refresh (Reed) ────────────────────────────────────────
# Lets users refresh an expired Reed session from any device (incl. phone) —
# runs a Playwright login on the server with their stored credentials.
# LinkedIn/Indeed still need the extension: LinkedIn blocks datacenter logins,
# Indeed requires Google OAuth + OTP.

def _session_status_file(user_id: int, platform: str) -> Path:
    d = Path(DATA_DIR) / "users" / str(user_id) / "sessions"
    d.mkdir(parents=True, exist_ok=True)
    return d / f"refresh_{platform}_status.json"

def _run_reed_refresh(user_id: int, email: str, password: str):
    import subprocess
    status_file = _session_status_file(user_id, "reed")
    udir = Path(DATA_DIR) / "users" / str(user_id)
    (udir / ".tmp").mkdir(parents=True, exist_ok=True)
    env = dict(os.environ)
    env.update({"TMP_DIR": str(udir / ".tmp"),
                "REED_EMAIL": email, "REED_PASSWORD": password})

    def _write(state, detail=""):
        status_file.write_text(json.dumps(
            {"state": state, "detail": detail[:400],
             "ts": datetime.utcnow().isoformat()}))

    _write("running")
    try:
        r = subprocess.run(["node", str(ROOT / "tools" / "reed_login_once.js")],
                           capture_output=True, text=True, timeout=180,
                           cwd=str(ROOT), env=env)
        tail = (r.stdout or "").strip()[-300:] or (r.stderr or "").strip()[-300:]
        if r.returncode == 0:
            _write("ok", "Session refreshed.")
            log.info(f"[refresh_session] Reed session refreshed for user {user_id}")
        elif r.returncode == 3:
            _write("captcha", "Reed showed a CAPTCHA. Try again in a few hours, "
                              "or refresh via the desktop Chrome extension.")
        else:
            _write("failed", tail)
            log.warning(f"[refresh_session] Reed refresh failed for user {user_id}: {tail[-150:]}")
    except subprocess.TimeoutExpired:
        _write("failed", "Login timed out after 3 minutes.")
    except Exception as e:
        _write("failed", str(e))

@app.route("/profile/refresh-session", methods=["POST"])
@login_required
def refresh_session():
    u = current_user()
    platform = request.values.get("platform", "")
    if platform != "reed":
        return jsonify({"ok": False, "reason": "Server-side refresh is available for Reed only. "
                        "Use the Chrome extension for LinkedIn and Indeed."}), 400
    password = _dec(u.get("reed_pass", ""))
    if not password:
        return jsonify({"ok": False, "reason": "No Reed password saved — add it under "
                        "'Email & job board credentials' first."}), 400
    sf = _session_status_file(u["id"], "reed")
    if sf.exists():
        try:
            st = json.loads(sf.read_text())
            cutoff = (datetime.utcnow() - timedelta(minutes=4)).isoformat()
            if st.get("state") == "running" and st.get("ts", "") > cutoff:
                return jsonify({"ok": True, "state": "running"})
        except Exception:
            pass
    email = u.get("reed_email") or u["email"]
    threading.Thread(target=_run_reed_refresh, args=(u["id"], email, password),
                     daemon=True).start()
    return jsonify({"ok": True, "state": "running"})

@app.route("/profile/refresh-session-status")
@login_required
def refresh_session_status():
    u = current_user()
    platform = request.args.get("platform", "reed")
    if platform not in ("reed", "linkedin", "indeed"):
        return jsonify({"error": "bad platform"}), 400
    out = {"state": "idle", "detail": ""}
    sf = _session_status_file(u["id"], platform)
    if sf.exists():
        try:
            out = json.loads(sf.read_text())
        except Exception:
            pass
    p = Path(DATA_DIR) / "users" / str(u["id"]) / "sessions" / f"{platform}.json"
    if p.exists():
        out["session_age_days"] = round((time.time() - p.stat().st_mtime) / 86400, 1)
    return jsonify(out)


@app.route("/api/email_test", methods=["POST"])
@api_key_required
def api_email_test():
    """Send a test email to verify SMTP configuration. Call with your API key."""
    u = g.api_user
    smtp_pwd = _dec(u.get("smtp_password", ""))
    if not smtp_pwd:
        return jsonify({"ok": False, "reason": "No Gmail app password set — go to Profile → Settings and add your Gmail App Password"}), 400
    try:
        _send_email(
            u["email"],
            "ApplyExpress: email test ✓",
            "<p>Your email notifications are working correctly. "
            "You will receive a report after each pipeline run.</p>",
            smtp_password=smtp_pwd,
        )
        return jsonify({"ok": True, "sent_to": u["email"]})
    except Exception as e:
        return jsonify({"ok": False, "reason": str(e)}), 500


# ── Unsubscribe ───────────────────────────────────────────────────────────────

@app.route("/unsubscribe")
def unsubscribe():
    """One-click email opt-out. Link format: /unsubscribe?token=<api_key>"""
    token = request.args.get("token", "")
    db = get_db()
    u = db.execute("SELECT id, email, first_name FROM users WHERE api_key=?", (token,)).fetchone()
    if not u:
        return _page("Unsubscribe", "<p style='color:#dc2626'>Invalid or expired unsubscribe link.</p>")

    # Set smtp_password to sentinel so pipeline completion email is skipped
    db.execute("UPDATE users SET smtp_password='__unsubscribed__' WHERE id=?", (u["id"],))
    db.commit()
    return _page("Unsubscribed", f"""
    <div style="max-width:480px;margin:80px auto;font-family:sans-serif;text-align:center">
      <h2 style="margin-bottom:12px">You've been unsubscribed</h2>
      <p style="color:#6b7280">Hi {u['first_name'] or u['email']}, you'll no longer receive
      pipeline completion emails from ApplyExpress.</p>
      <p style="margin-top:24px"><a href="/profile" style="color:#4f46e5">Manage preferences →</a></p>
    </div>""")


# ── Admin panel ───────────────────────────────────────────────────────────────

ADMIN_EMAIL = os.getenv("ADMIN_EMAIL", "")

def _admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            return redirect(url_for("login_page"))
        db = get_db()
        u = db.execute("SELECT email FROM users WHERE id=?", (session["user_id"],)).fetchone()
        if not u or u["email"] != ADMIN_EMAIL:
            return _page("Forbidden", "<p style='color:#dc2626;text-align:center'>Admin only.</p>"), 403
        return f(*args, **kwargs)
    return decorated


@app.route("/admin")
@_admin_required
def admin_panel():
    db = get_db()

    users = db.execute("""
        SELECT u.id, u.email, u.first_name, u.last_name, u.created_at,
               u.is_paid, u.trial_ends_at, u.keywords,
               COUNT(DISTINCT r.id)  AS total_runs,
               COALESCE(SUM(r.jobs_applied),0) AS total_applied,
               MAX(r.started_at) AS last_run
        FROM users u
        LEFT JOIN runs r ON r.user_id = u.id
        GROUP BY u.id
        ORDER BY u.created_at DESC
    """).fetchall()

    total_users    = len(users)
    active_users   = sum(1 for u in users if u["is_paid"] or
                         (u["trial_ends_at"] and u["trial_ends_at"] > datetime.utcnow().isoformat()))
    total_apps     = sum(u["total_applied"] for u in users)
    running_now    = db.execute("SELECT COUNT(*) FROM runs WHERE status='running'").fetchone()[0]

    rows = ""
    for u in users:
        trial_end = u["trial_ends_at"] or ""
        status_badge = (
            "<span style='color:#16a34a;font-weight:600'>Paid</span>" if u["is_paid"] else
            ("<span style='color:#d97706;font-weight:600'>Trial</span>"
             if trial_end > datetime.utcnow().isoformat() else
             "<span style='color:#dc2626;font-weight:600'>Expired</span>")
        )
        last_run_str = (u["last_run"] or "")[:16].replace("T", " ")
        rows += f"""
        <tr>
          <td>{u['id']}</td>
          <td><a href="/admin/user/{u['id']}" style="color:#4f46e5">{u['email']}</a></td>
          <td>{u['first_name']} {u['last_name']}</td>
          <td>{status_badge}</td>
          <td style="text-align:center">{u['total_runs']}</td>
          <td style="text-align:center">{u['total_applied']}</td>
          <td style="color:#6b7280;font-size:12px">{last_run_str}</td>
          <td style="color:#6b7280;font-size:12px">{(u['created_at'] or '')[:10]}</td>
        </tr>"""

    body = f"""
    <style>
      body {{ font-family: system-ui,sans-serif; background:#f9fafb; margin:0; padding:24px }}
      h1 {{ margin:0 0 24px; font-size:22px }}
      .stat-grid {{ display:grid; grid-template-columns:repeat(4,1fr); gap:16px; margin-bottom:28px }}
      .stat-card {{ background:#fff; border:1px solid #e5e7eb; border-radius:10px; padding:18px 20px }}
      .stat-card .num {{ font-size:28px; font-weight:700; color:#111 }}
      .stat-card .lbl {{ font-size:12px; color:#6b7280; margin-top:4px }}
      table {{ width:100%; border-collapse:collapse; background:#fff; border-radius:10px;
               overflow:hidden; box-shadow:0 1px 3px rgba(0,0,0,.07) }}
      th {{ background:#f3f4f6; text-align:left; padding:10px 14px; font-size:12px;
            color:#6b7280; text-transform:uppercase; letter-spacing:.5px }}
      td {{ padding:10px 14px; border-top:1px solid #f3f4f6; font-size:13px }}
      tr:hover td {{ background:#fafafa }}
    </style>
    <h1>ApplyExpress Admin</h1>
    <div class="stat-grid">
      <div class="stat-card"><div class="num">{total_users}</div><div class="lbl">Total users</div></div>
      <div class="stat-card"><div class="num">{active_users}</div><div class="lbl">Active (paid/trial)</div></div>
      <div class="stat-card"><div class="num">{total_apps}</div><div class="lbl">Applications submitted</div></div>
      <div class="stat-card"><div class="num">{running_now}</div><div class="lbl">Pipelines running now</div></div>
    </div>
    <table>
      <thead><tr>
        <th>ID</th><th>Email</th><th>Name</th><th>Status</th>
        <th>Runs</th><th>Applied</th><th>Last Run</th><th>Joined</th>
      </tr></thead>
      <tbody>{rows}</tbody>
    </table>
    <p style="margin-top:16px;font-size:12px;color:#9ca3af">
      <a href="/dashboard">← Back to dashboard</a>
    </p>
    """
    return _page("Admin — ApplyExpress", body)


@app.route("/admin/user/<int:uid>")
@_admin_required
def admin_user_detail(uid):
    db = get_db()
    u = db.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
    if not u:
        return _page("Not Found", "<p>User not found</p>"), 404

    u = dict(u)
    runs = db.execute(
        "SELECT * FROM runs WHERE user_id=? ORDER BY started_at DESC LIMIT 20", (uid,)
    ).fetchall()

    apps = db.execute(
        "SELECT * FROM applications WHERE user_id=? ORDER BY applied_at DESC LIMIT 30", (uid,)
    ).fetchall()

    run_rows = "".join(f"""
    <tr>
      <td>{r['run_id']}</td>
      <td>{r['status']}</td>
      <td>{r['jobs_found']}</td>
      <td>{r['jobs_applied']}</td>
      <td style='color:#6b7280;font-size:12px'>{(r['started_at'] or '')[:16].replace('T',' ')}</td>
    </tr>""" for r in runs)

    app_rows = "".join(f"""
    <tr>
      <td><a href="{a['url']}" target="_blank" style="color:#4f46e5">{a['title'][:40] if a['title'] else a['url'][:40]}</a></td>
      <td>{a['company'] or ''}</td>
      <td>{'✓' if a['status']=='applied' else '✗'}</td>
      <td style='color:#6b7280;font-size:12px'>{(a['applied_at'] or '')[:16].replace('T',' ')}</td>
    </tr>""" for a in apps)

    body = f"""
    <style>
      body {{ font-family:system-ui,sans-serif; background:#f9fafb; margin:0; padding:24px }}
      h2 {{ margin:0 0 6px }}
      table {{ width:100%; border-collapse:collapse; background:#fff; border-radius:10px;
               overflow:hidden; box-shadow:0 1px 3px rgba(0,0,0,.07); margin-bottom:28px }}
      th {{ background:#f3f4f6; text-align:left; padding:8px 14px; font-size:12px; color:#6b7280 }}
      td {{ padding:8px 14px; border-top:1px solid #f3f4f6; font-size:13px }}
    </style>
    <p><a href="/admin" style="color:#4f46e5">← All users</a></p>
    <h2>{u['first_name']} {u['last_name']}</h2>
    <p style="color:#6b7280;margin:0 0 20px">{u['email']} — joined {(u['created_at'] or '')[:10]}</p>

    <h3>Recent Runs</h3>
    <table><thead><tr><th>Run ID</th><th>Status</th><th>Found</th><th>Applied</th><th>Started</th></tr></thead>
    <tbody>{run_rows or '<tr><td colspan=5 style=color:#9ca3af>No runs yet</td></tr>'}</tbody></table>

    <h3>Applications</h3>
    <table><thead><tr><th>Job</th><th>Company</th><th>Result</th><th>Date</th></tr></thead>
    <tbody>{app_rows or '<tr><td colspan=4 style=color:#9ca3af>No applications yet</td></tr>'}</tbody></table>
    """
    return _page(f"Admin: {u['email']}", body)


# ── Scheduled pipeline (3× per 12 h = every 4 h) ─────────────────────────────
def _scheduled_pipeline_run():
    """Fire the pipeline for every active user. Runs inside a file-lock so that
    only one gunicorn worker acts when multiple workers share the same process."""
    lock_path = DATA_DIR / ".scheduler.lock"
    try:
        lf = open(str(lock_path), "w")
        fcntl.flock(lf, fcntl.LOCK_EX | fcntl.LOCK_NB)
    except (IOError, OSError):
        return  # another worker already holding the lock

    try:
        db = sqlite3.connect(str(DB_PATH))
        db.row_factory = sqlite3.Row
        import subprocess as _sp
        users = db.execute(
            "SELECT id FROM users WHERE keywords IS NOT NULL AND keywords != ''"
        ).fetchall()
        log.info(f"[scheduler] Firing pipeline for {len(users)} user(s)")
        seq_runner = Path(__file__).parent / "tools" / "run_sequential_runner.py"

        def _create_run(uid, spec_id=None):
            """Create a run record and return its db id, or None if already running."""
            if spec_id:
                already = db.execute(
                    "SELECT id FROM runs WHERE user_id=? AND specialty_id=? AND status='running'",
                    (uid, spec_id)
                ).fetchone()
            else:
                already = db.execute(
                    "SELECT id FROM runs WHERE user_id=? AND specialty_id IS NULL AND status='running'",
                    (uid,)
                ).fetchone()
            if already:
                log.info(f"[scheduler] user {uid} specialty={spec_id} already running — skip")
                return None
            run_id = datetime.utcnow().strftime("%Y%m%d_%H%M%S") + (f"_s{spec_id}" if spec_id else "")
            row = db.execute(
                "INSERT INTO runs (user_id, run_id, started_at, status, specialty_id) VALUES (?,?,?,?,?) RETURNING id",
                (uid, run_id, datetime.utcnow().isoformat(), "running", spec_id)
            ).fetchone()
            db.commit()
            return row["id"]

        for u in users:
            uid = u["id"]
            specialties = db.execute(
                "SELECT id FROM specialties WHERE user_id=?", (uid,)
            ).fetchall()
            # Create all run records upfront, then run sequentially
            pairs = []
            for spec_id in [None] + [s["id"] for s in specialties]:
                rid = _create_run(uid, spec_id=spec_id)
                if rid:
                    pairs.append(f"{rid}:{spec_id or ''}")
            if pairs:
                _sp.Popen(
                    [sys.executable, str(seq_runner), str(uid), str(DB_PATH), str(ROOT), str(DATA_DIR)]
                    + pairs,
                    start_new_session=True, stdout=_sp.DEVNULL, stderr=_sp.DEVNULL,
                )
                log.info(f"[scheduler] Sequential runner started for user {uid}: {len(pairs)} run(s)")

        db.close()
    except Exception as e:
        log.error(f"[scheduler] Error: {e}")
    finally:
        try:
            fcntl.flock(lf, fcntl.LOCK_UN)
            lf.close()
        except Exception:
            pass


def _start_scheduler():
    """Start background scheduler. Uses APScheduler if installed; falls back to a
    plain threading loop so scheduled runs work even without the package."""

    FIRE_HOURS = {7, 11, 15, 19}  # UTC hours to fire the pipeline

    # ── Try APScheduler first ──────────────────────────────────────────────────
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        scheduler = BackgroundScheduler(daemon=True)
        scheduler.add_job(_scheduled_pipeline_run, "cron", hour=",".join(map(str, FIRE_HOURS)),
                          minute=0, id="pipeline_cron", replace_existing=True)
        scheduler.start()
        log.info("[scheduler] APScheduler started — pipeline at %s UTC",
                 ", ".join(f"{h:02d}:00" for h in sorted(FIRE_HOURS)))
        return
    except ImportError:
        log.warning("[scheduler] APScheduler not installed — using threading fallback")
    except Exception as e:
        log.warning("[scheduler] APScheduler failed (%s) — using threading fallback", e)

    # ── Threading fallback — no external deps needed ───────────────────────────
    import threading
    import time as _time

    def _loop():
        last_fired = {}  # {user_id: (date, hour)} so we never double-fire
        while True:
            try:
                now = datetime.utcnow()
                if now.hour in FIRE_HOURS:
                    # Use DB to check: did any run start for this user in the last 90 min?
                    _scheduled_pipeline_run()
                _time.sleep(300)  # wake every 5 min — catch the start of each fire hour
            except Exception as exc:
                log.error("[scheduler-thread] Error: %s", exc)
                _time.sleep(60)

    t = threading.Thread(target=_loop, daemon=True, name="pipeline-scheduler")
    t.start()
    log.info("[scheduler] Threading scheduler started — pipeline at %s UTC",
             ", ".join(f"{h:02d}:00" for h in sorted(FIRE_HOURS)))


# ── Boot ──────────────────────────────────────────────────────────────────────
# init_db() runs at import time so gunicorn workers also create the schema
init_db()

# Mark any runs that were "running" when the server last died as interrupted
def _cleanup_orphaned_runs():
    try:
        db = sqlite3.connect(str(DB_PATH))
        db.execute(
            "UPDATE runs SET status='interrupted', finished_at=? WHERE status='running'",
            (datetime.utcnow().isoformat(),)
        )
        db.commit()
        db.close()
    except Exception:
        pass

_cleanup_orphaned_runs()
_start_scheduler()

if __name__ == "__main__":
    port = int(os.getenv("PORT", "5001"))
    log.info(f"ApplyExpress SaaS starting on http://localhost:{port}")
    app.run(host="0.0.0.0", port=port, debug=False)
